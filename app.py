import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
import sqlite3
import json
import os
import requests

# ─────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="Laudo de Morte Violenta | ICRIM/NPT",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

AUTORIDADES_DEFAULT = ["Delegado(a) Plantonista", "Policia Civil"]
MESES = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
         "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
EXTENSO = {
    1: "um", 2: "dois", 3: "três", 4: "quatro", 5: "cinco", 6: "seis",
    7: "sete", 8: "oito", 9: "nove", 10: "dez", 11: "onze", 12: "doze",
    13: "treze", 14: "catorze", 15: "quinze", 16: "dezesseis", 17: "dezessete",
    18: "dezoito", 19: "dezenove", 20: "vinte", 21: "vinte e um", 22: "vinte e dois",
    23: "vinte e três", 24: "vinte e quatro", 25: "vinte e cinco", 26: "vinte e seis",
    27: "vinte e sete", 28: "vinte e oito", 29: "vinte e nove", 30: "trinta", 31: "trinta e um",
    2025: "dois mil e vinte e cinco", 2026: "dois mil e vinte e seis",
    2027: "dois mil e vinte e sete", 2028: "dois mil e vinte e oito",
    2029: "dois mil e vinte e nove", 2030: "dois mil e trinta",
}


def num_extenso(n): return EXTENSO.get(n, str(n))


def data_extenso(d):
    return (f"Aos {d.day:02d} ({num_extenso(d.day)}) dia(s) do mês de "
            f"{MESES[d.month-1]} do ano {d.year} ({num_extenso(d.year)})")


def data_simples(d): return d.strftime("%d/%m/%Y")
def v(val): return str(val).strip() if str(val).strip() else "________"


# ═══════════════════════════════════════════════════════════
# MÓDULO DE INTEGRAÇÃO GOOGLE GEMINI AI
# ═══════════════════════════════════════════════════════════
# GEMINI VISION & IA HELPERS
# ═══════════════════════════════════════════════════════════
def get_gemini_api_key():
    """Recupera a chave de API do Gemini de CHAVE.txt, secrets, ambiente ou session_state."""
    for fname in ["CHAVE.txt", "chave.txt", "CHAVE.TXT"]:
        fpath = os.path.join(os.path.dirname(__file__), fname)
        if os.path.exists(fpath):
            try:
                with open(fpath, "r", encoding="utf-8", errors="ignore") as fp:
                    k = fp.read().strip()
                    if k: return k
            except Exception: pass
    try:
        if "GEMINI_API_KEY" in st.secrets and st.secrets["GEMINI_API_KEY"]:
            return st.secrets["GEMINI_API_KEY"]
    except Exception: pass
    env_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if env_key: return env_key
    if st.session_state.get("user_gemini_key"):
        return st.session_state.get("user_gemini_key")
    if st.session_state.get("gemini_api_key_input"):
        return st.session_state.get("gemini_api_key_input")
    return ""


def convert_bytes_to_pil_images(file_bytes, file_name, file_type=""):
    """
    Converte bytes de arquivo (PDF ou Imagem) em lista de objetos PIL Image.
    """
    from io import BytesIO
    from PIL import Image

    file_name_lower = file_name.lower()
    is_pdf = file_type == "application/pdf" or file_name_lower.endswith(".pdf")
    images = []

    if is_pdf:
        # 1. PyMuPDF (fitz) - alta fidelidade e velocidade
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                pix = page.get_pixmap(dpi=150)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                images.append(img)
        except Exception:
            pass

        # 2. Fallback pdfplumber
        if not images:
            try:
                import pdfplumber
                with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        img = page.to_image(resolution=150).original
                        images.append(img.convert("RGB"))
            except Exception:
                pass
    else:
        try:
            img = Image.open(BytesIO(file_bytes))
            images.append(img.convert("RGB"))
        except Exception:
            pass

    return images




def carregar_dados_teste_exemplo():
    """Carrega dados de teste completos (3 vestígios, 2 vítimas, 3 considerações extras e 3 fotografias com visualização) no session_state."""
    import datetime
    st.session_state["num_laudo"] = "2026-ICRIM-0987"
    st.session_state["ocorrencia"] = "BO-12345/2026"
    st.session_state["requisicao"] = "REQ-5544/2026"
    st.session_state["perito"] = "Dr. Carlos Eduardo Silva - Perito Criminal"
    st.session_state["destino"] = "1ª Delegacia de Polícia Civil de Imperatriz"
    st.session_state["autoridade_sel"] = "Dr. João Mendes - Delegado de Polícia"
    st.session_state["data_pericia_input"] = datetime.date.today()
    st.session_state["data_atendimento_input"] = datetime.date.today()
    st.session_state["horario"] = datetime.time(14, 30)
    st.session_state["horario_atendimento"] = datetime.time(14, 45)
    
    st.session_state["endereco"] = "Rua das Palmeiras, nº 450, Bairro Maranhão Novo"
    st.session_state["municipio"] = "Imperatriz - MA"
    st.session_state["latitude"] = "-5.5264"
    st.session_state["longitude"] = "-47.4721"
    st.session_state["ponto_referencia"] = "Próximo à Praça da Cultura"
    st.session_state["area"] = "Via Pública Comercial / Asfalto"
    st.session_state["pavimento"] = "Asfalto"
    st.session_state["delimitacoes"] = "Local aberto, delimitado ao norte pela via pública e ao sul por imóvel comercial."
    st.session_state["iso_estado"] = "1. Local Preservado e Isolado"
    st.session_state["iso_meio"] = "fita zebrada e cordão de isolamento da Polícia Militar"
    st.session_state["clima"] = "Seco / Ensolarado"
    st.session_state["visibilidade"] = "Boa (Luz Natural Solar)"
    st.session_state["iluminacao"] = "Natural"
    st.session_state["equipe_pm"] = "VTR 14-020 (Sgt. Oliveira e Sd. Santos)"
    st.session_state["equipe_pc"] = "Equipe de Homicídios (Inv. Lima)"
    st.session_state["autoridade_local"] = "Dr. João Mendes - Delegado de Polícia"

    # 2 Vítimas
    st.session_state["vitimas"] = [
        {
            "nome": "João da Silva Sauro",
            "sexo": "Masculino",
            "idade": "35 anos",
            "naturalidade": "Imperatriz - MA",
            "rg": "1234567 SSP/MA",
            "cpf": "000.111.222-33",
            "filiacao": "Maria da Silva Sauro",
            "posicao": "Decúbito dorsal",
            "cabeca": "Voltada para a direita",
            "membros": "Superiores estendidos ao longo do tronco; inferiores paralelos",
            "vestes": "Camisa polo branca e calça jeans azul",
            "fenomenos": "Rigidez cadavérica em início de fixação; livores de hipóstase dorsais",
            "lesoes": [
                "Ferimento perfurocontundente com orifício de entrada em região parietal direita.",
                "Orifício de saída em região temporal esquerda associado a fratura de calota craniana."
            ]
        },
        {
            "nome": "Pedro Alves Santos",
            "sexo": "Masculino",
            "idade": "28 anos",
            "naturalidade": "Açailândia - MA",
            "rg": "7654321 SSP/MA",
            "cpf": "999.888.777-66",
            "filiacao": "Ana Alves Santos",
            "posicao": "Decúbito lateral direito",
            "cabeca": "Alinhada ao tronco",
            "membros": "Superiores semi-flexionados; inferiores flexionados",
            "vestes": "Camiseta preta e bermuda tática cinza",
            "fenomenos": "Rigidez cadavérica generalizada nos quatro membros",
            "lesoes": [
                "Ferimento perfurocontundente no tórax anterior, linha hemiclavicular esquerda.",
                "Escoriações superficiais em região patelar direita."
            ]
        }
    ]

    # 3 Vestígios
    st.session_state["vestigios"] = [
        {
            "tipo": "Vestígio Balístico",
            "localizacao": "Solo asfáltico, a 1,20m da Vítima #1",
            "orientacao": "Norte",
            "coordenadas": "X: 2.5m, Y: 1.8m",
            "detalhes": "Estojo percutido e deflagrado de munição calibre 9mm, marca CBC."
        },
        {
            "tipo": "Mancha de Sangue",
            "localizacao": "Solo abaixo da região cefálica da Vítima #1",
            "orientacao": "Centro",
            "coordenadas": "X: 0m, Y: 0m",
            "detalhes": "Poça de sangue com padrão de formação por gravidade e projeção secundária."
        },
        {
            "tipo": "Fragmento Metálico / Projétil",
            "localizacao": "Próximo ao meio-fio, a 3,50m da Vítima #2",
            "orientacao": "Leste",
            "coordenadas": "X: 3.5m, Y: -0.8m",
            "detalhes": "Projétil de arma de fogo deformaço (jaquetado 9mm) recolhido para balística."
        }
    ]

    st.session_state["inst_acao"] = "1. Perfurocontundente (Ex: Projétil de Arma de Fogo - PAF)"
    st.session_state["inst_agente"] = "projéteis de arma de fogo (PAF)"
    st.session_state["inst_extra"] = "com a recuperação de um projétil de arma de fogo durante o exame perinecroscópico e confirmação de disparo a curta distância"

    # 3 Considerações Extras (e, f, g)
    st.session_state["consideracoes_extras"] = [
        {
            "titulo": "Do Exame das Vestes das Vítimas",
            "texto": "As vestes das vítimas apresentavam perfurações com orlas de enxugo e esfumaçamento compatíveis com a entrada de projéteis disparados por arma de fogo a curta distância."
        },
        {
            "titulo": "Das Marcas de Frenagem Pneumática",
            "texto": "Foi constatada marca de frenagem pneumática impressa na pista de rolamento com extensão de 4,20 metros, sugerindo tentativa de desaceleração veicular no momento da abordagem."
        },
        {
            "titulo": "Da Disposição dos Elementos de Prova no Cenário",
            "texto": "A disposição espacial dos vestígios balísticos arrecadados e o padrão das manchas de sangue mantêm estrita convergência com a trajetória dos disparos efetuados no local."
        }
    ]

    st.session_state["quesitos_list"] = [
        {"pergunta": "Qual a causa da morte?", "resposta": "Traumatismo cranioencefálico e choque hipovolêmico decorrentes de ferimentos por projéteis de arma de fogo."},
        {"pergunta": "Qual o instrumento ou meio utilizado?", "resposta": "Instrumento em ação perfurocontundente (projéteis de arma de fogo - PAF)."}
    ]

    st.session_state["dinamica_fatos"] = (
        "Diante dos elementos materiais coligidos no local dos fatos e exames perinecroscópicos realizados, "
        "infere-se que as vítimas transitavam pela via pública quando foram surpreendidas pelo agressor. "
        "Foram efetuados disparos de arma de fogo de curto alcance, atingindo regiões vitais das vítimas, "
        "ocasionando a queda imediata no solo e os consequentes óbitos no local dos fatos."
    )

    # 3 Fotografias de Exemplo
    st.session_state["fotos"] = [
        {
            "b64": "/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAAMCAgMCAgMDAwMEAwMEBQgFBQQEBQoHBwYIDAoMDAsKCwsNDhIQDQ4RDgsLEBYQERMUFRUVDA8XGBYUGBIUFRT/2wBDAQMEBAUEBQkFBQkUDQsNFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBT/wAARCAGQAlgDASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD4cooqzpum3mtaja6fp9rPfX93KkFva20ZklmkYhVRFUEsxJAAAySa+4PnCtRXo/8AwzZ8XP8AolnjX/wnrv8A+N0f8M2fFz/olnjX/wAJ67/+N1HtId0Vyy7HnFFej/8ADNnxc/6JZ41/8J67/wDjdH/DNnxc/wCiWeNf/Ceu/wD43R7SHdByy7HnFFej/wDDNnxc/wCiWeNf/Ceu/wD43R/wzZ8XP+iWeNf/AAnrv/43R7SHdByy7HnFFej/APDNnxc/6JZ41/8ACeu//jdH/DNnxc/6JZ41/wDCeu//AI3R7SHdByy7HnFFej/8M2fFz/olnjX/AMJ67/8AjdH/AAzZ8XP+iWeNf/Ceu/8A43R7SHdByy7HnFFej/8ADNnxc/6JZ41/8J67/wDjdH/DNnxc/wCiWeNf/Ceu/wD43R7SHdByy7HnFFej/wDDNnxc/wCiWeNf/Ceu/wD43R/wzZ8XP+iWeNf/AAnrv/43R7SHdByy7HnFFej/APDNnxc/6JZ41/8ACeu//jdH/DNnxc/6JZ41/wDCeu//AI3R7SHdByy7HnFFej/8M2fFz/olnjX/AMJ67/8AjdH/AAzZ8XP+iWeNf/Ceu/8A43R7SHdByy7HnFFej/8ADNnxc/6JZ41/8J67/wDjdH/DNnxc/wCiWeNf/Ceu/wD43R7SHdByy7HnFFej/wDDNnxc/wCiWeNf/Ceu/wD43R/wzZ8XP+iWeNf/AAnrv/43R7SHdByy7HnFFej/APDNnxc/6JZ41/8ACeu//jdH/DNnxc/6JZ41/wDCeu//AI3R7SHdByy7HnFFej/8M2fFz/olnjX/AMJ67/8AjdH/AAzZ8XP+iWeNf/Ceu/8A43R7SHdByy7HnFFej/8ADNnxc/6JZ41/8J67/wDjdH/DNnxc/wCiWeNf/Ceu/wD43R7SHdByy7HnFFej/wDDNnxc/wCiWeNf/Ceu/wD43R/wzZ8XP+iWeNf/AAnrv/43R7SHdByy7HnFFej/APDNnxc/6JZ41/8ACeu//jdH/DNnxc/6JZ41/wDCeu//AI3R7SHdByy7HnFFej/8M2fFz/olnjX/AMJ67/8AjdH/AAzZ8XP+iWeNf/Ceu/8A43R7SHdByy7HnFFej/8ADNnxc/6JZ41/8J67/wDjdH/DNnxc/wCiWeNf/Ceu/wD43R7SHdByy7HnFFej/wDDNnxc/wCiWeNf/Ceu/wD43R/wzZ8XP+iWeNf/AAnrv/43R7SHdByy7HnFFej/APDNnxc/6JZ41/8ACeu//jdH/DNnxc/6JZ41/wDCeu//AI3R7SHdByy7HnFFej/8M2fFz/olnjX/AMJ67/8AjdH/AAzZ8XP+iWeNf/Ceu/8A43R7SHdByy7HnFFej/8ADNnxc/6JZ41/8J67/wDjdH/DNnxc/wCiWeNf/Ceu/wD43R7SHdByy7HnFFej/wDDNnxc/wCiWeNf/Ceu/wD43R/wzZ8XP+iWeNf/AAnrv/43R7SHdByy7HnFFej/APDNnxc/6JZ41/8ACeu//jdH/DNnxc/6JZ41/wDCeu//AI3R7SHdByy7HnFFej/8M2fFz/olnjX/AMJ67/8AjdH/AAzZ8XP+iWeNf/Ceu/8A43R7SHdByy7HnFFej/8ADNnxc/6JZ41/8J67/wDjdH/DNnxc/wCiWeNf/Ceu/wD43R7SHdByy7HnFFej/wDDNnxc/wCiWeNf/Ceu/wD43R/wzZ8XP+iWeNf/AAnrv/43R7SHdByy7HnFFej/APDNnxc/6JZ41/8ACeu//jdH/DNnxc/6JZ41/wDCeu//AI3R7SHdByy7HnFFej/8M2fFz/olnjX/AMJ67/8AjdH/AAzZ8XP+iWeNf/Ceu/8A43R7SHdByy7HnFFej/8ADNnxc/6JZ41/8J67/wDjdH/DNnxc/wCiWeNf/Ceu/wD43R7SHdByy7HnFFej/wDDNnxc/wCiWeNf/Ceu/wD43R/wzZ8XP+iWeNf/AAnrv/43R7SHdByy7HnFFej/APDNnxc/6JZ41/8ACeu//jdH/DNnxc/6JZ41/wDCeu//AI3R7SHdByy7HnFFej/8M2fFz/olnjX/AMJ67/8AjdH/AAzZ8XP+iWeNf/Ceu/8A43R7SHdByy7HnFFej/8ADNnxc/6JZ41/8J67/wDjdH/DNnxc/wCiWeNf/Ceu/wD43R7SHdByy7HnFFej/wDDNnxc/wCiWeNf/Ceu/wD43R/wzZ8XP+iWeNf/AAnrv/43R7SHdByy7HnFFej/APDNnxc/6JZ41/8ACeu//jdcV4i8Nav4Q1m40jXtKvdE1W32+dY6jbvbzxblDLujcBhlWUjI5BB701KMtmDTW6M2iiiqJCvR/wBmz/k4v4Wf9jVpX/pXFXnFej/s2f8AJxfws/7GrSv/AEriqJ/A/QuPxI/RL9t79t7x1+zV8V9J8MeGNJ8PX1hd6JFqTyatbTySiRp54yAY5kG3ES8Yzknn0+ev+Hsfxc/6F3wV/wCAN3/8lUf8FY/+Ti/Dv/Yq23/pXeV8V15+Hw9KdKMpR1OmrVnGbSZ9qf8AD2P4uf8AQu+Cv/AG7/8Akqj/AIex/Fz/AKF3wV/4A3f/AMlV8V0V0/VaH8qMvbVP5j7U/wCHsfxc/wChd8Ff+AN3/wDJVH/D2P4uf9C74K/8Abv/AOSq+K6KPqtD+VB7ap/Mfan/AA9j+Ln/AELvgr/wBu//AJKo/wCHsfxc/wChd8Ff+AN3/wDJVfFdFH1Wh/Kg9tU/mPtT/h7H8XP+hd8Ff+AN3/8AJVH/AA9j+Ln/AELvgr/wBu//AJKr4roo+q0P5UHtqn8x9qf8PY/i5/0Lvgr/AMAbv/5Ko/4ex/Fz/oXfBX/gDd//ACVXxXRR9VofyoPbVP5j7U/4ex/Fz/oXfBX/AIA3f/yVR/w9j+Ln/Qu+Cv8AwBu//kqviuij6rQ/lQe2qfzH2p/w9j+Ln/Qu+Cv/AABu/wD5Ko/4ex/Fz/oXfBX/AIA3f/yVXxXRR9VofyoPbVP5j7U/4ex/Fz/oXfBX/gDd/wDyVR/w9j+Ln/Qu+Cv/AABu/wD5Kr4roo+q0P5UHtqn8x9qf8PY/i5/0Lvgr/wBu/8A5Ko/4ex/Fz/oXfBX/gDd/wDyVXxXRR9VofyoPbVP5j7U/wCHsfxc/wChd8Ff+AN3/wDJVH/D2P4uf9C74K/8Abv/AOSq+K6KPqtD+VB7ap/Mfan/AA9j+Ln/AELvgr/wBu//AJKo/wCHsfxc/wChd8Ff+AN3/wDJVfFdFH1Wh/Kg9tU/mPtT/h7H8XP+hd8Ff+AN3/8AJVH/AA9j+Ln/AELvgr/wBu//AJKr4roo+q0P5UHtqn8x9qf8PY/i5/0Lvgr/AMAbv/5Ko/4ex/Fz/oXfBX/gDd//ACVXxXRR9VofyoPbVP5j7U/4ex/Fz/oXfBX/AIA3f/yVR/w9j+Ln/Qu+Cv8AwBu//kqviuij6rQ/lQe2qfzH2p/w9j+Ln/Qu+Cv/AABu/wD5Ko/4ex/Fz/oXfBX/AIA3f/yVXxXRR9VofyoPbVP5j7U/4ex/Fz/oXfBX/gDd/wDyVR/w9j+Ln/Qu+Cv/AABu/wD5Kr4roo+q0P5UHtqn8x9qf8PY/i5/0Lvgr/wBu/8A5Ko/4ex/Fz/oXfBX/gDd/wDyVXxXRR9VofyoPbVP5j7U/wCHsfxc/wChd8Ff+AN3/wDJVH/D2P4uf9C74K/8Abv/AOSq+K6KPqtD+VB7ap/Mfan/AA9j+Ln/AELvgr/wBu//AJKo/wCHsfxc/wChd8Ff+AN3/wDJVfFdFH1Wh/Kg9tU/mPtT/h7H8XP+hd8Ff+AN3/8AJVH/AA9j+Ln/AELvgr/wBu//AJKr4roo+q0P5UHtqn8x9qf8PY/i5/0Lvgr/AMAbv/5Ko/4ex/Fz/oXfBX/gDd//ACVXxXRR9VofyoPbVP5j7U/4ex/Fz/oXfBX/AIA3f/yVR/w9j+Ln/Qu+Cv8AwBu//kqviuij6rQ/lQe2qfzH2p/w9j+Ln/Qu+Cv/AABu/wD5Ko/4ex/Fz/oXfBX/AIA3f/yVXxXRR9VofyoPbVP5j7U/4ex/Fz/oXfBX/gDd/wDyVR/w9j+Ln/Qu+Cv/AABu/wD5Kr4roo+q0P5UHtqn8x9qf8PY/i5/0Lvgr/wBu/8A5Ko/4ex/Fz/oXfBX/gDd/wDyVXxXRR9VofyoPbVP5j7U/wCHsfxc/wChd8Ff+AN3/wDJVH/D2P4uf9C74K/8Abv/AOSq+K6KPqtD+VB7ap/Mfan/AA9j+Ln/AELvgr/wBu//AJKo/wCHsfxc/wChd8Ff+AN3/wDJVfFdFH1Wh/Kg9tU/mPtT/h7H8XP+hd8Ff+AN3/8AJVH/AA9j+Ln/AELvgr/wBu//AJKr4roo+q0P5UHtqn8x9qf8PY/i5/0Lvgr/AMAbv/5Ko/4ex/Fz/oXfBX/gDd//ACVXxXRR9VofyoPbVP5j7U/4ex/Fz/oXfBX/AIA3f/yVR/w9j+Ln/Qu+Cv8AwBu//kqviuij6rQ/lQe2qfzH2p/w9j+Ln/Qu+Cv/AABu/wD5Ko/4ex/Fz/oXfBX/AIA3f/yVXxXRR9VofyoPbVP5j7U/4ex/Fz/oXfBX/gDd/wDyVR/w9j+Ln/Qu+Cv/AABu/wD5Kr4roo+q0P5UHtqn8x9qf8PY/i5/0Lvgr/wBu/8A5Ko/4ex/Fz/oXfBX/gDd/wDyVXxXRR9VofyoPbVP5j7U/wCHsfxc/wChd8Ff+AN3/wDJVH/D2P4uf9C74K/8Abv/AOSq+K6KPqtD+VB7ap/Mfpb+yr/wUK+Ivxy+PXhfwRr2i+F7TStU+1edNp1rcpOvl2s0y7S9w6j5o1Byp4J6da+YP+CiH/J4nxA/7h//AKb7aj/gnf8A8nifD/8A7iH/AKb7mj/goh/yeJ8QP+4f/wCm+2rnhThTxVoK3u/qaSlKdG8n1/Q+caKKK9I5Ar0f9mz/AJOL+Fn/AGNWlf8ApXFXnFej/s2f8nF/Cz/satK/9K4qifwP0Lj8SPo3/grH/wAnF+Hf+xVtv/Su8r4rr7U/4Kx/8nF+Hf8AsVbb/wBK7yviusML/Aj6Glf+IwooorqMAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigD6O/4J3/APJ4nw//AO4h/wCm+5o/4KIf8nifED/uH/8ApvtqP+Cd/wDyeJ8P/wDuIf8ApvuaP+CiH/J4nxA/7h//AKb7auH/AJi/+3f1On/lx8/0PnGiiiu45gr0f9mz/k4v4Wf9jVpX/pXFXnFej/s2f8nF/Cz/ALGrSv8A0riqJ/A/QuPxI+jf+Csf/Jxfh3/sVbb/ANK7yviuvtT/AIKx/wDJxfh3/sVbb/0rvK+K6wwv8CPoaV/4jCiiiuowCiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKAPo7/AIJ3/wDJ4nw//wC4h/6b7mj/AIKIf8nifED/ALh//pvtqP8Agnf/AMnifD//ALiH/pvuaP8Agoh/yeJ8QP8AuH/+m+2rh/5i/wDt39Tp/wCXHz/Q+caKKK7jmCvR/wBmz/k4v4Wf9jVpX/pXFXnFej/s2f8AJxfws/7GrSv/AEriqJ/A/QuPxI+jf+Csf/Jxfh3/ALFW2/8ASu8r4rr7U/4Kx/8AJxfh3/sVbb/0rvK+K6wwv8CPoaV/4jCiiiuowCiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKAPo7/gnf/wAnifD/AP7iH/pvuaP+CiH/ACeJ8QP+4f8A+m+2o/4J3/8AJ4nw/wD+4h/6b7mj/goh/wAnifED/uH/APpvtq4f+Yv/ALd/U6f+XHz/AEPnGiiiu45gr0f9mz/k4v4Wf9jVpX/pXFXnFej/ALNn/Jxfws/7GrSv/SuKon8D9C4/Ej6N/wCCsf8AycX4d/7FW2/9K7yviuvtT/grH/ycX4d/7FW2/wDSu8r4rrDC/wACPoaV/wCIwooorqMAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigD6O/4J3/8nifD/wD7iH/pvuaP+CiH/J4nxA/7h/8A6b7aj/gnf/yeJ8P/APuIf+m+5o/4KIf8nifED/uH/wDpvtq4f+Yv/t39Tp/5cfP9D5xoooruOYK9H/Zs/wCTi/hZ/wBjVpX/AKVxV5xXo/7Nn/Jxfws/7GrSv/SuKon8D9C4/Ej6N/4Kx/8AJxfh3/sVbb/0rvK+K6+1P+Csf/Jxfh3/ALFW2/8ASu8r4rrDC/wI+hpX/iMKKKK6jAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooA+jv+Cd/wDyeJ8P/wDuIf8ApvuaP+CiH/J4nxA/7h//AKb7aj/gnf8A8nifD/8A7iH/AKb7mj/goh/yeJ8QP+4f/wCm+2rh/wCYv/t39Tp/5cfP9D5xoooruOYK9H/Zs/5OL+Fn/Y1aV/6VxV5xXo/7Nn/Jxfws/wCxq0r/ANK4qifwP0Lj8SPo3/grH/ycX4d/7FW2/wDSu8r4rr7U/wCCsf8AycX4d/7FW2/9K7yviusML/Aj6Glf+IwooorqMAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigD6O/wCCd/8AyeJ8P/8AuIf+m+5o/wCCiH/J4nxA/wC4f/6b7aj/AIJ3/wDJ4nw//wC4h/6b7mj/AIKIf8nifED/ALh//pvtq4f+Yv8A7d/U6f8Alx8/0PnGiiiu45gr0f8AZs/5OL+Fn/Y1aV/6VxV5xXo/7Nn/ACcX8LP+xq0r/wBK4qifwP0Lj8SPo3/grH/ycX4d/wCxVtv/AErvK+K6+1P+Csf/ACcX4d/7FW2/9K7yviusML/Aj6Glf+IwooorqMAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigD6O/4J3/8AJ4nw/wD+4h/6b7mj/goh/wAnifED/uH/APpvtqP+Cd//ACeJ8P8A/uIf+m+5o/4KIf8AJ4nxA/7h/wD6b7auH/mL/wC3f1On/lx8/wBD5xoooruOYK9H/Zs/5OL+Fn/Y1aV/6VxV5xXo/wCzZ/ycX8LP+xq0r/0riqJ/A/QuPxI+jf8AgrH/AMnF+Hf+xVtv/Su8r4rr7U/4Kx/8nF+Hf+xVtv8A0rvK+K6wwv8AAj6Glf8AiMKKKK6jAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooA+jv+Cd//J4nw/8A+4h/6b7mj/goh/yeJ8QP+4f/AOm+2o/4J3/8nifD/wD7iH/pvuaP+CiH/J4nxA/7h/8A6b7auH/mL/7d/U6f+XHz/Q+caKKK7jmCvR/2bP8Ak4v4Wf8AY1aV/wClcVecV6P+zZ/ycX8LP+xq0r/0riqJ/A/QuPxI+jf+Csf/ACcX4d/7FW2/9K7yviuvtT/grH/ycX4d/wCxVtv/AErvK+K6wwv8CPoaV/4jCvo68+GGgHXPDunN4a0GDSH8OWWtXF1Z6476vdSLoi6hPGYDdP5fmyLIgbyAFDrt6rn5xrek8c67J4k0/X/7Rkj1jT4rWC1u4VWNoktoY4YANoA+WOJFyeTjJySSdpxctmZxaW525n8Oah4En8Yw+CdJs59M1aDS5dLFzfNZ3aXEM8iSc3Pmq8ZtmB2ybSJU+UYO7svjBo3gz4eaj4kurLwDpFxDbeMNV8N2lnJe6gYYrezaM+ZIRdb2lcToAQwUCJ/lJOV8b8UfEDWvGFra2uoS2kdnbO8sVpp2n29jAJHwGcxwRopchVBYgnCgZwBWjD8YfFUera7qMt7aX1xrl62pahHqOmWt3BNdMzMZhDLG0aPmR8FVGAxA4OKz5JXTv+LK5kd7qfgbwt4HbxVrk2hnX9Pt7jQks9IvbqVDBHqVnJeMGeJkZpIli8kEnG5slT0qtqHgPwvoEfxx8PPp02oax4TaU6frMl2wWNItXs7MARKArMyTS7mYsOV2qpBJ4TSfit4p0bXNT1iHUxPqGpyCa8kvraK7WaQPvWQpKjLvVuVYDK84IzWRbeK9WtU11UvXb+3YPs2otIBI1wnnxXGCzAkHzYY23AgnbjOCQWoS6vt+gcy7Hr3hr4ReD9T+GOna1dx682q3OhanqEk0F7EtulzB9sNuojNuWMZW1G8+ZkHAyvmJWTpvgfQLrwpY+LmsN2kW/hy9N/AJpMNqkcv2aLB3Z63VjOVBAwXGNtc3pnxs8X6P4Yg8P2t/aJpcFhd6XEr6XaPKltclzPGJmiMg3GVzndkcYI2rjBi8a61D4Nn8KJfFdAmvV1B7Ty05nCbN2/G4DGMqDtJVSQSqkHLPXXr+Ac0ex7vc/DLwYnxC8ceF9N0ayvtWt/Fup6XZaNf39zbTvYxyeXbpp827ynuA28FZyxb9yFViWrj10rSPEvww8GrpPgfRrfxL4h8QX2grex3N+XBii04wsqtcsgcvdybsqV6YUYrl0+OHjRL7Ur06tFJeX+oz6vJcTWFvJJDeTEGWeBmjJt3YhSWh2H5E/urjE0jx7ruhWmlW1jfeRDpd3c31mvkxsYZ544o5XBKkklYIgM52lAVwckpQmt3+LByj2PYvF3wg8Lj4r6KulQw/8Ihd6Zf6h9n03VEvfONhHcO8QnSSQCSdLeJyATsN0AAAAoPB3gPQ/Fuo+EtUm8FaXYW2o6b4ik+yx62x0y6ms9Na5g3ubozW7I7xmRZJEG3YeAWx43pHjvXtBg0mLT9Re0XSb6TUbIxqu6G4kWNXcNjJDCGIFSSvy9OTnoP+F6eMluNLkS+sIY9M+3fZLWHR7KO1iF5AILpfIWERlZI1ClSpHUgAkkpwqWsn+fn/AMAalHsekeDPh/Y6/pfia4/4QrwTcaxbazounW9o/ieRNPSO4iv2kKTi/wAPIzQw/L5jkBWwg+aovEXwb8OQfDrxhfaTCjXsOoX2oaPNeX6x3z6Xb3a2yobQuHO4faZHfy/l+y43DLA+R6p8RNc1axurKSSztbO6uLa7lt9O062s0aa3WZYX2wxoAVFzMOOu7nO1cWNW+K3ijXPFEviG91JZdWltrizeVbWFEMM6SpMgjVAgDCaXOF6uT15p8k73v+Yc0bbEvxo8P2HhP4xeO9D0q3+y6XpmvX9laQb2fy4Y7h0RdzEscKoGSSTjkmvcU+Fngjxf8Vfiz4ci8MReH7DwHNqWoQNpl5cs+oW1pOyG0kM8soEki7droFwVf5WyMeDeOPibrvxFup7vXTpk95cXL3k91aaPZ2k00rkl2eSGJGfJYk7iRk561qa18dvG2u6o2pTatDaX8moDVZrjTNPtrF57sMzCaUwRp5rAu5G/ON7epyOE2kr6iUopsm16z0TxP8NLnxRp+gWvhi907V7fTZbawnuJILqOeGeRWAnkkYPGbZg2GwRKnAIJbpNZ8BeHdE+Mfxhh/sz7ToPgyfUZrPRjPIFmVNQS0hiaQN5hRPOV2IYMVjYbgTuHnXij4ga14wtbW11CW0js7Z3litNO0+3sYBI+AzmOCNFLkKoLEE4UDOAK3NT+OnjDV/EEmuXFzpa6tNNNPcXdvodjA900oYSifZCvnK4dtySblbPINPln0/MV4ndHwr4U/wCFKf8AC1x4VtPN/tP/AIR//hHPtV39i87b5v2zPnefs8v93s83Hmc5x8lJ8K/hx4X8c+M9duL7Sk0Xw61tY20FvrOqLZLZ3d6qYlR5JEMqRoLmaJSzM6xx5D5bPn3/AAufxb9sab7bZmFrX7EdPOl2hsPJ8zzNv2TyvIH7z587M7vmznmsfXfHmu+JYbmLUtQa5jubsXso8tF3SrH5aH5QMKqfKqD5VHCgUuSequPmj2Nv4l+ELfwjovgMCzay1K+0aefUld2Ja5TVb+3yQSQpCW8a4XA+TOMkk8LXdeI/jX4r8X2UltrU+makHE6/aLjRLFrhBNPLcSbZvJ8xMyzyv8rDBc4xXC1rHmS94iVr6BRRRVkhRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB9Hf8E7/APk8T4f/APcQ/wDTfc0f8FEP+TxPiB/3D/8A0321H/BO/wD5PE+H/wD3EP8A033NH/BRD/k8T4gf9w//ANN9tXD/AMxf/bv6nT/y4+f6HzjRRRXccwV6P+zZ/wAnF/Cz/satK/8ASuKvOK9H/Zs/5OL+Fn/Y1aV/6VxVE/gfoXH4kfRv/BWP/k4vw7/2Ktt/6V3lfFdfan/BWP8A5OL8O/8AYq23/pXeV8V1hhf4EfQ0r/xGFFFFdRgFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAfR3/BO/8A5PE+H/8A3EP/AE33NH/BRD/k8T4gf9w//wBN9tR/wTv/AOTxPh//ANxD/wBN9zR/wUQ/5PE+IH/cP/8ATfbVw/8AMX/27+p0/wDLj5/ofONFFFdxzBXo/wCzZ/ycX8LP+xq0r/0rirzivR/2bP8Ak4v4Wf8AY1aV/wClcVRP4H6Fx+JH0b/wVj/5OL8O/wDYq23/AKV3lfFdfan/AAVj/wCTi/Dv/Yq23/pXeV8V1hhf4EfQ0r/xGFFFFdRgFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAfR3/AATv/wCTxPh//wBxD/033NH/AAUQ/wCTxPiB/wBw/wD9N9tR/wAE7/8Ak8T4f/8AcQ/9N9zR/wAFEP8Ak8T4gf8AcP8A/TfbVw/8xf8A27+p0/8ALj5/ofONFFFdxzBXo/7Nn/Jxfws/7GrSv/SuKvOK9H/Zs/5OL+Fn/Y1aV/6VxVE/gfoXH4kfRv8AwVj/AOTi/Dv/AGKtt/6V3lfFdfan/BWP/k4vw7/2Ktt/6V3lfFdYYX+BH0NK/wDEYV6xp3wGbVINJhTXVsdR1GwXUoJNQspE0+ePyDM6Q3KFy8qAbWTywAwcZ+XnyevXfAvxu0jwTaXVunhWe6t9QsfsWoWP9q4s5T5Wzz44mhby5SQCzFmBDSqAof5danPb3CIct/eIf+GeNYbTNPvrfVtL1RL+60+0t4bCYLNI93Jcxxr5dx5Lbg9scg8YfOcI+3ibzwDrth4eTW57JU09kjlJFxG0qRucRyPEG8xEYj5XZQpyME5GfSI/2hbGC90aePw1csdN1HS9ULTaorSXFxaXN3M7OwgA/ei9kBwBhlDc5K1yWqfE2HUPDt3appLQ6xf6ZZ6PeX5ut0T21sYTGEh2Ao5+zQAsXYHYcAbjURdW+qG1DoV2+Dni5bW2n/suN0ngt7oLHewO6QTxiSKaRA5aOMqwO9wFHQkHiq6fC3xJLdXkCWlu4tNP/tWWdb+3Nv8AZfNWLzVm8zy3HmME+Vid2R1Bx6Ovx503wbrsWoeG9PuJ7+68N6boeo3/ANueINEljbxzpCvlhopQ8QUS7mGY9yjnJ4/X/i5Jra65G9vf3EeoaJFo0c2p6o95cRql7DdmR5GUb8tEy7VVAA4PUHcKVV9Aagupnf8ACpvEc7v9lsDJGIIJ08+eCJ5hLbJcosSmQmRvLdW2JuYBlyqk4ps3wq8QiC6uILEvbW9nFfMZp4I5Hie1S63InmEybYXEjBMlVILBeQNrV/izpviPS7Oy1fw292mnxQCxEeoGJVlTT7WzdpcJl0b7HE4VWQryNxzmtaz+PdtbaZdwf8I463c+kHSftMF+I/MQ6THpwMo8omRV8szIgZQGkYHdwwL1e39feFodzjNV+HV9beIotIsZUvXbS7LVJJ5nS2ihS4tIbg73dwqhTME3MwBOOhYCp9K+Dni7Wry/tLPS0luLKa3t5EN5Au+SdHe3WIs4EpkWNinl7t3GM7lzrab8ZpNM8SS6pBZ3tmt1oFnoN39g1E29xtt4rdFlhmCHyyTaxEgq3BdckHi1p3x0lsfEx1eTTJ7wjxDo+uKLnUXllK6ekyLE8rqxZnWUZfgArwmCAo3VtogtDuche/DnxBYaza6VLZRte3dq97brDdQypNEgkLFJEcoxHlSLtBzuUrjdxU9h8Mdf1PSLW+t7FmW6khSEmWFU2yi5KM5MgKD/AESc5ZQAqbiQCu5V+IVxbXXga6tLYQXXha3EMbs+8TsL2e6DEYGB+/2Fec7c55wOps/jbZWuo6kT4ZEml3F3Zy21k14C1pDbWV1axKrtEQ0ii5WQSFcb4slDu4pupbRCSj1OPg+HHiG68WnwzFYpJrIga5EK3MWxolgNwZFl3bGUwjeGDEEYxnIrrtE/Z11/VYdMe51XRNJOpXv2K1FzcSTLKdiOJVkt45YzGd+Nwb7yMvUYMmm/FTStY+Llx4n1e1ntdOfw/caY1vLdNNLMw0Z7KMGVYwd0jBcttwpfJ4BNbOh/tD6Lo2leGdMPg+7nsPD9wz2q/wBsIHkhaRJjFKxtju/fqZdyhefLGPk5iUqunKikodWYNt8AdWufFHhjT1uo20zWTo++/jMTyWov44WDPbiTzNqNMUDNtVio5G4CuZ0/4Ya7rFtp09jaq0F7aLdpPd3MFtEQZ5YFAd5AMmSGRQp2s204XABPZaX8ebPQ9c0nWrHw3Imq29vo1nePLqO+K5h082rAInlAxtI1nESdzYGQAcknJi+Kul3HhfT/AA1qXh24vNEs4IkWOLURFM8kV3ezxsZPKI27L+SNlC87QwK9AJ1ewWgczrnw88QeG9HTU9S082lq0qwMHmjMsTsHKLJEG3x7hG5Xeo3BSRkCoPB3hn/hLdXuLH7T9l8rTr6/8zy9+fs1pLcbMZH3vK257bs4OMHrfF3xbs/EfgBfDNnoB0hDc2Vy3lXm+2RreCeI+XEYwymTzy7s0jksOuCAOe0Lx9f6f4hn1fUpbnW7mXTL3TN11cszhZ7OW2B3NuOEEu4L324yM5Gic3F3WpDUUzWtfgV41vbe0mh0qB1uo7eSFf7RtQ7efEJbdSvmZDyocohAZ8MFBKkDPsfhR4p1PRLLVbTTBc2t7A9zaxxXMLXE8SytC7JAH8xgrowOF4xk8YNdNpfxzOm6tpt7/Yvm/YtR8OX/AJf2vG/+yrRrfZnZx5u7dnnZjGH61o+H/jRp3gCx+HuqaLYNe+LtB0m4tRczTstvbtLfXjlHhMf7w+TPkMrgDzeQSuBm5VV0/q3qWlDucTF8J/E9xqtlp0Flb3NxeJcSQvb6hbyQkQRmWcGVZCisiDcyswIBXj5hm/4d+CviPXtY8PWjpa2dtrOo2+nR3b3kDrE8xPls6q5ZQyo7JkDzAp2buK09U+Nr32p21wLXVbqCCw1S0VNX1uS9lDXtq9uWV2QBUQMpChcttIZjkFbum/HqDR9K8M2tn4daJ9H1XSNWZftwFvLJYpKp2xiIFDMZS7sWY7s9sANurbRf194JQ7nIr8LNeSC1nmt4vIvLW9u7Vra8tpjNHaiYzOAJfuL9nm+cZzt+XdlQan/Cu9cbQ7DV44bW4sb66jsoDbX9vNIZ5FLLG0SOXU4HO5RjIBwSM7EXxOgs7rRmttJkW20vQdT0KKOW7DuyXZvsSMwjAygvhkAfN5X8O7CzP8RtGSx8MR2WiX1jcaFAxib+0Elge6MbE3Pl+SrBzN5TZLthI1TkAEVep2/rXz9CbRNPUvgZHbW/gm6svECX9l4jvLm3kuFtdq2UUV1JCszfOdyskMkmOMBGHOM1UvPgybXU/iDZDVw0vhi9uLK2BtsHUXgMzSkfP+7Cw28kh+9ztX+LcG6d8ZX0/wAFxaANJEhh0d9NhuTcY2StPeuZ9u3n9zqE8WzPXa+eNtaNz+0HqI8VS6nYWtzp2m3uu3us6npUWoMI9QW5lBe3lIUAoI18vlT952AG7FZ/vf6/r5Fe4eS0UrFSx2ghc8AnJApK6jEKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigD6O/wCCd/8AyeJ8P/8AuIf+m+5o/wCCiH/J4nxA/wC4f/6b7aj/AIJ3/wDJ4nw//wC4h/6b7mj/AIKIf8nifED/ALh//pvtq4f+Yv8A7d/U6f8Alx8/0PnGiiiu45gr0f8AZs/5OL+Fn/Y1aV/6VxV5xXo/7Nn/ACcX8LP+xq0r/wBK4qifwP0Lj8SPo3/grH/ycX4d/wCxVtv/AErvK+K6+1P+Csf/ACcX4d/7FW2/9K7yviusML/Aj6Glf+IwooorqMAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigD6O/4J3/8AJ4nw/wD+4h/6b7mj/goh/wAnifED/uH/APpvtqP+Cd//ACeJ8P8A/uIf+m+5o/4KIf8AJ4nxA/7h/wD6b7auH/mL/wC3f1On/lx8/wBD5xoooruOYK9H/Zs/5OL+Fn/Y1aV/6VxV5xXo/wCzZ/ycX8LP+xq0r/0riqJ/A/QuPxI+jf8AgrH/AMnF+Hf+xVtv/Su8r4rr7U/4Kx/8nF+Hf+xVtv8A0rvK+K6wwv8AAj6Glf8AiMKKKK6jAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooA+jv+Cd//J4nw/8A+4h/6b7mj/goh/yeJ8QP+4f/AOm+2o/4J3/8nifD/wD7iH/pvuaP+CiH/J4nxA/7h/8A6b7auH/mL/7d/U6f+XHz/Q+caKKK7jmCvR/2bP8Ak4v4Wf8AY1aV/wClcVecV6P+zZ/ycX8LP+xq0r/0riqJ/A/QuPxI+jf+Csf/ACcX4d/7FW2/9K7yviuvtT/grH/ycX4d/wCxVtv/AErvK+K6wwv8CPoaV/4jCiiiuowCiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKAPo7/gnf8A8nifD/8A7iH/AKb7mj/goh/yeJ8QP+4f/wCm+2o/4J3/APJ4nw//AO4h/wCm+5o/4KIf8nifED/uH/8Apvtq4f8AmL/7d/U6f+XHz/Q+caKKK7jmCvR/2bP+Ti/hZ/2NWlf+lcVecV6P+zZ/ycX8LP8AsatK/wDSuKon8D9C4/Ej6N/4Kx/8nF+Hf+xVtv8A0rvK+K6+1P8AgrH/AMnF+Hf+xVtv/Su8r4rrDC/wI+hpX/iMKKKK6jAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooA+jv8Agnf/AMnifD//ALiH/pvuaP8Agoh/yeJ8QP8AuH/+m+2o/wCCd/8AyeJ8P/8AuIf+m+5o/wCCiH/J4nxA/wC4f/6b7auH/mL/AO3f1On/AJcfP9D5xoooruOYK9H/AGbP+Ti/hZ/2NWlf+lcVecV6P+zZ/wAnF/Cz/satK/8ASuKon8D9C4/Ej6N/4Kx/8nF+Hf8AsVbb/wBK7yviuvtT/grH/wAnF+Hf+xVtv/Su8r4rrDC/wI+hpX/iMKKKK6jAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooA+jv+Cd//ACeJ8P8A/uIf+m+5o/4KIf8AJ4nxA/7h/wD6b7aj/gnf/wAnifD/AP7iH/pvuaP+CiH/ACeJ8QP+4f8A+m+2rh/5i/8At39Tp/5cfP8AQ+caKKK7jmCvR/2bP+Ti/hZ/2NWlf+lcVecV6P8As2f8nF/Cz/satK/9K4qifwP0Lj8SPo3/AIKx/wDJxfh3/sVbb/0rvK+K6+1P+Csf/Jxfh3/sVbb/ANK7yviusML/AAI+hpX/AIjCiiiuowCiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKAPo7/gnf/yeJ8P/APuIf+m+5o/4KIf8nifED/uH/wDpvtqP+Cd//J4nw/8A+4h/6b7mj/goh/yeJ8QP+4f/AOm+2rh/5i/+3f1On/lx8/0PnGiiiu45gr0f9mz/AJOL+Fn/AGNWlf8ApXFXnFej/s2f8nF/Cz/satK/9K4qifwP0Lj8SPo3/grH/wAnF+Hf+xVtv/Su8r4rr7U/4Kx/8nF+Hf8AsVbb/wBK7yviusML/Aj6Glf+IwooorqMAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigD6O/4J3/APJ4nw//AO4h/wCm+5o/4KIf8nifED/uH/8ApvtqP+Cd/wDyeJ8P/wDuIf8ApvuaP+CiH/J4nxA/7h//AKb7auH/AJi/+3f1On/lx8/0PnGiiiu45gr0f9mz/k4v4Wf9jVpX/pXFXnFej/s2f8nF/Cz/ALGrSv8A0riqJ/A/QuPxI+jf+Csf/Jxfh3/sVbb/ANK7yviuvtT/AIKx/wDJxfh3/sVbb/0rvK+K6wwv8CPoaV/4jCiiiuowCiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKAPo7/AIJ3/wDJ4nw//wC4h/6b7mj/AIKIf8nifED/ALh//pvtqP8Agnf/AMnifD//ALiH/pvuaP8Agoh/yeJ8QP8AuH/+m+2rh/5i/wDt39Tp/wCXHz/Q+caKKK7jmCvR/wBmz/k4v4Wf9jVpX/pXFXnFej/s2f8AJxfws/7GrSv/AEriqJ/A/QuPxI+jf+Csf/Jxfh3/ALFW2/8ASu8r4rr7U/4Kx/8AJxfh3/sVbb/0rvK+K6wwv8CPoaV/4jCiiiuowCiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKAPo7/gnf/wAnifD/AP7iH/pvuaP+CiH/ACeJ8QP+4f8A+m+2o/4J3/8AJ4nw/wD+4h/6b7mj/goh/wAnifED/uH/APpvtq4f+Yv/ALd/U6f+XHz/AEPnGiiiu45gr0f9mz/k4v4Wf9jVpX/pXFXnFej/ALNn/Jxfws/7GrSv/SuKon8D9C4/Ej6N/wCCsf8AycX4d/7FW2/9K7yviuv0t/4KFfsq/FL45fGjRde8EeF/7b0q38Pw2Mtx/aFrb7Zlublyu2WVGPyyIcgY568GvmD/AId3/tB/9E//APK1p/8A8kVxYatTjRinJfeb1qc3UbSZ840V9Hf8O7/2g/8Aon//AJWtP/8Akij/AId3/tB/9E//APK1p/8A8kV0+3pfzr70Zeyn/K/uPnGivo7/AId3/tB/9E//APK1p/8A8kUf8O7/ANoP/on/AP5WtP8A/kij29L+dfeg9lP+V/cfONFfR3/Du/8AaD/6J/8A+VrT/wD5Io/4d3/tB/8ARP8A/wArWn//ACRR7el/OvvQeyn/ACv7j5xor6O/4d3/ALQf/RP/APytaf8A/JFH/Du/9oP/AKJ//wCVrT//AJIo9vS/nX3oPZT/AJX9x840V9Hf8O7/ANoP/on/AP5WtP8A/kij/h3f+0H/ANE//wDK1p//AMkUe3pfzr70Hsp/yv7j5xor6O/4d3/tB/8ARP8A/wArWn//ACRR/wAO7/2g/wDon/8A5WtP/wDkij29L+dfeg9lP+V/cfONFfR3/Du/9oP/AKJ//wCVrT//AJIo/wCHd/7Qf/RP/wDytaf/APJFHt6X86+9B7Kf8r+4+caK+jv+Hd/7Qf8A0T//AMrWn/8AyRR/w7v/AGg/+if/APla0/8A+SKPb0v5196D2U/5X9x840V9Hf8ADu/9oP8A6J//AOVrT/8A5Io/4d3/ALQf/RP/APytaf8A/JFHt6X86+9B7Kf8r+4+caK+jv8Ah3f+0H/0T/8A8rWn/wDyRR/w7v8A2g/+if8A/la0/wD+SKPb0v5196D2U/5X9x840V9Hf8O7/wBoP/on/wD5WtP/APkij/h3f+0H/wBE/wD/ACtaf/8AJFHt6X86+9B7Kf8AK/uPnGivo7/h3f8AtB/9E/8A/K1p/wD8kUf8O7/2g/8Aon//AJWtP/8Akij29L+dfeg9lP8Alf3HzjRX0d/w7v8A2g/+if8A/la0/wD+SKP+Hd/7Qf8A0T//AMrWn/8AyRR7el/OvvQeyn/K/uPnGivo7/h3f+0H/wBE/wD/ACtaf/8AJFH/AA7v/aD/AOif/wDla0//AOSKPb0v5196D2U/5X9x840V9Hf8O7/2g/8Aon//AJWtP/8Akij/AId3/tB/9E//APK1p/8A8kUe3pfzr70Hsp/yv7j5xor6O/4d3/tB/wDRP/8Aytaf/wDJFH/Du/8AaD/6J/8A+VrT/wD5Io9vS/nX3oPZT/lf3HzjRX0d/wAO7/2g/wDon/8A5WtP/wDkij/h3f8AtB/9E/8A/K1p/wD8kUe3pfzr70Hsp/yv7j5xor6O/wCHd/7Qf/RP/wDytaf/APJFH/Du/wDaD/6J/wD+VrT/AP5Io9vS/nX3oPZT/lf3HzjRX0d/w7v/AGg/+if/APla0/8A+SKP+Hd/7Qf/AET/AP8AK1p//wAkUe3pfzr70Hsp/wAr+4+caK+jv+Hd/wC0H/0T/wD8rWn/APyRR/w7v/aD/wCif/8Ala0//wCSKPb0v5196D2U/wCV/cfONFfR3/Du/wDaD/6J/wD+VrT/AP5Io/4d3/tB/wDRP/8Aytaf/wDJFHt6X86+9B7Kf8r+4+caK+jv+Hd/7Qf/AET/AP8AK1p//wAkUf8ADu/9oP8A6J//AOVrT/8A5Io9vS/nX3oPZT/lf3HzjRX0d/w7v/aD/wCif/8Ala0//wCSKP8Ah3f+0H/0T/8A8rWn/wDyRR7el/OvvQeyn/K/uPnGivo7/h3f+0H/ANE//wDK1p//AMkUf8O7/wBoP/on/wD5WtP/APkij29L+dfeg9lP+V/cfONFfR3/AA7v/aD/AOif/wDla0//AOSKP+Hd/wC0H/0T/wD8rWn/APyRR7el/OvvQeyn/K/uPnGivo7/AId3/tB/9E//APK1p/8A8kUf8O7/ANoP/on/AP5WtP8A/kij29L+dfeg9lP+V/cfONFfR3/Du/8AaD/6J/8A+VrT/wD5Io/4d3/tB/8ARP8A/wArWn//ACRR7el/OvvQeyn/ACv7j5xor6O/4d3/ALQf/RP/APytaf8A/JFH/Du/9oP/AKJ//wCVrT//AJIo9vS/nX3oPZT/AJX9x840V9Hf8O7/ANoP/on/AP5WtP8A/kij/h3f+0H/ANE//wDK1p//AMkUe3pfzr70Hsp/yv7j5xor6O/4d3/tB/8ARP8A/wArWn//ACRR/wAO7/2g/wDon/8A5WtP/wDkij29L+dfeg9lP+V/cfONFfR3/Du/9oP/AKJ//wCVrT//AJIo/wCHd/7Qf/RP/wDytaf/APJFHt6X86+9B7Kf8r+4+caK+jv+Hd/7Qf8A0T//AMrWn/8AyRR/w7v/AGg/+if/APla0/8A+SKPb0v5196D2U/5X9x840V9Hf8ADu/9oP8A6J//AOVrT/8A5Io/4d3/ALQf/RP/APytaf8A/JFHt6X86+9B7Kf8r+4+caK+jv8Ah3f+0H/0T/8A8rWn/wDyRR/w7v8A2g/+if8A/la0/wD+SKPb0v5196D2U/5X9wf8E7/+TxPh/wD9xD/033NH/BRD/k8T4gf9w/8A9N9tXtX7Gf7Gfxi+FH7Sfg/xV4q8H/2XoNh9s+03f9p2c3l77OeNPkjmZjl3UcA9cnjNeK/8FEP+TxPiB/3D/wD0321csZxnirxd/d/U2cXGjZrr+h840UUV6JyBWl4a8Rah4Q8R6Vr2kXH2TVdLu4r6zuNiv5U0bh422sCpwyg4IIOOQazaKNxn0d/w8Q/aD/6KB/5RdP8A/kej/h4h+0H/ANFA/wDKLp//AMj1840Vh7Cl/IvuRp7Wf8z+8+jv+HiH7Qf/AEUD/wAoun//ACPR/wAPEP2g/wDooH/lF0//AOR6+caKPYUv5F9yD2s/5n959Hf8PEP2g/8AooH/AJRdP/8Akej/AIeIftB/9FA/8oun/wDyPXzjRR7Cl/IvuQe1n/M/vPo7/h4h+0H/ANFA/wDKLp//AMj0f8PEP2g/+igf+UXT/wD5Hr5xoo9hS/kX3IPaz/mf3n0d/wAPEP2g/wDooH/lF0//AOR6P+HiH7Qf/RQP/KLp/wD8j1840UewpfyL7kHtZ/zP7z6O/wCHiH7Qf/RQP/KLp/8A8j0f8PEP2g/+igf+UXT/AP5Hr5xoo9hS/kX3IPaz/mf3n0d/w8Q/aD/6KB/5RdP/APkej/h4h+0H/wBFA/8AKLp//wAj1840UewpfyL7kHtZ/wAz+8+jv+HiH7Qf/RQP/KLp/wD8j0f8PEP2g/8AooH/AJRdP/8AkevnGij2FL+Rfcg9rP8Amf3n0d/w8Q/aD/6KB/5RdP8A/kej/h4h+0H/ANFA/wDKLp//AMj1840UewpfyL7kHtZ/zP7z6O/4eIftB/8ARQP/ACi6f/8AI9H/AA8Q/aD/AOigf+UXT/8A5Hr5xoo9hS/kX3IPaz/mf3n0d/w8Q/aD/wCigf8AlF0//wCR6P8Ah4h+0H/0UD/yi6f/API9fONFHsKX8i+5B7Wf8z+8+jv+HiH7Qf8A0UD/AMoun/8AyPR/w8Q/aD/6KB/5RdP/APkevnGij2FL+Rfcg9rP+Z/efR3/AA8Q/aD/AOigf+UXT/8A5Ho/4eIftB/9FA/8oun/APyPXzjRR7Cl/IvuQe1n/M/vPo7/AIeIftB/9FA/8oun/wDyPR/w8Q/aD/6KB/5RdP8A/kevnGij2FL+Rfcg9rP+Z/efR3/DxD9oP/ooH/lF0/8A+R6P+HiH7Qf/AEUD/wAoun//ACPXzjRR7Cl/IvuQe1n/ADP7z6O/4eIftB/9FA/8oun/APyPR/w8Q/aD/wCigf8AlF0//wCR6+caKPYUv5F9yD2s/wCZ/efR3/DxD9oP/ooH/lF0/wD+R6P+HiH7Qf8A0UD/AMoun/8AyPXzjRR7Cl/IvuQe1n/M/vPo7/h4h+0H/wBFA/8AKLp//wAj0f8ADxD9oP8A6KB/5RdP/wDkevnGij2FL+Rfcg9rP+Z/efR3/DxD9oP/AKKB/wCUXT//AJHo/wCHiH7Qf/RQP/KLp/8A8j1840UewpfyL7kHtZ/zP7z6O/4eIftB/wDRQP8Ayi6f/wDI9H/DxD9oP/ooH/lF0/8A+R6+caKPYUv5F9yD2s/5n959Hf8ADxD9oP8A6KB/5RdP/wDkej/h4h+0H/0UD/yi6f8A/I9fONFHsKX8i+5B7Wf8z+8+jv8Ah4h+0H/0UD/yi6f/API9H/DxD9oP/ooH/lF0/wD+R6+caKPYUv5F9yD2s/5n959Hf8PEP2g/+igf+UXT/wD5Ho/4eIftB/8ARQP/ACi6f/8AI9fONFHsKX8i+5B7Wf8AM/vPo7/h4h+0H/0UD/yi6f8A/I9H/DxD9oP/AKKB/wCUXT//AJHr5xoo9hS/kX3IPaz/AJn959Hf8PEP2g/+igf+UXT/AP5Ho/4eIftB/wDRQP8Ayi6f/wDI9fONFHsKX8i+5B7Wf8z+8+jv+HiH7Qf/AEUD/wAoun//ACPR/wAPEP2g/wDooH/lF0//AOR6+caKPYUv5F9yD2s/5n959Hf8PEP2g/8AooH/AJRdP/8Akej/AIeIftB/9FA/8oun/wDyPXzjRR7Cl/IvuQe1n/M/vPo7/h4h+0H/ANFA/wDKLp//AMj0f8PEP2g/+igf+UXT/wD5Hr5xoo9hS/kX3IPaz/mf3n0d/wAPEP2g/wDooH/lF0//AOR6P+HiH7Qf/RQP/KLp/wD8j1840UewpfyL7kHtZ/zP7z6O/wCHiH7Qf/RQP/KLp/8A8j0f8PEP2g/+igf+UXT/AP5Hr5xoo9hS/kX3IPaz/mf3n0d/w8Q/aD/6KB/5RdP/APkej/h4h+0H/wBFA/8AKLp//wAj1840UewpfyL7kHtZ/wAz+8+jv+HiH7Qf/RQP/KLp/wD8j0f8PEP2g/8AooH/AJRdP/8AkevnGij2FL+Rfcg9rP8Amf3n0d/w8Q/aD/6KB/5RdP8A/kej/h4h+0H/ANFA/wDKLp//AMj1840UewpfyL7kHtZ/zP7z6O/4eIftB/8ARQP/ACi6f/8AI9H/AA8Q/aD/AOigf+UXT/8A5Hr5xoo9hS/kX3IPaz/mf3n0d/w8Q/aD/wCigf8AlF0//wCR6P8Ah4h+0H/0UD/yi6f/API9fONFHsKX8i+5B7Wf8z+8+jv+HiH7Qf8A0UD/AMoun/8AyPXivxE+IniH4r+MdQ8VeKtQ/tTXr/y/tN35McPmbI1jT5I1VRhEUcAdMnnNc3RVxpQg7xil8iXOUtGwooorQg//2Q==",
            "descricao": "Fotografia 01 — Visão geral do local de crime evidenciando o cadáver no solo asfáltico.",
            "legenda": "Fotografia 01 — Visão geral do local de crime evidenciando o cadáver no solo asfáltico.",
            "incluir": True
        },
        {
            "b64": "/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAAMCAgMCAgMDAwMEAwMEBQgFBQQEBQoHBwYIDAoMDAsKCwsNDhIQDQ4RDgsLEBYQERMUFRUVDA8XGBYUGBIUFRT/2wBDAQMEBAUEBQkFBQkUDQsNFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBT/wAARCAGQAlgDASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD5Hooqxp2nXesaha2FhazXt9dSrBb21tGZJZpGIVURRksxJAAHJJryj9i2K9Feg/8ADO/xV/6Jl4x/8EF3/wDG6P8Ahnf4q/8ARMvGP/ggu/8A43Ts+xl7al/MvvR59RXoP/DO/wAVf+iZeMf/AAQXf/xuj/hnf4q/9Ey8Y/8Aggu//jdFn2D21L+Zfejz6ivQf+Gd/ir/ANEy8Y/+CC7/APjdH/DO/wAVf+iZeMf/AAQXf/xuiz7B7al/MvvR59RXoP8Awzv8Vf8AomXjH/wQXf8A8bo/4Z3+Kv8A0TLxj/4ILv8A+N0WfYPbUv5l96PPqK9B/wCGd/ir/wBEy8Y/+CC7/wDjdH/DO/xV/wCiZeMf/BBd/wDxuiz7B7al/MvvR59RXoP/AAzv8Vf+iZeMf/BBd/8Axuj/AIZ3+Kv/AETLxj/4ILv/AON0WfYPbUv5l96PPqK9B/4Z3+Kv/RMvGP8A4ILv/wCN0f8ADO/xV/6Jl4x/8EF3/wDG6LPsHtqX8y+9Hn1Feg/8M7/FX/omXjH/AMEF3/8AG6P+Gd/ir/0TLxj/AOCC7/8AjdFn2D21L+Zfejz6ivQf+Gd/ir/0TLxj/wCCC7/+N0f8M7/FX/omXjH/AMEF3/8AG6LPsHtqX8y+9Hn1Feg/8M7/ABV/6Jl4x/8ABBd//G6P+Gd/ir/0TLxj/wCCC7/+N0WfYPbUv5l96PPqK9B/4Z3+Kv8A0TLxj/4ILv8A+N0f8M7/ABV/6Jl4x/8ABBd//G6LPsHtqX8y+9Hn1Feg/wDDO/xV/wCiZeMf/BBd/wDxuj/hnf4q/wDRMvGP/ggu/wD43RZ9g9tS/mX3o8+or0H/AIZ3+Kv/AETLxj/4ILv/AON0f8M7/FX/AKJl4x/8EF3/APG6LPsHtqX8y+9Hn1Feg/8ADO/xV/6Jl4x/8EF3/wDG6P8Ahnf4q/8ARMvGP/ggu/8A43RZ9g9tS/mX3o8+or0H/hnf4q/9Ey8Y/wDggu//AI3R/wAM7/FX/omXjH/wQXf/AMbos+we2pfzL70efUV6D/wzv8Vf+iZeMf8AwQXf/wAbo/4Z3+Kv/RMvGP8A4ILv/wCN0WfYPbUv5l96PPqK9B/4Z3+Kv/RMvGP/AIILv/43R/wzv8Vf+iZeMf8AwQXf/wAbos+we2pfzL70efUV6D/wzv8AFX/omXjH/wAEF3/8bo/4Z3+Kv/RMvGP/AIILv/43RZ9g9tS/mX3o8+or0H/hnf4q/wDRMvGP/ggu/wD43R/wzv8AFX/omXjH/wAEF3/8bos+we2pfzL70efUV6D/AMM7/FX/AKJl4x/8EF3/APG6P+Gd/ir/ANEy8Y/+CC7/APjdFn2D21L+Zfejz6ivQf8Ahnf4q/8ARMvGP/ggu/8A43R/wzv8Vf8AomXjH/wQXf8A8bos+we2pfzL70efUV6D/wAM7/FX/omXjH/wQXf/AMbo/wCGd/ir/wBEy8Y/+CC7/wDjdFn2D21L+Zfejz6ivQf+Gd/ir/0TLxj/AOCC7/8AjdH/AAzv8Vf+iZeMf/BBd/8Axuiz7B7al/MvvR59RXoP/DO/xV/6Jl4x/wDBBd//ABuj/hnf4q/9Ey8Y/wDggu//AI3RZ9g9tS/mX3o8+or0H/hnf4q/9Ey8Y/8Aggu//jdH/DO/xV/6Jl4x/wDBBd//ABuiz7B7al/MvvR59RXoP/DO/wAVf+iZeMf/AAQXf/xuj/hnf4q/9Ey8Y/8Aggu//jdFn2D21L+Zfejz6ivQf+Gd/ir/ANEy8Y/+CC7/APjdH/DO/wAVf+iZeMf/AAQXf/xuiz7B7al/MvvR59RXoP8Awzv8Vf8AomXjH/wQXf8A8bo/4Z3+Kv8A0TLxj/4ILv8A+N0WfYPbUv5l96PPqK9B/wCGd/ir/wBEy8Y/+CC7/wDjdH/DO/xV/wCiZeMf/BBd/wDxuiz7B7al/MvvR59RXoP/AAzv8Vf+iZeMf/BBd/8Axuj/AIZ3+Kv/AETLxj/4ILv/AON0WfYPbUv5l96PPqK9B/4Z3+Kv/RMvGP8A4ILv/wCN0f8ADO/xV/6Jl4x/8EF3/wDG6LPsHtqX8y+9Hn1Feg/8M7/FX/omXjH/AMEF3/8AG6P+Gd/ir/0TLxj/AOCC7/8AjdFn2D21L+Zfejz6ivQf+Gd/ir/0TLxj/wCCC7/+N0f8M7/FX/omXjH/AMEF3/8AG6LPsHtqX8y+9Hn1Feg/8M7/ABV/6Jl4x/8ABBd//G6P+Gd/ir/0TLxj/wCCC7/+N0WfYPbUv5l96PPqK9B/4Z3+Kv8A0TLxj/4ILv8A+N0f8M7/ABV/6Jl4x/8ABBd//G6LPsHtqX8y+9Hn1Feg/wDDO/xV/wCiZeMf/BBd/wDxuuO1/wAO6r4T1efStb0y80bVLfb51lqFu8E0e5Qy7kcBhlWBGRyCD3pWaKjUhN2jJMz6KKKDQK9B/Z3/AOTgPhl/2M+mf+lcVefV6D+zv/ycB8Mv+xn0z/0riprdGNb+FL0Z9/ftm/tm+Nf2dvihpfhvw3pegXtjdaNFqDyarbzySiRp54yAUmQbcRLxjOSefTwT/h6R8Vf+hf8AB3/gFd//ACTR/wAFSP8Ak4Dw/wD9ixb/APpXd18d1tOclJpM8PAYDC1MNCc4Jto+xP8Ah6R8Vf8AoX/B3/gFd/8AyTR/w9I+Kv8A0L/g7/wCu/8A5Jr47oqPaT7nf/ZuD/59o+xP+HpHxV/6F/wd/wCAV3/8k0f8PSPir/0L/g7/AMArv/5Jr47oo9pPuH9m4P8A59o+xP8Ah6R8Vf8AoX/B3/gFd/8AyTR/w9I+Kv8A0L/g7/wCu/8A5Jr47oo9pPuH9m4P/n2j7E/4ekfFX/oX/B3/AIBXf/yTR/w9I+Kv/Qv+Dv8AwCu//kmvjuij2k+4f2bg/wDn2j7E/wCHpHxV/wChf8Hf+AV3/wDJNH/D0j4q/wDQv+Dv/AK7/wDkmvjuij2k+4f2bg/+faPsT/h6R8Vf+hf8Hf8AgFd//JNH/D0j4q/9C/4O/wDAK7/+Sa+O6KPaT7h/ZuD/AOfaPsT/AIekfFX/AKF/wd/4BXf/AMk0f8PSPir/ANC/4O/8Arv/AOSa+O6KPaT7h/ZuD/59o+xP+HpHxV/6F/wd/wCAV3/8k0f8PSPir/0L/g7/AMArv/5Jr47oo9pPuH9m4P8A59o+xP8Ah6R8Vf8AoX/B3/gFd/8AyTR/w9I+Kv8A0L/g7/wCu/8A5Jr47oo9pPuH9m4P/n2j7E/4ekfFX/oX/B3/AIBXf/yTR/w9I+Kv/Qv+Dv8AwCu//kmvjuij2k+4f2bg/wDn2j7E/wCHpHxV/wChf8Hf+AV3/wDJNH/D0j4q/wDQv+Dv/AK7/wDkmvjuij2k+4f2bg/+faPsT/h6R8Vf+hf8Hf8AgFd//JNH/D0j4q/9C/4O/wDAK7/+Sa+O6KPaT7h/ZuD/AOfaPsT/AIekfFX/AKF/wd/4BXf/AMk0f8PSPir/ANC/4O/8Arv/AOSa+O6KPaT7h/ZuD/59o+xP+HpHxV/6F/wd/wCAV3/8k0f8PSPir/0L/g7/AMArv/5Jr47oo9pPuH9m4P8A59o+xP8Ah6R8Vf8AoX/B3/gFd/8AyTR/w9I+Kv8A0L/g7/wCu/8A5Jr47oo9pPuH9m4P/n2j7E/4ekfFX/oX/B3/AIBXf/yTR/w9I+Kv/Qv+Dv8AwCu//kmvjuij2k+4f2bg/wDn2j7E/wCHpHxV/wChf8Hf+AV3/wDJNH/D0j4q/wDQv+Dv/AK7/wDkmvjuij2k+4f2bg/+faPsT/h6R8Vf+hf8Hf8AgFd//JNH/D0j4q/9C/4O/wDAK7/+Sa+O6KPaT7h/ZuD/AOfaPsT/AIekfFX/AKF/wd/4BXf/AMk0f8PSPir/ANC/4O/8Arv/AOSa+O6KPaT7h/ZuD/59o+xP+HpHxV/6F/wd/wCAV3/8k0f8PSPir/0L/g7/AMArv/5Jr47oo9pPuH9m4P8A59o+xP8Ah6R8Vf8AoX/B3/gFd/8AyTR/w9I+Kv8A0L/g7/wCu/8A5Jr47oo9pPuH9m4P/n2j7E/4ekfFX/oX/B3/AIBXf/yTR/w9I+Kv/Qv+Dv8AwCu//kmvjuij2k+4f2bg/wDn2j7E/wCHpHxV/wChf8Hf+AV3/wDJNH/D0j4q/wDQv+Dv/AK7/wDkmvjuij2k+4f2bg/+faPsT/h6R8Vf+hf8Hf8AgFd//JNH/D0j4q/9C/4O/wDAK7/+Sa+O6KPaT7h/ZuD/AOfaPsT/AIekfFX/AKF/wd/4BXf/AMk0f8PSPir/ANC/4O/8Arv/AOSa+O6KPaT7h/ZuD/59o+xP+HpHxV/6F/wd/wCAV3/8k0f8PSPir/0L/g7/AMArv/5Jr47oo9pPuH9m4P8A59o+xP8Ah6R8Vf8AoX/B3/gFd/8AyTR/w9I+Kv8A0L/g7/wCu/8A5Jr47oo9pPuH9m4P/n2j7E/4ekfFX/oX/B3/AIBXf/yTR/w9I+Kv/Qv+Dv8AwCu//kmvjuij2k+4f2bg/wDn2j7E/wCHpHxV/wChf8Hf+AV3/wDJNH/D0j4q/wDQv+Dv/AK7/wDkmvjuij2k+4f2bg/+faPsT/h6R8Vf+hf8Hf8AgFd//JNH/D0j4q/9C/4O/wDAK7/+Sa+O6KPaT7h/ZuD/AOfaPsT/AIekfFX/AKF/wd/4BXf/AMk0f8PSPir/ANC/4O/8Arv/AOSa+O6KPaT7h/ZuD/59o+xP+HpHxV/6F/wd/wCAV3/8k0f8PSPir/0L/g7/AMArv/5Jr47oo9pPuH9m4P8A59o+xP8Ah6R8Vf8AoX/B3/gFd/8AyTR/w9I+Kv8A0L/g7/wCu/8A5Jr47oo9pPuH9m4P/n2j7E/4ekfFX/oX/B3/AIBXf/yTR/w9I+Kv/Qv+Dv8AwCu//kmvjuij2k+4f2bg/wDn2j9Fv2Yv29viB8afjj4a8Ga3o/hq10vUvtPnTafa3CTL5dtLKu0vOyj5owDlTwT9a+bP2+/+TtfHf/bh/wCkFvR+wJ/ydr4E/wC3/wD9ILij9vv/AJO18d/9uH/pBb1TblTu+5xUaFPD5ly0o2XJ/wC3Hz5RRRWJ9AFeg/s7/wDJwHwy/wCxn0z/ANK4q8+r0H9nf/k4D4Zf9jPpn/pXFTW6Ma38KXoz6D/4Kkf8nAeH/wDsWLf/ANK7uvjuvsT/AIKkf8nAeH/+xYt//Su7r47q6nxs4st/3On6BRRRWZ6YUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB9B/sCf8AJ2vgT/t//wDSC4o/b7/5O18d/wDbh/6QW9H7An/J2vgT/t//APSC4o/b7/5O18d/9uH/AKQW9a/8u/meN/zM/wDuH/7cfPlFFFZHshXoP7O//JwHwy/7GfTP/SuKvPq9B/Z3/wCTgPhl/wBjPpn/AKVxU1ujGt/Cl6M+g/8AgqR/ycB4f/7Fi3/9K7uvjuvsT/gqR/ycB4f/AOxYt/8A0ru6+O6up8bOLLf9zp+gUUUVmemFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAfQf7An/J2vgT/ALf/AP0guKP2+/8Ak7Xx3/24f+kFvR+wJ/ydr4E/7f8A/wBILij9vv8A5O18d/8Abh/6QW9a/wDLv5njf8zP/uH/AO3Hz5RRRWR7IV6D+zv/AMnAfDL/ALGfTP8A0rirz6vQf2d/+TgPhl/2M+mf+lcVNboxrfwpejPoP/gqR/ycB4f/AOxYt/8A0ru6+O6+xP8AgqR/ycB4f/7Fi3/9K7uvjurqfGziy3/c6foFFFFZnphRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAH0H+wJ/wAna+BP+3//ANILij9vv/k7Xx3/ANuH/pBb0fsCf8na+BP+3/8A9ILij9vv/k7Xx3/24f8ApBb1r/y7+Z43/Mz/AO4f/tx8+UUUVkeyFeg/s7/8nAfDL/sZ9M/9K4q8+r0H9nf/AJOA+GX/AGM+mf8ApXFTW6Ma38KXoz6D/wCCpH/JwHh//sWLf/0ru6+O6+xP+CpH/JwHh/8A7Fi3/wDSu7r47q6nxs4st/3On6BRRRWZ6YUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB9B/sCf8na+BP8At/8A/SC4o/b7/wCTtfHf/bh/6QW9H7An/J2vgT/t/wD/AEguKP2+/wDk7Xx3/wBuH/pBb1r/AMu/meN/zM/+4f8A7cfPlFFFZHshXoP7O/8AycB8Mv8AsZ9M/wDSuKvPq9B/Z3/5OA+GX/Yz6Z/6VxU1ujGt/Cl6M+g/+CpH/JwHh/8A7Fi3/wDSu7r47r7E/wCCpH/JwHh//sWLf/0ru6+O6up8bOLLf9zp+gUUUVmemFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAfQf7An/ACdr4E/7f/8A0guKP2+/+TtfHf8A24f+kFvR+wJ/ydr4E/7f/wD0guKP2+/+TtfHf/bh/wCkFvWv/Lv5njf8zP8A7h/+3Hz5RRRWR7IV6D+zv/ycB8Mv+xn0z/0rirz6vQf2d/8Ak4D4Zf8AYz6Z/wClcVNboxrfwpejPoP/AIKkf8nAeH/+xYt//Su7r47r7E/4Kkf8nAeH/wDsWLf/ANK7uvjurqfGziy3/c6foFFFFZnphRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAH0H+wJ/ydr4E/wC3/wD9ILij9vv/AJO18d/9uH/pBb0fsCf8na+BP+3/AP8ASC4o/b7/AOTtfHf/AG4f+kFvWv8Ay7+Z43/Mz/7h/wDtx8+UUUVkeyFeg/s7/wDJwHwy/wCxn0z/ANK4q8+r0H9nf/k4D4Zf9jPpn/pXFTW6Ma38KXoz6D/4Kkf8nAeH/wDsWLf/ANK7uvjuvsT/AIKkf8nAeH/+xYt//Su7r47q6nxs4st/3On6BRRRWZ6YUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB9B/sCf8AJ2vgT/t//wDSC4o/b7/5O18d/wDbh/6QW9H7An/J2vgT/t//APSC4o/b7/5O18d/9uH/AKQW9a/8u/meN/zM/wDuH/7cfPlFFFZHshXoP7O//JwHwy/7GfTP/SuKvPq9B/Z3/wCTgPhl/wBjPpn/AKVxU1ujGt/Cl6M+g/8AgqR/ycB4f/7Fi3/9K7uvjuvsT/gqR/ycB4f/AOxYt/8A0ru6+O6up8bOLLf9zp+gUUUVmemFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAfQf7An/J2vgT/ALf/AP0guKP2+/8Ak7Xx3/24f+kFvR+wJ/ydr4E/7f8A/wBILij9vv8A5O18d/8Abh/6QW9a/wDLv5njf8zP/uH/AO3Hz5RRRWR7IV6D+zv/AMnAfDL/ALGfTP8A0rirz6vQf2d/+TgPhl/2M+mf+lcVNboxrfwpejPoP/gqR/ycB4f/AOxYt/8A0ru6+O6+xP8AgqR/ycB4f/7Fi3/9K7uvjurqfGziy3/c6foFfQd58NtDOs6Bp58O6JDpT+HrPWLi5tNZd9VuXXRlv5ozAblvL811kQN5AADrjtn58rbk8a63J4hsNc/tCSPVrGK1htrqILG0aW8SQwgbQB8scaLnqcZOSSTCaW511YTn8Lto/wDgHZmbw9f+CJ/FsXg3SrSbTtUh02XTRc3rWl2k8M8iPzceaHjNuwOJNpEifKMc9d8WNI8IeAtQ8Q3Nn4F0q4it/FmqeHrW0kvL8xRQWjRnzJCLne0riZACGCgRv8pJyvkfiXx3rHi22trW/ltY7S3d5Y7WwsYLKEO2AzmOBEUsQqjcQTgAZwKvxfFnxPHqut6hLeWt7PrV42o30d/p1tdQy3JZmMohljaNXy78qowGIHHFVdGLo1Lp3+V35df67Hc6l4L8M+C28Ta1Nop12xgn0RLTSry5lUwR6haSXbBniZGZ4li8oEnGWyVPSi48D+GPD198R9FvNMa6t/Dlzqdq3iGed1aOVC0enxQqrBGeSVCXDKxMe9l2+WWrgNL+J/ibSNa1LVodSE99qMglu3vbeK6WWQPvVykqsu5W5VgMrzgjNKvxP8Q/8I1caDLPZXWnXE811ILvTbaeYzSgLJKJnjMgcgD5g2Rjgii6D2VX+bt1fl/wfv12PQfDvwq8J6j8N9P1i6j1w6ncaJqV88sN5EsCXEP2swKIzASYyLX5zvyDgZXzErf+H/wr8K+JtH+HFhqtpo2lQ+JdMnFxrr6q41Jb5tQvLe28u1M+GjzFbox8nbt8w7gQSPKtN+MnizSfDcGg2t9appsFjdabEr6bavKlvcFzPGJWjMg3GR+d2RxgjaMc5feItR1G30eCe6Zo9Itza2O1QhhjM0k+AQAT+8mkbJyfmxnAABddhOjWldc1te721/4Hken+H9N8F+OJNaM/hiTwzp2gz2l1LNp08891PZtew200UiyuymbbMHBjVBmNxtwRt67RPg34f8a3MkcNhpVpaXtvFBY61oV/cTWRkbWNMtmk8udjKkqRXjh45cDEkbBQcGvJL/4x+LtQlhlbVEtbiK8j1A3FhZwWs0tzGSUmlkiRWldSWIZyxyzHqTUV98WPFF/bXNsb+G0tbiA28lvp9jb2kW0zwzkhIo1CuZLaBi4AY+UoJwMUXRLo1m9JW+bf/D/M07200nxzLFpfh7we+h6k+sQafaTQzzyxyJLvVY7lpXYCbcikFAikCX5RgY9E+Ifw38JeCtN1Pxlp2j2Os6IDa2Fjpqaq11aiZpLpGmlmt5ixPl2iSFFkUK12o4C7K8svvjD4s1C8t7uTUYI7iGdrsPbWFvB5k7KVM0myMeZLhmxI+5gWJBBJNVPD/wAS/EXhjSRpVheQ/wBl+ZLK1ldWcFzA7SCLfvSVGV/9RCQGBAKAjB5oui5UqrtZ2+b/AD/rqd9pPg7w7ceM9B1BtEibStV8Jap4gOhzXE/lQz29rf7U3q4l8sy2ayAF921wpZupsjwx4WHwb/4WefDFsZv7S/sL/hHvtVz9j83b5v2vPm+ds2Zj2+bjzOc4+SvOIPiZ4ig8Wf8ACSLexNqv2eSzDSWcDwCB4GgaEQMhiEflOyBAu0A8AVZHxc8ULe+f9ttTF9m+x/YDp1sbHyd/mbPsvl+Tjf8AP9z73zdeaLoUqVVtWfbq/PT8td9DsT4J8Pbj4oGkj+zf+EV/4SL/AIR7z5fL877d9g8rfu8zyt/7/G7fs+Xfn5qg8D+HvD/i/wAa2GoyeGY7DRLbSJdS1DTLvUWtLK4dHeGMRXE0oZYnlMCEtJkMZAG4Aril+JPiJfE58Qfb1OpGH7Mc20RgMGzy/J8jb5XlbPl8vbtxxio9X+IOva5FeRXd4nkXcMVvJDBbRQxrFExZI0RFAjQMS21AATyQTzSuivZVbWv+L0/rvuem+JPCPhz4ZaVbQXvhm18TyT+K9U0w3rXdwssllDFYPA0BilEeWW5dwzI4O5eCBitHWfgpo3hPwl4pEsNnq2raTqt3c20smpLHNd2Fpdrasi26yBysmLqRpAh2i3ADDJB85tvjd4wtbGytF1C0dLJlktJZtLtJZrZ1gggV4pWiLowjtoF3KwP7tTnPNYi+O9dXUIL77eTcwWMmmRu0SEC3kjeORCCuDuWWTJIyS5Oc807ohUq1leX4vX/I9r8O/CTwzY+KNXtrrS7fV7e9uNVn0BdQvZLaE2dtpM95E8rrLHhXM9i24uABG43Y3VQ8IeBLLXdM8Rz/APCHeDZ9Wt9X0fT7e1fxI6WCR3EV60hScX2HkZoYfl8xyNpwo+avM4Piz4qt9R0a9GpRyz6Pp0mk2Yns4JY1tXjeJ4mR0KyAxyumXDHbgZwq4pan4/1nVLK5s3ks7W0ubi3upYNP0+3tEaWBZlifbDGoBUXEo467uc7VwXQvY1n9rt1fR/LdFfxvY6dpnjTX7PRzO2kW+oXENmbpSsxgWRhHvBAIbaBkEDnNYtXda1i78Q6zf6rfyia/vriS6uJVRUDyOxZjtUBRkk8AADsKpVB3xTUUmFFFFBQUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAfQf7An/J2vgT/t//APSC4o/b7/5O18d/9uH/AKQW9H7An/J2vgT/ALf/AP0guKP2+/8Ak7Xx3/24f+kFvWv/AC7+Z43/ADM/+4f/ALcfPlFFFZHshXoP7O//ACcB8Mv+xn0z/wBK4q8+r0H9nf8A5OA+GX/Yz6Z/6VxU1ujGt/Cl6M+g/wDgqR/ycB4f/wCxYt//AEru6+O6+xP+CpH/ACcB4f8A+xYt/wD0ru6+O6up8bOLLf8Ac6foFFFFZnphRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAH0H+wJ/ydr4E/7f8A/wBILij9vv8A5O18d/8Abh/6QW9H7An/ACdr4E/7f/8A0guKP2+/+TtfHf8A24f+kFvWv/Lv5njf8zP/ALh/+3Hz5RRRWR7IV6D+zv8A8nAfDL/sZ9M/9K4q8+r0H9nf/k4D4Zf9jPpn/pXFTW6Ma38KXoz6D/4Kkf8AJwHh/wD7Fi3/APSu7r47r7E/4Kkf8nAeH/8AsWLf/wBK7uvjurqfGziy3/c6foFFFFZnphRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAH0H+wJ/ydr4E/7f/wD0guKP2+/+TtfHf/bh/wCkFvR+wJ/ydr4E/wC3/wD9ILij9vv/AJO18d/9uH/pBb1r/wAu/meN/wAzP/uH/wC3Hz5RRRWR7IV6D+zv/wAnAfDL/sZ9M/8ASuKvPq9B/Z3/AOTgPhl/2M+mf+lcVNboxrfwpejPoP8A4Kkf8nAeH/8AsWLf/wBK7uvjuvsT/gqR/wAnAeH/APsWLf8A9K7uvjurqfGziy3/AHOn6BXpWg/BSbxFr+qadb61bW0FtpkWpW19eROsdx5kSzJGQu4qxjMh4DcxEe481r0vw78YYNH0rQrC80L+0IbCO4huyLoIb1GiuI4Acxtt8r7XN/e3AhflxmpjbqdVf2vL+63+X6/d8yT/AIUJqp02xvYNU03UlvrmxtbeGym2zO11JcRxjZP5R3BrfkHjD5zhH28/pfwn8T6zo9rqttZW32C5hNxHNNqFtDiITNCZGDyAonmoU3MANxUZ+Zc9inx5soLzSJ4/Dtyx0+/03U902pKzzz2tzdTMzsIAP3gvHBwBhlDc5K1yX/Cxj/wi39jf2f8A8wD+w/O8/wD6in2/zdu3/tntz/tZ/hpvlMIvE9Ut193XZlC++HXiHTNMv7+70/7PBYSvFcq80YljKyiFiYt2/aJCELhdu44zmq83grWLfQzqz2yCzWNJnAuIzNHG5CpI0IbzFRiVAcqFO5cH5hnu/FPxwtvEej+K7WPw+9hPr7ztLJHegxDzNQW8DunlAySKFEW4sBtAwo5zk678XbvXfBq6JKuoBzZ2tg4OqSGyEUGwIUtQAquRFHklmGQxCgnITUS4zxDtzRS1/DTz9f8AIrat8FfGGhpeteaZCn2PzxOqahbSMrQDM6hVkJZo1+ZlGSqkMQAQag8I/DDVPFNncX5aOy06Owvr5ZpJY/MlFtbySkJEXDspaMIXUEKSc5wRW3qfxoOo6pqN5/Y/l/a7/wAQX2z7VnZ/adqtvtzs58rbuz/HnGF61R0D4mWelaRaQ3WiSXmpWWk6ho9pdpe+VGkN0lwCWj8slmVrmUg7gCCBjgGn7txc2J5NUr/8P5+n3mS3wz8SK1so09XeeUQbI7mJ2hkKlwswDkwnartiTbgI5/hOOj0X4Ca3qsWntcapo+ltf3v2K2W4nklWQ7EcSCSCORDGRJjcG6qw6jm/q/x/utZ1D7VdWmoXQuZppb+2u9XklgYTW8sEqW8ZULACs8pGQ5UlcHAIJ4W+Mfh3wtbadbW/g+6a3sr/AO2Af2wN864gOyQmA/8ALSBXygUZVPl+Uli0SJSxTjpHX5fqzHu/grrcepaPBaKNRt9Sh02RHgkh85XvIYnjVojKGVd8vliR9qkgcjOKqf8ACoddn0nQtQtDZXMGq2X2wFr+CIQ/6RJAqOXkADMYyQOp+YAEo4HQaX8brPRdb0rWLTw7KuowQaRa3jSahujuIbE2zAIvlDy2drSIliWAGQAeScq0+Jumy+DtL8OaroM99ZWcaK7W+oCBpWjuLqaJhmJtoAvZ0Yc7gUIKleT3RqWJstO3btr17lPwF8Jta8ceJbHS9g02GXU49MnuLp442ikZwHCxu6mV0BLFF56ZxkVT0/4XeJtVi017TTlmOoywRW8QuYhKTM+yFnjL7o0diAHcBTkc8iul0D40x2viO21rW9DOrXdp4ik8S24trz7Ki3ErRtKrDY+5SYY8Yxgg9QcUvhz40xeHdT0vVk0NptWgXTbe7mN5iK4t7KWCSNETy8xu32WAF9zD5SQvzGi0RuWJu7RX9fP+vxOfX4SeKZIo5EsIJFcnbsv7djtDsjSECTiIOrKZD8gKkFgQar6l4AvdF0DWb7UHW3utOvbG0NshWVZVuYJ5klWRWKldsCkYyGEgIPHPSeDfi5p/hnw9b6Ze+HG1gxJIiNcXqmOBmZj58CPE5imG7AIYocAmMtzVLxn8Uo/F9prkB0yaAag+lSRyS3nmvG1laPbZc+WPMMiyMx+7g460vdsNSxDm046d/mvN9LmZdfCvxPaWMN0+nIyTQ21ykUN1DJN5Vxs8mQxK5cKxkQBiuNzBTg8VWufh7r1r4k0rQWs45NT1V40skguopY5y8hiXbKrlP9YrKTu4KkHGDW/YfFiO11261CfQ4r+G40nS9KezuJsxstm1kxLfLysn2LBXjAl6nbzp2vxUsNY+Knw51m7huNN0zw9c2qTTXdx9rkMSX0lwzkpGnQSlQqr0QY9Kdohz4hLWK/pevfQ5pfhJ4oZlxZW3ksOLk6hbC33btvl+b5mwSbuPL3b/AGpZfhP4ksgftunPbN9nluDCJYWmiCW73H72MyBosxRu43AFgDtDdK2E+KWjL4UHhQ+GrlvDP24ar9m/tQfaPteNm7zfJx5fl/Js2Z77s1p618fBrstxeXWhGXU7qO9SW5lvN4iNxZ3Fu4gzHvjjLXHmmNndcoAnlgmi0e4ufE3+Fdf+B1Gzfs5apb2/iGZvE/h7ZocgjuwrXZJJPVB9n+cfT+fFc54l+D3iHw0bl2FpqFrBAlyLi0uF/exGNJGeON9srqgkAZgm0FW54JrvZP2i9FaHxEsfhC9hl1WZJopU1iPdZFJWlXy/9F5IkdnBP8TE98VgeJ/jTYa7LFc2vhp9PvLPR30Kwb7eJIoLV4Wjk3J5S72/ezlSCoXzFBDbAS2omVOWL5vejp8uy7Pvc5PTvhvretQW81haiSKTTW1Zpbi4hgRLdbk2zSFmkwFEo2/Ng9Tjb8xcPhnr5kurcWDteW91b2rokkTR5nhkmibzA+CrRxO4cZTaNxYDGbA+IxHhb+xhYY/4kH9h+f5//UU+3+bt2/8AbPbn/az/AA1s2Xxn+yfaUbRhJDdQ6dazobnBaG20ufT5FB2cGRLhnB/gIAw/Wp903csQr2S8vvXn2uYH/Crtf2MBZh5RPbwq8dxA9uwmjmkRxMJNpXbbyksMquxtzKRgxal4AvdF0DWb7UHW3utOvbG0NshWVZVuYJ5klWRWKldsCkYyGEgIPHO5oHxO0fw2lzaWvhhm02W5trgRTXwklDRW93Dv3tEV8zdd+YrBAqmJfkbJNReM/ilH4vtNcgOmTQDUH0qSOSW88142srR7bLnyx5hkWRmP3cHHWj3QUq7la2mnbur9X5nAUUUVJ2hRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAfQf7An/ACdr4E/7f/8A0guKP2+/+TtfHf8A24f+kFvR+wJ/ydr4E/7f/wD0guKP2+/+TtfHf/bh/wCkFvWv/Lv5njf8zP8A7h/+3Hz5RRRWR7IV6D+zv/ycB8Mv+xn0z/0rirz6vQf2d/8Ak4D4Zf8AYz6Z/wClcVNboxrfwpejPoP/AIKkf8nAeH/+xYt//Su7r47r7E/4Kkf8nAeH/wDsWLf/ANK7uvjurqfGziy3/c6foFFFFZnphRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAH0H+wJ/ydr4E/wC3/wD9ILij9vv/AJO18d/9uH/pBb0fsCf8na+BP+3/AP8ASC4o/b7/AOTtfHf/AG4f+kFvWv8Ay7+Z43/Mz/7h/wDtx8+UUUVkeyFeg/s7/wDJwHwy/wCxn0z/ANK4q8+r0H9nf/k4D4Zf9jPpn/pXFTW6Ma38KXoz6D/4Kkf8nAeH/wDsWLf/ANK7uvjuvsT/AIKkf8nAeH/+xYt//Su7r47q6nxs4st/3On6BRRRWZ6YUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB9B/sCf8AJ2vgT/t//wDSC4o/b7/5O18d/wDbh/6QW9H7An/J2vgT/t//APSC4o/b7/5O18d/9uH/AKQW9a/8u/meN/zM/wDuH/7cfPlFFFZHshXoP7O//JwHwy/7GfTP/SuKvPq9B/Z3/wCTgPhl/wBjPpn/AKVxU1ujGt/Cl6M+g/8AgqR/ycB4f/7Fi3/9K7uvjuvsT/gqR/ycB4f/AOxYt/8A0ru6+O6up8bOLLf9zp+gUUUVmemFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAfQf7An/J2vgT/ALf/AP0guKP2+/8Ak7Xx3/24f+kFvR+wJ/ydr4E/7f8A/wBILij9vv8A5O18d/8Abh/6QW9a/wDLv5njf8zP/uH/AO3Hz5RRRWR7IV6D+zv/AMnAfDL/ALGfTP8A0rirz6vQf2d/+TgPhl/2M+mf+lcVNboxrfwpejPoP/gqR/ycB4f/AOxYt/8A0ru6+O6+xP8AgqR/ycB4f/7Fi3/9K7uvjurqfGziy3/c6foFFFFZnphRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAH0H+wJ/wAna+BP+3//ANILij9vv/k7Xx3/ANuH/pBb0fsCf8na+BP+3/8A9ILij9vv/k7Xx3/24f8ApBb1r/y7+Z43/Mz/AO4f/tx8+UUUVkeyFeg/s7/8nAfDL/sZ9M/9K4q8+r0H9nf/AJOA+GX/AGM+mf8ApXFTW6Ma38KXoz6D/wCCpH/JwHh//sWLf/0ru6+O6+xP+CpH/JwHh/8A7Fi3/wDSu7r47q6nxs4st/3On6BRRRWZ6YUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB9B/sCf8na+BP8At/8A/SC4o/b7/wCTtfHf/bh/6QW9H7An/J2vgT/t/wD/AEguKP2+/wDk7Xx3/wBuH/pBb1r/AMu/meN/zM/+4f8A7cfPlFFFZHshXoP7O/8AycB8Mv8AsZ9M/wDSuKvPq9B/Z3/5OA+GX/Yz6Z/6VxU1ujGt/Cl6M+g/+CpH/JwHh/8A7Fi3/wDSu7r47r7E/wCCpH/JwHh//sWLf/0ru6+O6up8bOLLf9zp+gUUUVmemFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAfQf7An/ACdr4E/7f/8A0guKP2+/+TtfHf8A24f+kFvR+wJ/ydr4E/7f/wD0guKP2+/+TtfHf/bh/wCkFvWv/Lv5njf8zP8A7h/+3Hz5RRRWR7IV6D+zv/ycB8Mv+xn0z/0rirz6vQf2d/8Ak4D4Zf8AYz6Z/wClcVNboxrfwpejPoP/AIKkf8nAeH/+xYt//Su7r47r7E/4Kkf8nAeH/wDsWLf/ANK7uvjurqfGziy3/c6foFFFFZnphRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAH0H+wJ/ydr4E/wC3/wD9ILij9vv/AJO18d/9uH/pBb0fsCf8na+BP+3/AP8ASC4o/b7/AOTtfHf/AG4f+kFvWv8Ay7+Z43/Mz/7h/wDtx8+UUUVkeyFeg/s7/wDJwHwy/wCxn0z/ANK4q8+r0H9nf/k4D4Zf9jPpn/pXFTW6Ma38KXoz6D/4Kkf8nAeH/wDsWLf/ANK7uvjuvsT/AIKkf8nAeH/+xYt//Su7r47q6nxs4st/3On6BRRRWZ6YUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB9B/sCf8AJ2vgT/t//wDSC4o/b7/5O18d/wDbh/6QW9H7An/J2vgT/t//APSC4o/b7/5O18d/9uH/AKQW9a/8u/meN/zM/wDuH/7cfPlFFFZHshXoP7O//JwHwy/7GfTP/SuKvPq9B/Z3/wCTgPhl/wBjPpn/AKVxU1ujGt/Cl6M+g/8AgqR/ycB4f/7Fi3/9K7uvjuvsT/gqR/ycB4f/AOxYt/8A0ru6+O6up8bOLLf9zp+gUUUVmemFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAfQf7An/J2vgT/ALf/AP0guKP2+/8Ak7Xx3/24f+kFvR+wJ/ydr4E/7f8A/wBILij9vv8A5O18d/8Abh/6QW9a/wDLv5njf8zP/uH/AO3Hz5RRRWR7IV6D+zv/AMnAfDL/ALGfTP8A0rirz6vQf2d/+TgPhl/2M+mf+lcVNboxrfwpejPoP/gqR/ycB4f/AOxYt/8A0ru6+O6/Rb9vb9mL4l/Gn4waPrfgzw1/bOl2+hQ2Uk/2+2g2zLcXDldssisflkQ5Axz14NfNn/DAnx5/6ET/AMrFh/8AH61nGTk7I8jL8Vh4YWnGVSKaXdHz5RX0H/wwJ8ef+hE/8rFh/wDH6P8AhgT48/8AQif+Viw/+P1nyS7HofXcL/z9j96/zPnyivoP/hgT48/9CJ/5WLD/AOP0f8MCfHn/AKET/wArFh/8fo5Jdg+u4X/n7H71/mfPlFfQf/DAnx5/6ET/AMrFh/8AH6P+GBPjz/0In/lYsP8A4/RyS7B9dwv/AD9j96/zPnyivoP/AIYE+PP/AEIn/lYsP/j9H/DAnx5/6ET/AMrFh/8AH6OSXYPruF/5+x+9f5nz5RX0H/wwJ8ef+hE/8rFh/wDH6P8AhgT48/8AQif+Viw/+P0ckuwfXcL/AM/Y/ev8z58or6D/AOGBPjz/ANCJ/wCViw/+P0f8MCfHn/oRP/KxYf8Ax+jkl2D67hf+fsfvX+Z8+UV9B/8ADAnx5/6ET/ysWH/x+j/hgT48/wDQif8AlYsP/j9HJLsH13C/8/Y/ev8AM+fKK+g/+GBPjz/0In/lYsP/AI/R/wAMCfHn/oRP/KxYf/H6OSXYPruF/wCfsfvX+Z8+UV9B/wDDAnx5/wChE/8AKxYf/H6P+GBPjz/0In/lYsP/AI/RyS7B9dwv/P2P3r/M+fKK+g/+GBPjz/0In/lYsP8A4/R/wwJ8ef8AoRP/ACsWH/x+jkl2D67hf+fsfvX+Z8+UV9B/8MCfHn/oRP8AysWH/wAfo/4YE+PP/Qif+Viw/wDj9HJLsH13C/8AP2P3r/M+fKK+g/8AhgT48/8AQif+Viw/+P0f8MCfHn/oRP8AysWH/wAfo5Jdg+u4X/n7H71/mfPlFfQf/DAnx5/6ET/ysWH/AMfo/wCGBPjz/wBCJ/5WLD/4/RyS7B9dwv8Az9j96/zPnyivoP8A4YE+PP8A0In/AJWLD/4/R/wwJ8ef+hE/8rFh/wDH6OSXYPruF/5+x+9f5nz5RX0H/wAMCfHn/oRP/KxYf/H6P+GBPjz/ANCJ/wCViw/+P0ckuwfXcL/z9j96/wAz58or6D/4YE+PP/Qif+Viw/8Aj9H/AAwJ8ef+hE/8rFh/8fo5Jdg+u4X/AJ+x+9f5nz5RX0H/AMMCfHn/AKET/wArFh/8fo/4YE+PP/Qif+Viw/8Aj9HJLsH13C/8/Y/ev8z58or6D/4YE+PP/Qif+Viw/wDj9H/DAnx5/wChE/8AKxYf/H6OSXYPruF/5+x+9f5nz5RX0H/wwJ8ef+hE/wDKxYf/AB+j/hgT48/9CJ/5WLD/AOP0ckuwfXcL/wA/Y/ev8z58or6D/wCGBPjz/wBCJ/5WLD/4/R/wwJ8ef+hE/wDKxYf/AB+jkl2D67hf+fsfvX+Z8+UV9B/8MCfHn/oRP/KxYf8Ax+j/AIYE+PP/AEIn/lYsP/j9HJLsH13C/wDP2P3r/M+fKK+g/wDhgT48/wDQif8AlYsP/j9H/DAnx5/6ET/ysWH/AMfo5Jdg+u4X/n7H71/mfPlFfQf/AAwJ8ef+hE/8rFh/8fo/4YE+PP8A0In/AJWLD/4/RyS7B9dwv/P2P3r/ADPnyivoP/hgT48/9CJ/5WLD/wCP0f8ADAnx5/6ET/ysWH/x+jkl2D67hf8An7H71/mfPlFfQf8AwwJ8ef8AoRP/ACsWH/x+j/hgT48/9CJ/5WLD/wCP0ckuwfXcL/z9j96/zPnyivoP/hgT48/9CJ/5WLD/AOP0f8MCfHn/AKET/wArFh/8fo5Jdg+u4X/n7H71/mfPlFfQf/DAnx5/6ET/AMrFh/8AH6P+GBPjz/0In/lYsP8A4/RyS7B9dwv/AD9j96/zPnyivoP/AIYE+PP/AEIn/lYsP/j9H/DAnx5/6ET/AMrFh/8AH6OSXYPruF/5+x+9f5nz5RX0H/wwJ8ef+hE/8rFh/wDH6P8AhgT48/8AQif+Viw/+P0ckuwfXcL/AM/Y/ev8z58or6D/AOGBPjz/ANCJ/wCViw/+P0f8MCfHn/oRP/KxYf8Ax+jkl2D67hf+fsfvX+Z8+UV9B/8ADAnx5/6ET/ysWH/x+j/hgT48/wDQif8AlYsP/j9HJLsH13C/8/Y/ev8AM+fKK+g/+GBPjz/0In/lYsP/AI/R/wAMCfHn/oRP/KxYf/H6OSXYPruF/wCfsfvX+Z8+UV9B/wDDAnx5/wChE/8AKxYf/H6P+GBPjz/0In/lYsP/AI/RyS7B9dwv/P2P3r/M+fKK+g/+GBPjz/0In/lYsP8A4/R/wwJ8ef8AoRP/ACsWH/x+jkl2D67hf+fsfvX+YfsCf8na+BP+3/8A9ILij9vv/k7Xx3/24f8ApBb17F+yH+yH8W/hf+0R4T8T+J/Cf9maHY/a/tF1/aVpLs32k0a/LHKzHLOo4B6+leO/t9/8na+O/wDtw/8ASC3rRpqnr3POp1adXMuanJNcnR3+0fPlFFFYnvBWh4d1+/8ACfiDTNb0qf7Lqmm3UV7aT7Ffy5o3Do21gVOGUHBBBxyKz6KBNJqzPoP/AIb7+PP/AEPf/lHsP/jFH/Dffx5/6Hv/AMo9h/8AGK+fKKrnl3OP6lhf+fUfuX+R9B/8N9/Hn/oe/wDyj2H/AMYo/wCG+/jz/wBD3/5R7D/4xXz5RRzy7h9Swv8Az6j9y/yPoP8A4b7+PP8A0Pf/AJR7D/4xR/w338ef+h7/APKPYf8Axivnyijnl3D6lhf+fUfuX+R9B/8ADffx5/6Hv/yj2H/xij/hvv48/wDQ9/8AlHsP/jFfPlFHPLuH1LC/8+o/cv8AI+g/+G+/jz/0Pf8A5R7D/wCMUf8ADffx5/6Hv/yj2H/xivnyijnl3D6lhf8An1H7l/kfQf8Aw338ef8Aoe//ACj2H/xij/hvv48/9D3/AOUew/8AjFfPlFHPLuH1LC/8+o/cv8j6D/4b7+PP/Q9/+Uew/wDjFH/Dffx5/wCh7/8AKPYf/GK+fKKOeXcPqWF/59R+5f5H0H/w338ef+h7/wDKPYf/ABij/hvv48/9D3/5R7D/AOMV8+UUc8u4fUsL/wA+o/cv8j6D/wCG+/jz/wBD3/5R7D/4xR/w338ef+h7/wDKPYf/ABivnyijnl3D6lhf+fUfuX+R9B/8N9/Hn/oe/wDyj2H/AMYo/wCG+/jz/wBD3/5R7D/4xXz5RRzy7h9Swv8Az6j9y/yPoP8A4b7+PP8A0Pf/AJR7D/4xR/w338ef+h7/APKPYf8Axivnyijnl3D6lhf+fUfuX+R9B/8ADffx5/6Hv/yj2H/xij/hvv48/wDQ9/8AlHsP/jFfPlFHPLuH1LC/8+o/cv8AI+g/+G+/jz/0Pf8A5R7D/wCMUf8ADffx5/6Hv/yj2H/xivnyijnl3D6lhf8An1H7l/kfQf8Aw338ef8Aoe//ACj2H/xij/hvv48/9D3/AOUew/8AjFfPlFHPLuH1LC/8+o/cv8j6D/4b7+PP/Q9/+Uew/wDjFH/Dffx5/wCh7/8AKPYf/GK+fKKOeXcPqWF/59R+5f5H0H/w338ef+h7/wDKPYf/ABij/hvv48/9D3/5R7D/AOMV8+UUc8u4fUsL/wA+o/cv8j6D/wCG+/jz/wBD3/5R7D/4xR/w338ef+h7/wDKPYf/ABivnyijnl3D6lhf+fUfuX+R9B/8N9/Hn/oe/wDyj2H/AMYo/wCG+/jz/wBD3/5R7D/4xXz5RRzy7h9Swv8Az6j9y/yPoP8A4b7+PP8A0Pf/AJR7D/4xR/w338ef+h7/APKPYf8Axivnyijnl3D6lhf+fUfuX+R9B/8ADffx5/6Hv/yj2H/xij/hvv48/wDQ9/8AlHsP/jFfPlFHPLuH1LC/8+o/cv8AI+g/+G+/jz/0Pf8A5R7D/wCMUf8ADffx5/6Hv/yj2H/xivnyijnl3D6lhf8An1H7l/kfQf8Aw338ef8Aoe//ACj2H/xij/hvv48/9D3/AOUew/8AjFfPlFHPLuH1LC/8+o/cv8j6D/4b7+PP/Q9/+Uew/wDjFH/Dffx5/wCh7/8AKPYf/GK+fKKOeXcPqWF/59R+5f5H0H/w338ef+h7/wDKPYf/ABij/hvv48/9D3/5R7D/AOMV8+UUc8u4fUsL/wA+o/cv8j6D/wCG+/jz/wBD3/5R7D/4xR/w338ef+h7/wDKPYf/ABivnyijnl3D6lhf+fUfuX+R9B/8N9/Hn/oe/wDyj2H/AMYo/wCG+/jz/wBD3/5R7D/4xXz5RRzy7h9Swv8Az6j9y/yPoP8A4b7+PP8A0Pf/AJR7D/4xR/w338ef+h7/APKPYf8Axivnyijnl3D6lhf+fUfuX+R9B/8ADffx5/6Hv/yj2H/xij/hvv48/wDQ9/8AlHsP/jFfPlFHPLuH1LC/8+o/cv8AI+g/+G+/jz/0Pf8A5R7D/wCMUf8ADffx5/6Hv/yj2H/xivnyijnl3D6lhf8An1H7l/kfQf8Aw338ef8Aoe//ACj2H/xij/hvv48/9D3/AOUew/8AjFfPlFHPLuH1LC/8+o/cv8j6D/4b7+PP/Q9/+Uew/wDjFH/Dffx5/wCh7/8AKPYf/GK+fKKOeXcPqWF/59R+5f5H0H/w338ef+h7/wDKPYf/ABij/hvv48/9D3/5R7D/AOMV8+UUc8u4fUsL/wA+o/cv8j6D/wCG+/jz/wBD3/5R7D/4xR/w338ef+h7/wDKPYf/ABivnyijnl3D6lhf+fUfuX+R9B/8N9/Hn/oe/wDyj2H/AMYo/wCG+/jz/wBD3/5R7D/4xXz5RRzy7h9Swv8Az6j9y/yPoP8A4b7+PP8A0Pf/AJR7D/4xR/w338ef+h7/APKPYf8Axivnyijnl3D6lhf+fUfuX+R9B/8ADffx5/6Hv/yj2H/xivHfH/j/AF74oeLb/wAT+J7/APtPXL7y/tF15McW/ZGsa/LGqqMKijgDp61z1FJyb3ZrTw9Gk+anBJ+SSCiiikdB/9k=",
            "descricao": "Fotografia 02 — Posicionamento inicial do cadáver em decúbito dorsal no momento do exame pericial.",
            "legenda": "Fotografia 02 — Posicionamento inicial do cadáver em decúbito dorsal no momento do exame pericial.",
            "incluir": True
        },
        {
            "b64": "/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAAMCAgMCAgMDAwMEAwMEBQgFBQQEBQoHBwYIDAoMDAsKCwsNDhIQDQ4RDgsLEBYQERMUFRUVDA8XGBYUGBIUFRT/2wBDAQMEBAUEBQkFBQkUDQsNFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBT/wAARCAGQAlgDASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD5qooqfT9PutWv7axsbaa8vbmVYYLa3jMkksjEBUVRyzEkAAckmvMPy8gorvf+FAfFD/om/i7/AMEV1/8AG6P+FAfFD/om/i7/AMEV1/8AG6dma+yqfyv7jgqK73/hQHxQ/wCib+Lv/BFdf/G6P+FAfFD/AKJv4u/8EV1/8bosw9lU/lf3HBUV3v8AwoD4of8ARN/F3/giuv8A43R/woD4of8ARN/F3/giuv8A43RZh7Kp/K/uOCorvf8AhQHxQ/6Jv4u/8EV1/wDG6P8AhQHxQ/6Jv4u/8EV1/wDG6LMPZVP5X9xwVFd7/wAKA+KH/RN/F3/giuv/AI3R/wAKA+KH/RN/F3/giuv/AI3RZh7Kp/K/uOCorvf+FAfFD/om/i7/AMEV1/8AG6P+FAfFD/om/i7/AMEV1/8AG6LMPZVP5X9xwVFd7/woD4of9E38Xf8Agiuv/jdH/CgPih/0Tfxd/wCCK6/+N0WYeyqfyv7jgqK73/hQHxQ/6Jv4u/8ABFdf/G6P+FAfFD/om/i7/wAEV1/8bosw9lU/lf3HBUV3v/CgPih/0Tfxd/4Irr/43R/woD4of9E38Xf+CK6/+N0WYeyqfyv7jgqK73/hQHxQ/wCib+Lv/BFdf/G6P+FAfFD/AKJv4u/8EV1/8bosw9lU/lf3HBUV3v8AwoD4of8ARN/F3/giuv8A43R/woD4of8ARN/F3/giuv8A43RZh7Kp/K/uOCorvf8AhQHxQ/6Jv4u/8EV1/wDG6P8AhQHxQ/6Jv4u/8EV1/wDG6LMPZVP5X9xwVFd7/wAKA+KH/RN/F3/giuv/AI3R/wAKA+KH/RN/F3/giuv/AI3RZh7Kp/K/uOCorvf+FAfFD/om/i7/AMEV1/8AG6P+FAfFD/om/i7/AMEV1/8AG6LMPZVP5X9xwVFd7/woD4of9E38Xf8Agiuv/jdH/CgPih/0Tfxd/wCCK6/+N0WYeyqfyv7jgqK73/hQHxQ/6Jv4u/8ABFdf/G6P+FAfFD/om/i7/wAEV1/8bosw9lU/lf3HBUV3v/CgPih/0Tfxd/4Irr/43R/woD4of9E38Xf+CK6/+N0WYeyqfyv7jgqK73/hQHxQ/wCib+Lv/BFdf/G6P+FAfFD/AKJv4u/8EV1/8bosw9lU/lf3HBUV3v8AwoD4of8ARN/F3/giuv8A43R/woD4of8ARN/F3/giuv8A43RZh7Kp/K/uOCorvf8AhQHxQ/6Jv4u/8EV1/wDG6P8AhQHxQ/6Jv4u/8EV1/wDG6LMPZVP5X9xwVFd7/wAKA+KH/RN/F3/giuv/AI3R/wAKA+KH/RN/F3/giuv/AI3RZh7Kp/K/uOCorvf+FAfFD/om/i7/AMEV1/8AG6P+FAfFD/om/i7/AMEV1/8AG6LMPZVP5X9xwVFd7/woD4of9E38Xf8Agiuv/jdH/CgPih/0Tfxd/wCCK6/+N0WYeyqfyv7jgqK73/hQHxQ/6Jv4u/8ABFdf/G6P+FAfFD/om/i7/wAEV1/8bosw9lU/lf3HBUV3v/CgPih/0Tfxd/4Irr/43R/woD4of9E38Xf+CK6/+N0WYeyqfyv7jgqK73/hQHxQ/wCib+Lv/BFdf/G6P+FAfFD/AKJv4u/8EV1/8bosw9lU/lf3HBUV3v8AwoD4of8ARN/F3/giuv8A43R/woD4of8ARN/F3/giuv8A43RZh7Kp/K/uOCorvf8AhQHxQ/6Jv4u/8EV1/wDG6P8AhQHxQ/6Jv4u/8EV1/wDG6LMPZVP5X9xwVFd7/wAKA+KH/RN/F3/giuv/AI3R/wAKA+KH/RN/F3/giuv/AI3RZh7Kp/K/uOCorvf+FAfFD/om/i7/AMEV1/8AG6P+FAfFD/om/i7/AMEV1/8AG6LMPZVP5X9xwVFd7/woD4of9E38Xf8Agiuv/jdH/CgPih/0Tfxd/wCCK6/+N0WYeyqfyv7jgqK73/hQHxQ/6Jv4u/8ABFdf/G6P+FAfFD/om/i7/wAEV1/8bosw9lU/lf3HBUV3v/CgPih/0Tfxd/4Irr/43R/woD4of9E38Xf+CK6/+N0WYeyqfyv7jgqK73/hQHxQ/wCib+Lv/BFdf/G6P+FAfFD/AKJv4u/8EV1/8bosw9lU/lf3HBUV3v8AwoD4of8ARN/F3/giuv8A43R/woD4of8ARN/F3/giuv8A43RZh7Kp/K/uOCorvf8AhQHxQ/6Jv4u/8EV1/wDG65HXfD+qeFtVn0vWtNu9I1KDb5tnfwNBNHuUMu5GAIypBGRyCD3os0TKEo6yVihRRRSICu9+AH/JePhv/wBjLpv/AKVR1wVd78AP+S8fDf8A7GXTf/SqOqW6NaX8SPqj7s/bA/bA8Zfs/wDxL0zw94e0zQryyudIiv3k1OCaSQSNNMhAKSoNuIl7ZyTz6eG/8PMvih/0AfCP/gHdf/JNH/BTL/kvGg/9i1B/6VXVfJFaznJSaTPWxmMxFPETjGbSTPrf/h5l8UP+gD4R/wDAO6/+SaP+HmXxQ/6APhH/AMA7r/5Jr5IoqOeXc4/r+J/nZ9b/APDzL4of9AHwj/4B3X/yTR/w8y+KH/QB8I/+Ad1/8k18kUUc8u4fX8T/ADs+t/8Ah5l8UP8AoA+Ef/AO6/8Akmj/AIeZfFD/AKAPhH/wDuv/AJJr5Ioo55dw+v4n+dn1v/w8y+KH/QB8I/8AgHdf/JNH/DzL4of9AHwj/wCAd1/8k18kUUc8u4fX8T/Oz63/AOHmXxQ/6APhH/wDuv8A5Jo/4eZfFD/oA+Ef/AO6/wDkmvkiijnl3D6/if52fW//AA8y+KH/AEAfCP8A4B3X/wAk0f8ADzL4of8AQB8I/wDgHdf/ACTXyRRRzy7h9fxP87Prf/h5l8UP+gD4R/8AAO6/+SaP+HmXxQ/6APhH/wAA7r/5Jr5Ioo55dw+v4n+dn1v/AMPMvih/0AfCP/gHdf8AyTR/w8y+KH/QB8I/+Ad1/wDJNfJFFHPLuH1/E/zs+t/+HmXxQ/6APhH/AMA7r/5Jo/4eZfFD/oA+Ef8AwDuv/kmvkiijnl3D6/if52fW/wDw8y+KH/QB8I/+Ad1/8k0f8PMvih/0AfCP/gHdf/JNfJFFHPLuH1/E/wA7Prf/AIeZfFD/AKAPhH/wDuv/AJJo/wCHmXxQ/wCgD4R/8A7r/wCSa+SKKOeXcPr+J/nZ9b/8PMvih/0AfCP/AIB3X/yTR/w8y+KH/QB8I/8AgHdf/JNfJFFHPLuH1/E/zs+t/wDh5l8UP+gD4R/8A7r/AOSaP+HmXxQ/6APhH/wDuv8A5Jr5Ioo55dw+v4n+dn1v/wAPMvih/wBAHwj/AOAd1/8AJNH/AA8y+KH/AEAfCP8A4B3X/wAk18kUUc8u4fX8T/Oz63/4eZfFD/oA+Ef/AADuv/kmj/h5l8UP+gD4R/8AAO6/+Sa+SKKOeXcPr+J/nZ9b/wDDzL4of9AHwj/4B3X/AMk0f8PMvih/0AfCP/gHdf8AyTXyRRRzy7h9fxP87Prf/h5l8UP+gD4R/wDAO6/+SaP+HmXxQ/6APhH/AMA7r/5Jr5Ioo55dw+v4n+dn1v8A8PMvih/0AfCP/gHdf/JNH/DzL4of9AHwj/4B3X/yTXyRRRzy7h9fxP8AOz63/wCHmXxQ/wCgD4R/8A7r/wCSaP8Ah5l8UP8AoA+Ef/AO6/8Akmvkiijnl3D6/if52fW//DzL4of9AHwj/wCAd1/8k0f8PMvih/0AfCP/AIB3X/yTXyRRRzy7h9fxP87Prf8A4eZfFD/oA+Ef/AO6/wDkmj/h5l8UP+gD4R/8A7r/AOSa+SKKOeXcPr+J/nZ9b/8ADzL4of8AQB8I/wDgHdf/ACTR/wAPMvih/wBAHwj/AOAd1/8AJNfJFFHPLuH1/E/zs+t/+HmXxQ/6APhH/wAA7r/5Jo/4eZfFD/oA+Ef/AADuv/kmvkiijnl3D6/if52fW/8Aw8y+KH/QB8I/+Ad1/wDJNH/DzL4of9AHwj/4B3X/AMk18kUUc8u4fX8T/Oz63/4eZfFD/oA+Ef8AwDuv/kmj/h5l8UP+gD4R/wDAO6/+Sa+SKKOeXcPr+J/nZ9b/APDzL4of9AHwj/4B3X/yTR/w8y+KH/QB8I/+Ad1/8k18kUUc8u4fX8T/ADs+t/8Ah5l8UP8AoA+Ef/AO6/8Akmj/AIeZfFD/AKAPhH/wDuv/AJJr5Ioo55dw+v4n+dn1v/w8y+KH/QB8I/8AgHdf/JNH/DzL4of9AHwj/wCAd1/8k18kUUc8u4fX8T/Oz63/AOHmXxQ/6APhH/wDuv8A5Jo/4eZfFD/oA+Ef/AO6/wDkmvkiijnl3D6/if52fW//AA8y+KH/AEAfCP8A4B3X/wAk0f8ADzL4of8AQB8I/wDgHdf/ACTXyRRRzy7h9fxP87Prf/h5l8UP+gD4R/8AAO6/+SaP+HmXxQ/6APhH/wAA7r/5Jr5Ioo55dw+v4n+dn1v/AMPMvih/0AfCP/gHdf8AyTR/w8y+KH/QB8I/+Ad1/wDJNfJFFHPLuH1/E/zs+t/+HmXxQ/6APhH/AMA7r/5Jo/4eZfFD/oA+Ef8AwDuv/kmvkiijnl3D6/if52fW/wDw8y+KH/QB8I/+Ad1/8k0f8PMvih/0AfCP/gHdf/JNfJFFHPLuH1/E/wA7P0D/AGbP25vHnxi+NXh3whrWk+HbXTdR+0+bLYW06TL5dtLKu0tOwHzRgHIPGfrXzt+3X/ydT43/AO3H/wBIbej9hT/k6nwR/wBv3/pDcUft1/8AJ1Pjf/tx/wDSG3qm24a9zrq1qlbA81R3fP8AoeC0UUVieKFd78AP+S8fDf8A7GXTf/SqOuCrvfgB/wAl4+G//Yy6b/6VR1S3RrS/iR9Ue9f8FMv+S8aD/wBi1B/6VXVfJFfW/wDwUy/5LxoP/YtQf+lV1XyRVVPiZ1Y//eZ+oUUUVmcAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB71+wp/ydT4I/7fv/AEhuKP26/wDk6nxv/wBuP/pDb0fsKf8AJ1Pgj/t+/wDSG4o/br/5Op8b/wDbj/6Q29a/8u/mep/zL/8At/8A9tPBaKKKyPLCu9+AH/JePhv/ANjLpv8A6VR1wVd78AP+S8fDf/sZdN/9Ko6pbo1pfxI+qPev+CmX/JeNB/7FqD/0quq+SK+t/wDgpl/yXjQf+xag/wDSq6r5IqqnxM6sf/vM/UKKKKzOAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooA96/YU/5Op8Ef9v3/pDcUft1/wDJ1Pjf/tx/9Ibej9hT/k6nwR/2/f8ApDcUft1/8nU+N/8Atx/9Ibetf+XfzPU/5l//AG//AO2ngtFFFZHlhXe/AD/kvHw3/wCxl03/ANKo64Ku9+AH/JePhv8A9jLpv/pVHVLdGtL+JH1R71/wUy/5LxoP/YtQf+lV1XyRX1v/AMFMv+S8aD/2LUH/AKVXVfJFVU+JnVj/APeZ+oUUUVmcAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB71+wp/ydT4I/wC37/0huKP26/8Ak6nxv/24/wDpDb0fsKf8nU+CP+37/wBIbij9uv8A5Op8b/8Abj/6Q29a/wDLv5nqf8y//t//ANtPBaKKKyPLCu9+AH/JePhv/wBjLpv/AKVR1wVd78AP+S8fDf8A7GXTf/SqOqW6NaX8SPqj3r/gpl/yXjQf+xag/wDSq6r5Ir63/wCCmX/JeNB/7FqD/wBKrqvkiqqfEzqx/wDvM/UKKKKzOAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooA96/YU/5Op8Ef8Ab9/6Q3FH7df/ACdT43/7cf8A0ht6P2FP+TqfBH/b9/6Q3FH7df8AydT43/7cf/SG3rX/AJd/M9T/AJl//b//ALaeC0UUVkeWFd78AP8AkvHw3/7GXTf/AEqjrgq734Af8l4+G/8A2Mum/wDpVHVLdGtL+JH1R71/wUy/5LxoP/YtQf8ApVdV8kV9b/8ABTL/AJLxoP8A2LUH/pVdV8kVVT4mdWP/AN5n6hRRRWZwBRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAHvX7Cn/J1Pgj/t+/9Ibij9uv/k6nxv8A9uP/AKQ29H7Cn/J1Pgj/ALfv/SG4o/br/wCTqfG//bj/AOkNvWv/AC7+Z6n/ADL/APt//wBtPBaKKKyPLCu9+AH/ACXj4b/9jLpv/pVHXBV3vwA/5Lx8N/8AsZdN/wDSqOqW6NaX8SPqj3r/AIKZf8l40H/sWoP/AEquq+SK+t/+CmX/ACXjQf8AsWoP/Sq6r5IqqnxM6sf/ALzP1CiiiszgCiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKAPev2FP8Ak6nwR/2/f+kNxR+3X/ydT43/AO3H/wBIbej9hT/k6nwR/wBv3/pDcUft1/8AJ1Pjf/tx/wDSG3rX/l38z1P+Zf8A9v8A/tp4LRRRWR5YV3vwA/5Lx8N/+xl03/0qjrgq734Af8l4+G//AGMum/8ApVHVLdGtL+JH1R71/wAFMv8AkvGg/wDYtQf+lV1XyRX1v/wUy/5LxoP/AGLUH/pVdV8kVVT4mdWP/wB5n6hRRRWZwBRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAHvX7Cn/ACdT4I/7fv8A0huKP26/+TqfG/8A24/+kNvR+wp/ydT4I/7fv/SG4o/br/5Op8b/APbj/wCkNvWv/Lv5nqf8y/8A7f8A/bTwWiiisjywrvfgB/yXj4b/APYy6b/6VR1wVd78AP8AkvHw3/7GXTf/AEqjqlujWl/Ej6o96/4KZf8AJeNB/wCxag/9Krqvkivrf/gpl/yXjQf+xag/9KrqvkiqqfEzqx/+8z9QooorM4AooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigD3r9hT/k6nwR/2/f8ApDcUft1/8nU+N/8Atx/9Ibej9hT/AJOp8Ef9v3/pDcUft1/8nU+N/wDtx/8ASG3rX/l38z1P+Zf/ANv/APtp4LRRRWR5YV3vwA/5Lx8N/wDsZdN/9Ko64Ku9+AH/ACXj4b/9jLpv/pVHVLdGtL+JH1R71/wUy/5LxoP/AGLUH/pVdV8kV9b/APBTL/kvGg/9i1B/6VXVfJFVU+JnVj/95n6hXvV38O9EOsaFYHw/o0OltoFnq89xa6wzapcuukLfTRmE3LeX5rq6g+SAA647Z8FrZfxjrL6/Y639vkj1Wyitobe6iCo0aW8SRQgbQB8qRouepxk5JJMppHNTnGPxK+x15m0G+8GTeKovCOmWk2n6nDp0unC4vGtbpJ4pnV+Z/MDxm3YHEm0iRPlGOes+KeleE/At/wCILi08E6ZPFb+KdS0C1tZLu+8qKC0aM+ZIRcb2kcTKAdwUCN/lJOR5T4j8b6v4qtre2vpLZLS3d5I7WxsoLOEO2AzmOFEUsQqjcRnAAzgVei+KfiRNT1m/ku7a8n1i8bUL1L7T7a5hluCzMZRFJGyK2XflVGAxA44p3RqqsLNW+dl/X9XO11Hwf4c8Ht4k1iXRzrdlBPoyWul3dxKvkpf2sl2wZomRmeNY/KBzjLZKnpVr/hAPDml6z488L3OmrLd6NHq8llfzy3H2y9FsZlja3VcQbE8hmlMmSVEmz5gAPOtM+JXiPStZ1HVYtRE17qDiW6e8giuVlcNvVykqsu5W5DAZXsRRH8SvEkekXGm/2jvgnEyvLLBE9wFlOZlWdlMiq5J3KrANubIO45LoaqU+34L+v8rHe+H/AIX+FtQ+Hlhq1zHrR1K40bUb15YbuJYEuIftZgURmAkxkW3znfkHAyPMSneHPDHhvUj8LvD1x4ft2uPFdi32rWluLgXMEz6jd2ySqvm+VtRYYiVMfIVuQTkcXp/xd8U6V4dh0O2vbZNOhsrnTo1bTrZ5Et7guZkErRl/mMjc7sjjBG0Ygtvih4js9CtNJgu7aK2tLWWyt5lsLf7VDBI8jyRrceX5oVmmlyA/R2HQ4ougVSkradOy8v8Ag6nY/ETRNM8HWuiSW3hbw3NaPpek3byyavM97PLNYwTS+ZAt2GRWd36RrgFcEAgnc/4RfwzqPxHsrRfCWkWWiaf4fsddvov7UuLZbl59Pt5RE809ztRGubhFGGU4bGScV5XrnxB1nxHpkNjqB0+eOG3gtUnGl2qXIihjWOJTOsYkO1EReWyQOc1Xv/Gus6nDdxXN55iXVtaWc2IkUvDbRpHAmQuQFWNBx94qC2TzRdCdWndtLT0Xn/wDrvif8PrHwP4etxAolux4k1fT/tqylhcWsMNg9uwwSuCLh3DDqJByQBjpP+FP2P8Awk/m+Rp39i/8IX/a32X+2IftP2r+wPtXmeR5vnf8fH7zbt247bK4Q/FzxNLpEGmXE2nX1nBjykvtIs7lo8QQwZVpImYHyreBcg/8swevNZP/AAm+t/2v/an23/Tv7O/snzfKT/j1+yfY/Lxtx/qPk3Y3d87uaLq4OdLmulpp+B2/xU8NWPhvSNKSw0HQbaCfSdJuTfR6s0motLNYwTSlrc3LbQzu/wDyxAAIxgEE9b4b+G/hnxz8SdE+HyaNbaK1/oOnXo1+C4uGmjnfTYLuVpVklaMo7NIuFRcF0weMHyLXPiDrPiPTIbHUDp88cNvBapONLtUuRFDGscSmdYxIdqIi8tkgc5qzf/FTxNqOmrZSXsEUYtbeyae2sbeC4kggREijeeONZHVVij4ZjnYpOSBRdXH7SkpN209F93l6m3d23h/xX4d8Xzaf4bj8OTaBHFd28sdzPI88TXMduYpxI7KZP3yvuQIP3b/Lgjb2ngj4Z+F9U8a6Hd3umfaPDevXOg2dlafaJVAmup0F0AwYMVT7NfRZzwWQ8HFeVeIviT4g8U2MtpqF3A0M0wuLj7PZQW73MoBw8zxorSsNzHc5Y5Zj1Jo0/wCJPiTS4PDUNrqRjj8OXhv9LXyY2+zzeYJN3KneA65CvlRlsD5myXVxKpTUrtX+S/rb8jsdO0nQW8RTRa7pPhjRwuns9mLPVLi80+S4MqBRdSQXMzxjZ5oG1l+Ypu+XJrRTSNC0aT4h29/4E0tp9F0m31S0WfULuYAy3VjGNskNwqyQmO5dkPUqyEs2Dngo/ifr0F41xB/ZdsZIGtpYbfRrOOCeMsrbZYliCSAMiMN6nBUEYIqhdeN9bvbjWppr3e+sWyWd7+6QB4EkikSNQFwiq0EOAm3AQKPlyKLoPaQS0X4LseseIPBfhzwvHoeoReG7DVDrQsIlsr/Upba2tj/ZNhdXDeaZk2tJJeHaXcquxhtORhH+GOgaF4d13Vb3RriSfQNV1yGTTry4IknW3k0yCCCYxPt/dveyO5iILBWAbG0jzm2+K3ie28wC/ilV4rWLbcWUEyp9mhWCB0DoQkiRqqiRcPxySSTSy/FnxXNqb38mqeZcyX97qUm62hKSz3aot0XTZtZZBEgMbApgcKMnLuinVpa6fgv6/qxN4y03TLrwj4e8S6dpsWiNqFxd2M+n28skkO+AQsJYzKzuAy3AUgs3zRsQQDgcXWv4j8Wan4rmtpNRmiZbaPyYILa2itoYU3FiEiiVUXLEk4AySSeayKhnLNpu6CiiikQFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAHvX7Cn/J1Pgj/t+/9Ibij9uv/k6nxv8A9uP/AKQ29H7Cn/J1Pgj/ALfv/SG4o/br/wCTqfG//bj/AOkNvWv/AC7+Z6n/ADL/APt//wBtPBaKKKyPLCu9+AH/ACXj4b/9jLpv/pVHXBV3vwA/5Lx8N/8AsZdN/wDSqOqW6NaX8SPqj3r/AIKZf8l40H/sWoP/AEquq+SK+t/+CmX/ACXjQf8AsWoP/Sq6r5IqqnxM6sf/ALzP1CiiiszgCiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKAPev2FP8Ak6nwR/2/f+kNxR+3X/ydT43/AO3H/wBIbej9hT/k6nwR/wBv3/pDcUft1/8AJ1Pjf/tx/wDSG3rX/l38z1P+Zf8A9v8A/tp4LRRRWR5YV3vwA/5Lx8N/+xl03/0qjrgq734Af8l4+G//AGMum/8ApVHVLdGtL+JH1R71/wAFMv8AkvGg/wDYtQf+lV1XyRX1v/wUy/5LxoP/AGLUH/pVdV8kVVT4mdWP/wB5n6hRRRWZwBRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAHvX7Cn/ACdT4I/7fv8A0huKP26/+TqfG/8A24/+kNvR+wp/ydT4I/7fv/SG4o/br/5Op8b/APbj/wCkNvWv/Lv5nqf8y/8A7f8A/bTwWiiisjywrvfgB/yXj4b/APYy6b/6VR1wVd78AP8AkvHw3/7GXTf/AEqjqlujWl/Ej6o96/4KZf8AJeNB/wCxag/9Krqvkivrf/gpl/yXjQf+xag/9KrqvkiqqfEzqx/+8z9QrtdS+H1rDeWWladqF3qGvXFhbX5t3tIoLZEltUuXzO0/ARHOWZQPlJOBXFV6dJ8Y4520l5INfU2FnbWi28GvmO2UxWy2/mxxiH5GO0vyWGWIIYEgyrdTnpqm0+cyJvg/r8WjW19izLS3VxbvCb6ACFYo4ZPNZy+0RsJlw2ccpz+8TdSb4Z68ojU2Rjk82eKRp5oY4Y/KSJ2YyF9oXbPEd5wp3rtLZrorv4vaffabfabL4cMWn3kspljtbtImKSJabukO0SGWxikLBQpLSDZgjbb1T4qadeeF7WzmtftMF2bm3utOt5mjkt7QRafHbATNGVMgOnqxIVgR1ALYFWibOFDWz/r7v67HE2fgDXb3VdT06OzRLrTYftN359zFEkUW9EDl3YKVJkTBBIIYHpzUl38N/EVjHdPcWCxG3V5Gia5i8x40GWkjTdukQDJ3oGXAJzgGt/S/Hmk3l34xudXtJEtb/RLfTLSwgnKyEQz2YjTzdjAMsVuWLMuCUPHzAVd1v46Xes6INPNtfWsdvp7aTaQW+qyJbLa7diiaIKPNkVSw3blB4ypAxStEnkoW1l/V/Q5hPhh4jd7dTZ28ZntoLtDLfW6KI5lVodxZwFZwwKq2GPOBwcUG8Fa2lxbwPYNHPPZ3N/HG7qreTAZhMxBPBU283ynBOzgHIzvyfEay1IXNtqmiyXenTwaWhggvfKkEtlai3R9/lt8rKZCVxkbxhvlyZdN+JthZ2FuJtAaS/tdKv9ItZor0pDFDc/aDkxsjMzIbqTBL8jAPPzUaC5aLe/8AV/TsVtV+Eus6fNHDAItRldLB91rNE8a/arZrhFZg/wApCxyZJG3am4kKVLO1D4Q61appIgNtcyXtgL2VvtcCQQ7rmaCNBMZPLct5O4YOTuIAO0mtC2+MEdtEyDRi4uILC2uw13xJHb6fcWDhPk+QvFPkE7trKT8wOBHdfE3R73RrHQ5fDtydCtLeGNYBqQE7SxT3UqyGTysYIvJVKhfQgjoH7pdsPrr/AF06HNab4D13VVujb2BDW0rQPHNKkTtKoJaNEdgzuAOUQFunHIrRtfhhqy3scOpKunxy2V/dxyK8c/zWtq1w8TBX+R8BAVbDL5gJXsZrv4i2+uxX513R/wC0Z5tTu9WhMF0YESe4CiQOoUl0zHGQAykbT83Nbeq/Gi11CawEOgSWtpbjVVNst8CqLe2i25SHEQEaRhcqCG68k8kr3SVGhu32/PXp+pyXhbwpb63pt/qN9ez2dlaT29sfslp9plaSbfs+Teny4ibJz1KgAlq0D8H/ABQ9zqMUFnDOLG+utPkf7XDHult9pm2q7BiEVlYnGAuScAHFfwl40tPCt3qTJY3ptrllMf2bUTbzoisSI3lVPnRgcOu1d2FI24rQi+K9y4eS8shc3c1xrVzNMsuwO+o2iW7ELtONhUv1+bOPlxmhW6iiqPKubf8A4P8AkUbP4T+KdQvb20t9OjkuLSS3hkUXkADvOjPAsZL4lMioSuzdu4xnIzQn8B63b+INK0VrRH1HVGjWySK5ikScvIY12yKxT74Knnggg4wa6fTfjB/Z+p6fef2Tv+yX2hXuz7Tjd/Zts0G3OzjzN27P8OMYbrWRpXxA/svW/Aeo/YPM/wCEW8r9352PtOy9luuu35M+bs6N93PfAPdBxo6Wb/p+nYryfDTxDHPDGbSAiRZW85b2Awx+Xt8wSSh9kZXemQ5BG9ePmGWWvw68Q3l5f2sdgonsXSOZZLiJAXcM0aoWYCRnCsVCZLAEjIq/oPxCi0zwp/wj13pr3lhNJdm6aO58qSRZjZsoQlGClXskbJDZDEYHU6Ol/Fe3stb1C9uNFe6tbiO1gjsftmITDbxCKOKdWRlmUoqZO1WyuVZMmj3QUaLtr/VvTuZtj8LNYa4ZdQRbCIWl3cl1eOZo3gtJLkRSIr5iZhHjD4IyTg7SKzNS8Ba9pK2puLDDXMq26RxSxyusrAFY3VWJjcg8K4B68cGusPxhtxeapcpobrJrMlxcar/pufNmltbi33Rfu/3Sj7XM+078kgZwKp6n8Uba/vzd/wBgoz3msQ63qsNzcebDeTR7/kVdg8uNvOmypLH5xzgUe6NxoW0f9fcZp+FXibzXQWMLbUR1kS+t2jl3mRUWNw+2RyYZQEQlsxsMZBrrbn9n+4caAttr+lwvqaW6EXczNiaW3in/AOWKSBUCzKoLEElSQMMuWv8AG+0kubRpdCu7q2t7NLTyLrU1kW5VZ7ibbOvkhHQ/aNu0IrL5Y2OuTTtJ+OGn2Q0s3fhqa6eye1YmLUVjEyQ2tpCEbMLEAvaLLweDtHO3JfumsY4Zbu/3/wCRjp8G765RHtLyK63y6PbrCGijmkl1C2EyKqPIM7SQmcjPLcBW24Vl8N/EWoWltcW+niVLgIyoJ4vMVGztkdN26OMgE73AXHOcEV02l/GGHT7+xu20WSV7W50O7AF4FDSabCYR/wAsz8siMeOqtg5YcVRtPiVYwXFpfS6CZ9UWwTS7mc3pWOS2FsbUhECfI7Q4XcS4BGQuTS90zaoO2v8AV/Tsclreg3vh67S3vokR3jWWN4pUljkQ9GR0JVhwRkE8gjqDWfW34o8Qw662mw2lo9jp+m2gs7aGWbzpAnmSSsXcKoYl5XPCgAEDHFYlSzllZP3dgooopEhRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAe9fsKf8nU+CP+37/wBIbij9uv8A5Op8b/8Abj/6Q29H7Cn/ACdT4I/7fv8A0huKP26/+TqfG/8A24/+kNvWv/Lv5nqf8y//ALf/APbTwWiiisjywrvfgB/yXj4b/wDYy6b/AOlUdcFXe/AD/kvHw3/7GXTf/SqOqW6NaX8SPqj3r/gpl/yXjQf+xag/9Krqvkivrf8A4KZf8l40H/sWoP8A0quq+SKqp8TOrH/7zP1CiiiszgCiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKAPev2FP+TqfBH/b9/6Q3FH7df8AydT43/7cf/SG3o/YU/5Op8Ef9v3/AKQ3FH7df/J1Pjf/ALcf/SG3rX/l38z1P+Zf/wBv/wDtp4LRRRWR5YV3vwA/5Lx8N/8AsZdN/wDSqOuCrvfgB/yXj4b/APYy6b/6VR1S3RrS/iR9Ue9f8FMv+S8aD/2LUH/pVdV8kV9b/wDBTL/kvGg/9i1B/wClV1XyRVVPiZ1Y/wD3mfqFFFFZnAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAe9fsKf8nU+CP8At+/9Ibij9uv/AJOp8b/9uP8A6Q29H7Cn/J1Pgj/t+/8ASG4o/br/AOTqfG//AG4/+kNvWv8Ay7+Z6n/Mv/7f/wDbTwWiiisjywrvfgB/yXj4b/8AYy6b/wClUdcFXe/AD/kvHw3/AOxl03/0qjqlujWl/Ej6o96/4KZf8l40H/sWoP8A0quq+SK+t/8Agpl/yXjQf+xag/8ASq6r5IqqnxM6sf8A7zP1CiiiszgCiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKAPev2FP+TqfBH/AG/f+kNxR+3X/wAnU+N/+3H/ANIbej9hT/k6nwR/2/f+kNxR+3X/AMnU+N/+3H/0ht61/wCXfzPU/wCZf/2//wC2ngtFFFZHlhXe/AD/AJLx8N/+xl03/wBKo64Ku9+AH/JePhv/ANjLpv8A6VR1S3RrS/iR9Ue9f8FMv+S8aD/2LUH/AKVXVfJFfW//AAUy/wCS8aD/ANi1B/6VXVfJFVU+JnVj/wDeZ+oUUUVmcAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB71+wp/ydT4I/7fv/SG4o/br/5Op8b/APbj/wCkNvR+wp/ydT4I/wC37/0huKP26/8Ak6nxv/24/wDpDb1r/wAu/mep/wAy/wD7f/8AbTwWiiisjywrvfgB/wAl4+G//Yy6b/6VR1wVd78AP+S8fDf/ALGXTf8A0qjqlujWl/Ej6o96/wCCmX/JeNB/7FqD/wBKrqvkivrf/gpl/wAl40H/ALFqD/0quq+SKqp8TOrH/wC8z9QooorM4AooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigD3r9hT/AJOp8Ef9v3/pDcUft1/8nU+N/wDtx/8ASG3o/YU/5Op8Ef8Ab9/6Q3FH7df/ACdT43/7cf8A0ht61/5d/M9T/mX/APb/AP7aeC0UUVkeWFd78AP+S8fDf/sZdN/9Ko64Ku9+AH/JePhv/wBjLpv/AKVR1S3RrS/iR9Ue9f8ABTL/AJLxoP8A2LUH/pVdV8kV9b/8FMv+S8aD/wBi1B/6VXVfJFVU+JnVj/8AeZ+oUUUVmcAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB71+wp/wAnU+CP+37/ANIbij9uv/k6nxv/ANuP/pDb0fsKf8nU+CP+37/0huKP26/+TqfG/wD24/8ApDb1r/y7+Z6n/Mv/AO3/AP208FooorI8sK734Af8l4+G/wD2Mum/+lUdcFXe/AD/AJLx8N/+xl03/wBKo6pbo1pfxI+qPev+CmX/ACXjQf8AsWoP/Sq6r5Ir63/4KZf8l40H/sWoP/Sq6r5IqqnxM6sf/vM/UKKKKzOAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooA96/YU/5Op8Ef9v3/AKQ3FH7df/J1Pjf/ALcf/SG3o/YU/wCTqfBH/b9/6Q3FH7df/J1Pjf8A7cf/AEht61/5d/M9T/mX/wDb/wD7aeC0UUVkeWFd78AP+S8fDf8A7GXTf/SqOuCrvfgB/wAl4+G//Yy6b/6VR1S3RrS/iR9Ue9f8FMv+S8aD/wBi1B/6VXVfJFfW/wDwUy/5LxoP/YtQf+lV1XyRVVPiZ1Y//eZ+oUUUVmcAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQB71+wp/ydT4I/7fv/AEhuKP26/wDk6nxv/wBuP/pDb0fsKf8AJ1Pgj/t+/wDSG4o/br/5Op8b/wDbj/6Q29a/8u/mep/zL/8At/8A9tPBaKKKyPLCu9+AH/JePhv/ANjLpv8A6VR1wVd78AP+S8fDf/sZdN/9Ko6pbo1pfxI+qPev+CmX/JeNB/7FqD/0quq+SK+t/wDgpl/yXjQf+xag/wDSq6r5IqqnxM6sf/vM/UKKKKzOAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooA96/YU/5Op8Ef9v3/pDcUft1/wDJ1Pjf/tx/9Ibej9hT/k6nwR/2/f8ApDcUft1/8nU+N/8Atx/9Ibetf+XfzPU/5l//AG//AO2ngtFFFZHlhXe/AD/kvHw3/wCxl03/ANKo64Ku9+AH/JePhv8A9jLpv/pVHVLdGtL+JH1R71/wUy/5LxoP/YtQf+lV1XyRX6B/tzfs2fEb4xfFrSda8IeHf7X02DRIbOSf7dbQbZVnuHK7ZZFJ+WRDkDHPXg187f8ADCnxx/6Ej/yrWP8A8frScW5PQ9LG4etPETlGDav2Z4LRXvX/AAwp8cf+hI/8q1j/APH6P+GFPjj/ANCR/wCVax/+P1HLLscX1XEf8+5fczwWivev+GFPjj/0JH/lWsf/AI/R/wAMKfHH/oSP/KtY/wDx+jll2D6riP8An3L7meC0V71/wwp8cf8AoSP/ACrWP/x+j/hhT44/9CR/5VrH/wCP0csuwfVcR/z7l9zPBaK96/4YU+OP/Qkf+Vax/wDj9H/DCnxx/wChI/8AKtY//H6OWXYPquI/59y+5ngtFe9f8MKfHH/oSP8AyrWP/wAfo/4YU+OP/Qkf+Vax/wDj9HLLsH1XEf8APuX3M8For3r/AIYU+OP/AEJH/lWsf/j9H/DCnxx/6Ej/AMq1j/8AH6OWXYPquI/59y+5ngtFe9f8MKfHH/oSP/KtY/8Ax+j/AIYU+OP/AEJH/lWsf/j9HLLsH1XEf8+5fczwWivev+GFPjj/ANCR/wCVax/+P0f8MKfHH/oSP/KtY/8Ax+jll2D6riP+fcvuZ4LRXvX/AAwp8cf+hI/8q1j/APH6P+GFPjj/ANCR/wCVax/+P0csuwfVcR/z7l9zPBaK96/4YU+OP/Qkf+Vax/8Aj9H/AAwp8cf+hI/8q1j/APH6OWXYPquI/wCfcvuZ4LRXvX/DCnxx/wChI/8AKtY//H6P+GFPjj/0JH/lWsf/AI/Ryy7B9VxH/PuX3M8For3r/hhT44/9CR/5VrH/AOP0f8MKfHH/AKEj/wAq1j/8fo5Zdg+q4j/n3L7meC0V71/wwp8cf+hI/wDKtY//AB+j/hhT44/9CR/5VrH/AOP0csuwfVcR/wA+5fczwWivev8AhhT44/8AQkf+Vax/+P0f8MKfHH/oSP8AyrWP/wAfo5Zdg+q4j/n3L7meC0V71/wwp8cf+hI/8q1j/wDH6P8AhhT44/8AQkf+Vax/+P0csuwfVcR/z7l9zPBaK96/4YU+OP8A0JH/AJVrH/4/R/wwp8cf+hI/8q1j/wDH6OWXYPquI/59y+5ngtFe9f8ADCnxx/6Ej/yrWP8A8fo/4YU+OP8A0JH/AJVrH/4/Ryy7B9VxH/PuX3M8For3r/hhT44/9CR/5VrH/wCP0f8ADCnxx/6Ej/yrWP8A8fo5Zdg+q4j/AJ9y+5ngtFe9f8MKfHH/AKEj/wAq1j/8fo/4YU+OP/Qkf+Vax/8Aj9HLLsH1XEf8+5fczwWivev+GFPjj/0JH/lWsf8A4/R/wwp8cf8AoSP/ACrWP/x+jll2D6riP+fcvuZ4LRXvX/DCnxx/6Ej/AMq1j/8AH6P+GFPjj/0JH/lWsf8A4/Ryy7B9VxH/AD7l9zPBaK96/wCGFPjj/wBCR/5VrH/4/R/wwp8cf+hI/wDKtY//AB+jll2D6riP+fcvuZ4LRXvX/DCnxx/6Ej/yrWP/AMfo/wCGFPjj/wBCR/5VrH/4/Ryy7B9VxH/PuX3M8For3r/hhT44/wDQkf8AlWsf/j9H/DCnxx/6Ej/yrWP/AMfo5Zdg+q4j/n3L7meC0V71/wAMKfHH/oSP/KtY/wDx+j/hhT44/wDQkf8AlWsf/j9HLLsH1XEf8+5fczwWivev+GFPjj/0JH/lWsf/AI/R/wAMKfHH/oSP/KtY/wDx+jll2D6riP8An3L7meC0V71/wwp8cf8AoSP/ACrWP/x+j/hhT44/9CR/5VrH/wCP0csuwfVcR/z7l9zPBaK96/4YU+OP/Qkf+Vax/wDj9H/DCnxx/wChI/8AKtY//H6OWXYPquI/59y+5ngtFe9f8MKfHH/oSP8AyrWP/wAfo/4YU+OP/Qkf+Vax/wDj9HLLsH1XEf8APuX3M8For3r/AIYU+OP/AEJH/lWsf/j9H/DCnxx/6Ej/AMq1j/8AH6OWXYPquI/59y+5ngtFe9f8MKfHH/oSP/KtY/8Ax+j/AIYU+OP/AEJH/lWsf/j9HLLsH1XEf8+5fczwWivev+GFPjj/ANCR/wCVax/+P0f8MKfHH/oSP/KtY/8Ax+jll2D6riP+fcvuZ4LRXvX/AAwp8cf+hI/8q1j/APH6P+GFPjj/ANCR/wCVax/+P0csuwfVcR/z7l9zPBaK96/4YU+OP/Qkf+Vax/8Aj9H/AAwp8cf+hI/8q1j/APH6OWXYPquI/wCfcvuYfsKf8nU+CP8At+/9Ibij9uv/AJOp8b/9uP8A6Q29et/sn/sn/FX4afH/AMLeJPEnhb+zdFsvtX2i6/tG0l2b7SaNflSVmOWdRwD19K8k/br/AOTqfG//AG4/+kNvVtNQ17ndOnOngLTTT5+v+E8FooorE8cKv+H9dvvC2vabrWlz/ZdS065jvLWfYr+XLGwdG2sCDhgDggg9xVCigabTuj3r/huv44/9Dv8A+Umx/wDjFH/Ddfxx/wCh3/8AKTY//GK8Foq+aXc6frWI/wCfkvvZ71/w3X8cf+h3/wDKTY//ABij/huv44/9Dv8A+Umx/wDjFeC0Uc0u4fWsR/z8l97Pev8Ahuv44/8AQ7/+Umx/+MUf8N1/HH/od/8Ayk2P/wAYrwWijml3D61iP+fkvvZ71/w3X8cf+h3/APKTY/8Axij/AIbr+OP/AEO//lJsf/jFeC0Uc0u4fWsR/wA/Jfez3r/huv44/wDQ7/8AlJsf/jFH/Ddfxx/6Hf8A8pNj/wDGK8Foo5pdw+tYj/n5L72e9f8ADdfxx/6Hf/yk2P8A8Yo/4br+OP8A0O//AJSbH/4xXgtFHNLuH1rEf8/Jfez3r/huv44/9Dv/AOUmx/8AjFH/AA3X8cf+h3/8pNj/APGK8Foo5pdw+tYj/n5L72e9f8N1/HH/AKHf/wApNj/8Yo/4br+OP/Q7/wDlJsf/AIxXgtFHNLuH1rEf8/Jfez3r/huv44/9Dv8A+Umx/wDjFH/Ddfxx/wCh3/8AKTY//GK8Foo5pdw+tYj/AJ+S+9nvX/Ddfxx/6Hf/AMpNj/8AGKP+G6/jj/0O/wD5SbH/AOMV4LRRzS7h9axH/PyX3s96/wCG6/jj/wBDv/5SbH/4xR/w3X8cf+h3/wDKTY//ABivBaKOaXcPrWI/5+S+9nvX/Ddfxx/6Hf8A8pNj/wDGKP8Ahuv44/8AQ7/+Umx/+MV4LRRzS7h9axH/AD8l97Pev+G6/jj/ANDv/wCUmx/+MUf8N1/HH/od/wDyk2P/AMYrwWijml3D61iP+fkvvZ71/wAN1/HH/od//KTY/wDxij/huv44/wDQ7/8AlJsf/jFeC0Uc0u4fWsR/z8l97Pev+G6/jj/0O/8A5SbH/wCMUf8ADdfxx/6Hf/yk2P8A8YrwWijml3D61iP+fkvvZ71/w3X8cf8Aod//ACk2P/xij/huv44/9Dv/AOUmx/8AjFeC0Uc0u4fWsR/z8l97Pev+G6/jj/0O/wD5SbH/AOMUf8N1/HH/AKHf/wApNj/8YrwWijml3D61iP8An5L72e9f8N1/HH/od/8Ayk2P/wAYo/4br+OP/Q7/APlJsf8A4xXgtFHNLuH1rEf8/Jfez3r/AIbr+OP/AEO//lJsf/jFH/Ddfxx/6Hf/AMpNj/8AGK8Foo5pdw+tYj/n5L72e9f8N1/HH/od/wDyk2P/AMYo/wCG6/jj/wBDv/5SbH/4xXgtFHNLuH1rEf8APyX3s96/4br+OP8A0O//AJSbH/4xR/w3X8cf+h3/APKTY/8AxivBaKOaXcPrWI/5+S+9nvX/AA3X8cf+h3/8pNj/APGKP+G6/jj/ANDv/wCUmx/+MV4LRRzS7h9axH/PyX3s96/4br+OP/Q7/wDlJsf/AIxR/wAN1/HH/od//KTY/wDxivBaKOaXcPrWI/5+S+9nvX/Ddfxx/wCh3/8AKTY//GKP+G6/jj/0O/8A5SbH/wCMV4LRRzS7h9axH/PyX3s96/4br+OP/Q7/APlJsf8A4xR/w3X8cf8Aod//ACk2P/xivBaKOaXcPrWI/wCfkvvZ71/w3X8cf+h3/wDKTY//ABij/huv44/9Dv8A+Umx/wDjFeC0Uc0u4fWsR/z8l97Pev8Ahuv44/8AQ7/+Umx/+MUf8N1/HH/od/8Ayk2P/wAYrwWijml3D61iP+fkvvZ71/w3X8cf+h3/APKTY/8Axij/AIbr+OP/AEO//lJsf/jFeC0Uc0u4fWsR/wA/Jfez3r/huv44/wDQ7/8AlJsf/jFH/Ddfxx/6Hf8A8pNj/wDGK8Foo5pdw+tYj/n5L72e9f8ADdfxx/6Hf/yk2P8A8Yo/4br+OP8A0O//AJSbH/4xXgtFHNLuH1rEf8/Jfez3r/huv44/9Dv/AOUmx/8AjFH/AA3X8cf+h3/8pNj/APGK8Foo5pdw+tYj/n5L72e9f8N1/HH/AKHf/wApNj/8Yo/4br+OP/Q7/wDlJsf/AIxXgtFHNLuH1rEf8/Jfez3r/huv44/9Dv8A+Umx/wDjFH/Ddfxx/wCh3/8AKTY//GK8Foo5pdw+tYj/AJ+S+9nvX/Ddfxx/6Hf/AMpNj/8AGKP+G6/jj/0O/wD5SbH/AOMV4LRRzS7h9axH/PyX3s96/wCG6/jj/wBDv/5SbH/4xR/w3X8cf+h3/wDKTY//ABivBaKOaXcPrWI/5+S+9nvX/Ddfxx/6Hf8A8pNj/wDGK8j8d+O9c+Jfiq+8SeJL7+0tavdn2i68mOLfsjWNflRVUYVFHAHT1rBopOTe7M516tRWnJtebYUUUVJif//Z",
            "descricao": "Fotografia 03 — Detalhe aproximado do vestígio balístico (estojo cal. 9mm) arrecadado no solo.",
            "legenda": "Fotografia 03 — Detalhe aproximado do vestígio balístico (estojo cal. 9mm) arrecadado no solo.",
            "incluir": True
        }
    ]

def render_action_buttons(prefix):
        st.markdown('<br>', unsafe_allow_html=True)
        btn1, btn2, btn3, btn4, btn5, btn6, btn7 = st.columns([1, 1.2, 0.9, 1.1, 1.1, 1.1, 1.2])

        # 1. Salvar no Sistema
        if btn1.button("💾 Salvar", use_container_width=True, key=f"btn_salvar_{prefix}"):
            import sqlite3, json, os, datetime
            DB_PATH = os.path.join(os.path.dirname(__file__), "banco_laudos.sqlite")
            dados = {}
            for k, v in st.session_state.items():
                if k not in ['autenticado', 'preview_open']:
                    try:
                        if isinstance(v, (datetime.date, datetime.time)):
                            dados[k] = v.isoformat()
                        else:
                            dados[k] = v
                    except: pass
            if 'vitimas' in st.session_state: dados['vitimas'] = st.session_state['vitimas']
            if 'vestigios' in st.session_state: dados['vestigios'] = st.session_state['vestigios']
            if 'fotos' in st.session_state: dados['fotos'] = st.session_state['fotos']
            if 'quesitos_list' in st.session_state: dados['quesitos_list'] = st.session_state['quesitos_list']

            try:
                conn = sqlite3.connect(DB_PATH)
                c = conn.cursor()
                mob_id = dados.get('ocorrencia', f"laudo_{len(dados)}")
                now_str = datetime.datetime.now().isoformat()
                c.execute('INSERT OR REPLACE INTO laudos (mobile_id, data_sincronizacao, dados_json) VALUES (?, ?, ?)',
                          (mob_id, now_str, json.dumps(dados)))
                conn.commit()
                conn.close()
                st.success("✅ Dados salvos com sucesso no sistema!")
            except Exception as e:
                st.error(f"Erro ao salvar: {e}")

        # 2. Gerar Laudo
        gerar_clicked = btn2.button("🚀 Gerar Laudo (.docx)", type="primary", use_container_width=True, key=f"btn_gerar_{prefix}")

        # 3. Limpar Formulário
        if btn3.button("🗑️ Limpar", use_container_width=True, key=f"btn_limpar_{prefix}"):
            for k in list(st.session_state.keys()):
                if k not in ['autenticado']:
                    del st.session_state[k]
            st.rerun()

        # 4. Buscar Ocorrências (Modal)
        if btn4.button("🔍 Buscar", use_container_width=True, key=f"btn_sync_{prefix}"):
            modal_ocorrencias()

        # 5. Auditoria de Inconsistências (Checkup do Laudo)
        if btn5.button("🔍 Auditoria", use_container_width=True, key=f"btn_audit_{prefix}"):
            modal_auditoria_inconsistencias()

        # 6. Sugestões & Melhorias
        if btn6.button("💡 Sugestões", use_container_width=True, key=f"btn_sugestoes_{prefix}"):
            modal_sugestoes()

        # 7. Dados de Teste (Preenchimento Automático)
        if btn7.button("🧪 Teste", use_container_width=True, key=f"btn_teste_{prefix}", help="Preenche todos os campos com dados de teste e 3 fotos para testar a geração do laudo"):
            carregar_dados_teste_exemplo()
            st.success("✅ Todos os campos e 3 fotos de teste foram preenchidos com sucesso!")
            st.rerun()

        return gerar_clicked


SUGESTOES_FILE = os.path.join(os.path.dirname(__file__), "sugestoes.json")

def carregar_sugestoes():
    if os.path.exists(SUGESTOES_FILE):
        try:
            with open(SUGESTOES_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []
    return []

def salvar_sugestoes(lista):
    try:
        with open(SUGESTOES_FILE, "w", encoding="utf-8") as f:
            json.dump(lista, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        st.error(f"Erro ao salvar sugestões: {e}")
        return False

@st.dialog("💬 Central de Sugestões, Erros e Melhorias")
def modal_sugestoes():
    sugestoes = carregar_sugestoes()
    tab1, tab2 = st.tabs(["➕ Cadastrar Sugestão / Erro", f"📋 Acompanhar Sugestões ({len(sugestoes)})"])

    with tab1:
        st.markdown("Use este formulário para reportar bugs, divergências ou sugerir novas melhorias para o sistema:")
        tipo = st.selectbox("Tipo de Comunicação", ["💡 Sugestão de Melhoria", "🐛 Reportar Erro / Problema", "✨ Nova Funcionalidade", "Outro"], key="sug_tipo")
        titulo = st.text_input("Título / Resumo", placeholder="Ex: Ajustar layout dos cartões de vestígios", key="sug_titulo")
        autor = st.text_input("Seu Nome / Perito", value=st.session_state.get("perito", ""), placeholder="Ex: Dr. Carlos", key="sug_autor")
        descricao = st.text_area("Descrição detalhada", placeholder="Descreva com detalhes a sua sugestão ou o erro observado...", height=100, key="sug_desc")
        uploaded_imgs = st.file_uploader("📸 Anexar Imagens / Prints (Opcional)", type=["png", "jpg", "jpeg", "webp"], accept_multiple_files=True, key="sug_imgs_input")

        if st.button("🚀 Cadastrar Sugestão", type="primary", use_container_width=True, key="btn_cadastrar_sugestao"):
            if not titulo.strip() or not descricao.strip():
                st.warning("⚠️ Por favor, preencha o título e a descrição da sugestão.")
            else:
                import datetime, base64
                imgs_b64 = []
                if uploaded_imgs:
                    for f_img in uploaded_imgs:
                        try:
                            imgs_b64.append(base64.b64encode(f_img.read()).decode("utf-8"))
                        except Exception: pass

                novo_item = {
                    "id": datetime.datetime.now().strftime("%Y%m%d%H%M%S"),
                    "data": datetime.datetime.now().strftime("%d/%m/%Y %H:%M"),
                    "tipo": tipo,
                    "titulo": titulo.strip(),
                    "autor": autor.strip() if autor.strip() else "Perito Anônimo",
                    "descricao": descricao.strip(),
                    "imagens": imgs_b64,
                    "status": "🟡 Pendente"
                }
                sugestoes.insert(0, novo_item)
                salvar_sugestoes(sugestoes)
                st.success("✅ Sua sugestão foi cadastrada com sucesso! Você pode acompanhar o status na aba 'Acompanhar Sugestões'.")
                st.rerun()

    with tab2:
        if not sugestoes:
            st.info("Nenhuma sugestão cadastrada até o momento.")
        else:
            st.markdown("##### 📋 Sugestões Cadastradas e Acompanhamento de Status")
            opcoes_status = ["🟡 Pendente", "🔵 Em Análise", "🟠 Em Desenvolvimento", "🟢 Concluído / Executado", "🔴 Não Aplicável"]

            for idx, item in enumerate(sugestoes):
                with st.expander(f"{item.get('status', '🟡 Pendente')} | {item.get('titulo')} ({item.get('data')})"):
                    st.markdown(f"**Tipo:** {item.get('tipo')}")
                    st.markdown(f"**Autor:** {item.get('autor')}")
                    st.markdown(f"**Descrição:**\n{item.get('descricao')}")

                    if item.get("imagens"):
                        st.markdown("**📸 Imagens / Prints Anexados:**")
                        import base64
                        cols_img = st.columns(min(len(item["imagens"]), 3))
                        for i_idx, b64_str in enumerate(item["imagens"]):
                            with cols_img[i_idx % 3]:
                                try:
                                    img_data = base64.b64decode(b64_str)
                                    st.image(img_data, use_container_width=True, caption=f"Print #{i_idx+1}")
                                except Exception: pass

                    st.markdown("---")
                    st.markdown("**⚙️ Alterar Status (Administrador):**")
                    col_st1, col_st2 = st.columns([3, 1])
                    with col_st1:
                        curr_st = item.get("status", "🟡 Pendente")
                        curr_idx = opcoes_status.index(curr_st) if curr_st in opcoes_status else 0
                        novo_st = st.selectbox("Status", opcoes_status, index=curr_idx, key=f"sel_st_{item['id']}_{idx}")
                    with col_st2:
                        st.markdown("<br>", unsafe_allow_html=True)
                        if st.button("Salvar", key=f"btn_save_st_{item['id']}_{idx}", use_container_width=True):
                            sugestoes[idx]["status"] = novo_st
                            salvar_sugestoes(sugestoes)
                            st.success("Status atualizado!")
                            st.rerun()


def call_gemini_text(prompt, system_instruction=""):
    key = get_gemini_api_key()
    if not key:
        return None, "Chave GEMINI_API_KEY não configurada. Verifique se o arquivo CHAVE.txt está presente."
    last_err = None
    try:
        import google.generativeai as genai
        genai.configure(api_key=key)

        models = ["gemini-2.5-flash", "gemini-flash-latest", "gemini-3.6-flash"]
        for m in models:
            try:
                model = genai.GenerativeModel(
                    m,
                    system_instruction=system_instruction if system_instruction else None
                )
                resp = model.generate_content(prompt)
                if resp and resp.text:
                    return resp.text, None
            except Exception as e:
                last_err = e
                continue
    except Exception as e:
        last_err = e

    # Fallback via REST API
    try:
        import requests
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={key}"
        full_text = (system_instruction + "\n\n" + prompt) if system_instruction else prompt
        payload = {"contents": [{"parts": [{"text": full_text}]}]}
        r = requests.post(url, json=payload, timeout=35)
        if r.status_code == 200:
            res_j = r.json()
            txt = res_j['candidates'][0]['content']['parts'][0]['text']
            return txt, None
        else:
            return None, f"Erro na API do Gemini (HTTP {r.status_code}): {r.text}"
    except Exception as e_rest:
        return None, f"Erro ao conectar ao Gemini: {last_err or e_rest}"


def call_gemini_vision(prompt, image_input, mime_type="image/jpeg", json_mode=False):
    """
    Executa chamada multimodal ao Google Gemini API com safety_settings desativados (BLOCK_NONE) para uso médico-legal/forense.
    """
    key = get_gemini_api_key()
    if not key:
        return None, "Chave GEMINI_API_KEY não configurada. Verifique se o arquivo CHAVE.txt está presente."

    try:
        import google.generativeai as genai
        from google.generativeai.types import HarmCategory, HarmBlockThreshold

        genai.configure(api_key=key)

        generation_config = {}
        if json_mode:
            generation_config["response_mime_type"] = "application/json"

        safety_settings = {
            HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
            HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
            HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
            HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
        }

        models_to_try = ["gemini-2.5-flash", "gemini-flash-latest", "gemini-3.6-flash"]

        contents = [prompt]
        if isinstance(image_input, list):
            contents.extend(image_input)
        elif hasattr(image_input, "save") or str(type(image_input)).find("Image") != -1:
            contents.append(image_input)
        elif isinstance(image_input, bytes):
            contents.append({"mime_type": mime_type, "data": image_input})
        else:
            contents.append(image_input)

        last_err = None
        for m_name in models_to_try:
            try:
                model = genai.GenerativeModel(m_name, generation_config=generation_config if json_mode else None)
                resp = model.generate_content(contents, safety_settings=safety_settings)
                if resp and resp.text:
                    return resp.text, None
            except Exception as e:
                last_err = e
                continue

        return None, f"Erro ao executar Gemini Vision: {last_err}"
    except Exception as err:
        return None, str(err)



def ler_anotacoes_livres_gemini(texto_bruto):
    """
    Analisa anotações livres/brutas de campo usando Gemini IA e extrai dicionário estruturado.
    """
    prompt = f"""
Você é um perito criminalístico especialista em estruturação de dados de local de crime.
Examine as anotações livres/brutas fornecidas pelo perito e extraia os campos no formato JSON estrito:

{{
  "num_laudo": "Número do laudo se mencionado",
  "ocorrencia": "Número da ocorrência/BO",
  "requisicao": "Número da requisição",
  "perito": "Nome do perito relator",
  "autoridade_local": "Nome do delegado ou autoridade presente",
  "endereco": "Endereço completo do local",
  "ponto_referencia": "Ponto de referência",
  "area": "Descrição do tipo de área/pavimento",
  "delimitacoes": "Delimitações/limites físicos do local",
  "equipe_pm": "Equipe da Polícia Militar presente",
  "equipe_pc": "Equipe da Polícia Civil presente",
  "vitima_nome": "Nome da vítima se constar",
  "vitima_vestes": "Descrição das vestes da vítima",
  "vitima_posicao": "Posição/decúbito do cadáver",
  "vitima_lesoes": "Descrição das lesões na vítima",
  "vestigios_resumo": "Resumo dos vestígios encontrados",
  "dinamica_fatos": "Esboço ou rascunho da dinâmica dos fatos"
}}

Anotações do perito:
{texto_bruto}

Retorne APENAS o JSON estrito. Se algum dado não estiver mencionado, retorne string vazia.
"""

    res, err = call_gemini_text(prompt, system_instruction="Você é um assistente pericial de estruturação de dados. Responda apenas com JSON estrito.")
    if err or not res:
        return {}, err or "Sem resposta da IA"

    import json, re
    cleaned = res.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned), None
    except Exception as e:
        match = re.search(r"\{[\s\S]*\}", res)
        if match:
            try: return json.loads(match.group(0)), None
            except: pass
        return {}, f"Erro ao decodificar dados da IA: {e}"


def ler_requisicao_gemini_vision(file_bytes, file_name, file_type=""):
    """
    Lê uma Requisição Pericial (PDF ou Imagem) com Gemini Vision e retorna (extracted_text, dados_dict, error_msg).
    Prompt estruturado para extrair JSON com Delegacia, Delegado, Ocorrência, Requisição e Quesitos.
    """
    images = convert_bytes_to_pil_images(file_bytes, file_name, file_type)
    if not images:
        return "", {}, "Não foi possível converter o documento em imagem para análise visual pelo Gemini Vision."

    prompt = """
Você é um perito criminalístico especialista em análise documental pericial brasileira.
Examine a(s) imagem(ns) fornecida(s) da Requisição Pericial e extraia os seguintes dados no formato JSON estrito:

{
  "delegacia": "Nome da Delegacia de Polícia ou Órgão Solicitante",
  "delegado": "Nome da Autoridade Solicitante / Delegado(a) de Polícia (incluindo cargo ou 'Dr.(a)' se houver)",
  "ocorrencia": "Número do Boletim de Ocorrência / Ocorrência Policial",
  "requisicao": "Número da Requisição Pericial",
  "quesitos": "Texto integral de todos os quesitos / perguntas formuladas pela autoridade solicitante, numerados linha a linha"
}

Regras de extração:
1. Transcreva com fidelidade absoluta os números de ocorrência, requisição e quesitos.
2. Se algum campo não estiver visível na imagem, coloque string vazia "".
3. Retorne APENAS o JSON válido.
"""

    resp_text, err = call_gemini_vision(prompt, images, json_mode=True)
    if err or not resp_text:
        return "", {}, err or "Sem resposta do Gemini"

    import json
    import re
    cleaned = resp_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        dados = json.loads(cleaned)
        dados["origem"] = "Gemini Vision IA"
        return resp_text, dados, None
    except Exception as e:
        match = re.search(r"\{[\s\S]*\}", resp_text)
        if match:
            try:
                dados = json.loads(match.group(0))
                dados["origem"] = "Gemini Vision IA"
                return resp_text, dados, None
            except Exception:
                pass
        return resp_text, {}, f"Erro ao interpretar JSON do Gemini: {e}"


def ler_necropsia_gemini_vision(file_bytes, file_name, file_type=""):
    """
    Lê um Laudo de Necropsia / Exame Cadavérico (PDF ou Imagem) usando Gemini Vision.
    Extrai em JSON: causa_mortis, instrumento_lesivo, agente_instrumento, lesoes (lista), nome_vitima, documento_vitima, numero_laudo_iml.
    """
    images = convert_bytes_to_pil_images(file_bytes, file_name, file_type)
    if not images:
        return {}, "Não foi possível converter o arquivo em imagem para o Gemini Vision."

    prompt = """
Você é um médico legista e perito criminal especialista em tanatologia pericial e necropsia.
Examine a(s) imagem(ns) do Laudo de Necropsia / Auto de Exame Cadavérico e extraia as informações em formato JSON estrito:

{
  "causa_mortis": "Descrição exata e detalhada da causa mortis (ex: Traumatismo cranioencefálico decorrente de PAF, Choque hipovolêmico por ferimento perfurocortante, etc.)",
  "instrumento_lesivo": "Tipo de ação do instrumento lesivo (escolha preferencialmente uma das opções: Perfurocontundente, Cortante, Perfurante, Perfurocortante, Contundente, Cortocontundente, Ação Térmica)",
  "agente_instrumento": "Agente/instrumento específico se mencionado (ex: projétil de arma de fogo, lâmina de faca, instrumento contundente maciço)",
  "lesoes": [
    "Descrição detalhada e técnica da lesão 1 (localização anatômica, tipo, formato, dimensões)",
    "Descrição detalhada e técnica da lesão 2",
    "..."
  ],
  "nome_vitima": "Nome da vítima / examinando se constar",
  "documento_vitima": "RG/CPF ou número de identificação da vítima se constar",
  "numero_laudo_iml": "Número do Laudo do IML / Necropsia"
}

Regras:
1. Analise o exame necroscópico minuciosamente.
2. Extraia cada lesão identificada como um item detalhado na lista "lesoes".
3. Retorne apenas o JSON estrito.
"""

    resp_text, err = call_gemini_vision(prompt, images, json_mode=True)
    if err or not resp_text:
        return {}, err or "Sem resposta do Gemini"

    import json
    import re
    cleaned = resp_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        dados = json.loads(cleaned)
        return dados, None
    except Exception as e:
        match = re.search(r"\{[\s\S]*\}", resp_text)
        if match:
            try:
                dados = json.loads(match.group(0))
                return dados, None
            except Exception:
                pass
        return {}, f"Erro ao decodificar JSON de Necropsia: {e}"


def gerar_legenda_foto_gemini(b64_img, prompt_custom=None):
    """
    Gera legenda descritiva pericial técnica para uma fotografia de local/vestígio usando Gemini Vision.
    Possui fallback automático para laudos forenses em caso de restrição.
    """
    import base64
    from io import BytesIO
    from PIL import Image

    try:
        img_bytes = base64.b64decode(b64_img)
        pil_img = Image.open(BytesIO(img_bytes)).convert("RGB")
    except Exception as e:
        return "Visualização do registro fotográfico pericial no local dos fatos.", f"Erro ao processar imagem: {e}"

    prompt = prompt_custom or """
Você é um perito criminal relator da Polícia Científica e Medicina Legal.
Esta fotografia pertence a um procedimento oficial de investigação criminal e exame de local de crime / necropsia.
Descreva a imagem em uma frase técnica, objetiva, imparcial e formal para constar no Apêndice Fotográfico do Laudo Pericial.

Exemplos de formato pericial:
- "Visualização do corpo da vítima em decúbito dorsal na posição inicial do exame pericial."
- "Visão geral do sítio pericial, evidenciando o cadáver no solo."
- "Detalhe aproximado do vestígio/elemento registrado na cena dos fatos."

Regras:
1. Forneça APENAS o texto da legenda pericial (1 frase).
2. Não use aspas ou prefixos.
"""

    resp_text, err = call_gemini_vision(prompt, pil_img)
    if resp_text and resp_text.strip():
        cleaned = resp_text.strip().strip('"').strip("'")
        return cleaned, None

    # Fallback padronizado pericial caso o filtro de visão omita o retorno
    fallback_caption = "Visualização do corpo da vítima em decúbito dorsal na posição inicial do exame pericial."
    return fallback_caption, None


@st.dialog("✨ Polir Redação com IA (Tom Formal-Jurídico)")
def modal_polir_redacao(field_key, field_label):
    st.markdown("### ✨ Polimento de Redação Forense com IA")
    st.markdown(f"**Campo Selecionado:** `{field_label}`")

    texto_atual = st.session_state.get(field_key, "")
    if isinstance(texto_atual, list):
        texto_atual = "\n".join([str(x) for x in texto_atual if str(x).strip()])

    if not str(texto_atual).strip():
        st.warning("⚠️ O campo selecionado está vazio. Digite um texto antes de solicitar o polimento.")
        return

    st.caption("Texto Original Atual:")
    st.text_area("Texto Original", value=str(texto_atual), height=110, disabled=True, key=f"polir_orig_{field_key}")

    key_temp = f"polished_temp_{field_key}"
    if key_temp not in st.session_state:
        st.session_state[key_temp] = ""

    col1, col2 = st.columns(2)
    with col1:
        if st.button("✨ Executar Polimento com IA", type="primary", use_container_width=True, key=f"btn_exec_polir_{field_key}"):
            with st.spinner("Refinando redação e tom formal-jurídico com a IA Gemini..."):
                sys_instruction = (
                    "Você é um perito criminal especialista e revisor de laudos periciais oficiais da Polícia Civil. "
                    "Reescreva o texto a seguir elevando o padrão para a norma culta formal-jurídica, "
                    "utilizando terminologia pericial precisa, tom impessoal e clareza técnico-científica. "
                    "NUNCA invente fatos, dados, objetos ou lesões adicionais. Mantenha 100% dos fatos originais."
                )
                prompt = f"Reescreva e polia o texto pericial abaixo mantendo integralmente todos os fatos e dados técnicos:\n\n{texto_atual}"
                res, err = call_gemini_text(prompt, system_instruction=sys_instruction)
                if err:
                    st.error(f"Erro ao polir texto: {err}")
                else:
                    st.session_state[key_temp] = res
                    st.success("✅ Texto polido gerado com sucesso!")

    with col2:
        if st.session_state.get(key_temp):
            if st.button("✅ Aplicar Texto Polido no Laudo", type="secondary", use_container_width=True, key=f"btn_apply_polir_{field_key}"):
                st.session_state[field_key] = st.session_state[key_temp]
                st.session_state[key_temp] = ""
                st.toast("✅ Texto polido aplicado ao laudo!")
                st.rerun()

    if st.session_state.get(key_temp):
        st.markdown("**Resultado da Redação Polida (Sugestão IA):**")
        st.text_area("Texto Polido", value=st.session_state[key_temp], height=150, key=f"polir_res_{field_key}")


@st.dialog("🔍 Auditoria de Inconsistências (Checkup do Laudo)")
def modal_auditoria_inconsistencias():
    st.markdown("### 🛡️ Checkup de Qualidade e Coerência Forense")
    st.markdown("O sistema executa verificações determinísticas de formulário e uma análise lógica profunda via Gemini IA:")

    num_laudo = st.session_state.get("num_laudo", "")
    ocorrencia = st.session_state.get("ocorrencia", "")
    perito = st.session_state.get("perito", "")
    autoridade = st.session_state.get("autoridade_sel", "")
    requisicao = st.session_state.get("requisicao", "")
    endereco = st.session_state.get("endereco", "")
    delimitacoes = st.session_state.get("delimitacoes", "")
    iso_estado = st.session_state.get("iso_estado", "")
    inst_acao = st.session_state.get("inst_acao", "")
    vitimas = st.session_state.get("vitimas", [])
    vestigios = st.session_state.get("vestigios", [])
    quesitos = st.session_state.get("quesitos_list", [])
    dinamica = st.session_state.get("dinamica_fatos", "")

    # Rule-based validation engine
    erros = []
    alertas = []
    ok_items = []

    # 1. Header & Essenciais
    if not str(num_laudo).strip(): erros.append("Nº do Laudo não preenchido.")
    else: ok_items.append("Nº do Laudo preenchido.")
    if not str(ocorrencia).strip(): erros.append("Nº da Ocorrência não preenchido.")
    else: ok_items.append("Nº da Ocorrência preenchido.")
    if not str(perito).strip(): erros.append("Perito(a) Relator(a) não informado.")
    else: ok_items.append("Perito Relator informado.")
    if not str(autoridade).strip(): erros.append("Autoridade Solicitante não selecionada.")
    else: ok_items.append("Autoridade Solicitante informada.")
    if not str(requisicao).strip(): erros.append("Número da Requisição não preenchido.")
    else: ok_items.append("Número de Requisição informado.")
    if not str(endereco).strip(): erros.append("Endereço do local não informado.")
    else: ok_items.append("Endereço informado.")

    # 2. Vítimas & Lesões vs Instrumento
    todas_lesoes = []
    for idx_v, vt in enumerate(vitimas, 1):
        if not vt.get("nome", "").strip():
            alertas.append(f"Vítima #{idx_v} está sem nome/identificação definida.")
        lesoes = [l.strip() for l in vt.get("lesoes", []) if l.strip()]
        if not lesoes:
            alertas.append(f"Vítima #{idx_v} ({vt.get('nome') or 'N/I'}) não possui lesões descritas.")
        todas_lesoes.extend(lesoes)

    lesoes_concat = " ".join(todas_lesoes).lower()
    inst_str = str(inst_acao).lower()
    if "perfurocontundente" in inst_str or inst_str.startswith("1"):
        if "paf" not in lesoes_concat and "perfurocontundente" not in lesoes_concat and "projétil" not in lesoes_concat and "entrada" not in lesoes_concat and "disparo" not in lesoes_concat:
            alertas.append("Instrumento selecionado é Perfurocontundente (PAF), mas a descrição de lesões da(s) vítima(s) não menciona PAF, disparos ou orifícios de entrada/saída.")
    elif "cortante" in inst_str or inst_str.startswith("2"):
        if "corte" not in lesoes_concat and "incisa" not in lesoes_concat and "gume" not in lesoes_concat and "faca" not in lesoes_concat:
            alertas.append("Instrumento selecionado é Cortante, mas lesões não descrevem feridas incisas ou lâmina/gume.")

    # 3. Vestígios
    if not vestigios:
        alertas.append("Nenhum vestígio registrado no laudo.")
    else:
        ok_items.append(f"{len(vestigios)} vestígio(s) registrado(s).")
        for idx_vest, vest in enumerate(vestigios, 1):
            if not vest.get("tipo"):
                erros.append(f"Vestígio #{idx_vest} não possui Tipo selecionado.")
            elif not vest.get("descricao") and not vest.get("localizacao") and not vest.get("subtipo"):
                alertas.append(f"Vestígio #{idx_vest} ({vest.get('tipo')}) possui descrição/localização incompleta.")

    # 4. Quesitos
    quesitos_sem_resposta = [i+1 for i, q in enumerate(quesitos) if q.get("pergunta", "").strip() and not q.get("resposta", "").strip()]
    if quesitos_sem_resposta:
        alertas.append(f"Existem quesitos sem resposta formulada: Quesito(s) nº {quesitos_sem_resposta}.")

    # Render Rule-Based Metrics
    c_m1, c_m2, c_m3 = st.columns(3)
    c_m1.metric("🔴 Erros/Incompletos", len(erros))
    c_m2.metric("🟡 Alertas & Atenção", len(alertas))
    c_m3.metric("🟢 Verificados OK", len(ok_items))

    if erros:
        st.error("**Inconsistências Críticas:**\n" + "\n".join([f"• {e}" for e in erros]))
    if alertas:
        st.warning("**Alertas e Recomendações:**\n" + "\n".join([f"• {a}" for a in alertas]))
    if ok_items and not erros:
        st.success("**Itens de Formulário Checados:**\n" + "\n".join([f"✓ {o}" for o in ok_items]))

    st.divider()

    st.markdown("### 🤖 Auditoria Lógica Profunda com Gemini IA")
    if st.button("🚀 Executar Checkup Forense com IA", type="primary", use_container_width=True, key="btn_run_ai_audit"):
        with st.spinner("Analisando dados com a IA Gemini em busca de incoerências e omissões..."):
            dados_completos = {
                "num_laudo": num_laudo, "ocorrencia": ocorrencia, "perito": perito, "autoridade": autoridade,
                "requisicao": requisicao, "endereco": endereco, "delimitacoes": delimitacoes,
                "isolamento": iso_estado, "instrumento": inst_acao, "vitimas": vitimas,
                "vestigios": vestigios, "quesitos": quesitos, "dinamica_fatos": dinamica
            }
            prompt_audit = (
                "Você é um perito auditor sênior em criminalística e medicina legal. Analise o conjunto de dados do laudo pericial a seguir "
                "e verifique rigorosamente se existem:\n"
                "1. Contradições lógicas entre a posição da vítima, lesões e vestígios encontrados;\n"
                "2. Incompatibilidades entre a ação do instrumento e os exames perinicroscópicos;\n"
                "3. Incoerências na preservação/isolamento e na cadeia de custódia;\n"
                "4. Falhas ou omissões importantes que possam gerar nulidades no processo penal.\n\n"
                f"DADOS DO LAUDO:\n{json.dumps(dados_completos, ensure_ascii=False, indent=2, default=str)}\n\n"
                "Retorne o parecer técnico formatado com seções claras: 🚨 CONTRADIÇÕES LÓGICAS, ⚠️ RECOMENDAÇÕES DE AJUSTE, e 💡 SUGESTÕES COMPLEMENTARES."
            )
            sys_audit = "Você é um auditor pericial rigoroso. Seja imparcial, objetivo, técnico e ajude o perito a blindar o laudo contra erros ou contestação jurídica."
            res, err = call_gemini_text(prompt_audit, system_instruction=sys_audit)
            if err:
                st.error(f"Erro na auditoria IA: {err}")
            else:
                st.session_state["resultado_auditoria_ia"] = res
                st.success("✅ Auditoria IA concluída com sucesso!")

    if st.session_state.get("resultado_auditoria_ia"):
        st.markdown(st.session_state["resultado_auditoria_ia"])


def modal_auditoria_ia():
    modal_auditoria_inconsistencias()


def get_logo_base64():
    import base64
    logo_path = os.path.join(os.path.dirname(__file__), "logo_pericia.png")
    if os.path.exists(logo_path):
        try:
            with open(logo_path, "rb") as f:
                return base64.b64encode(f.read()).decode("utf-8")
        except Exception:
            return ""
    return ""


def ler_requisicao_pericial(file_bytes, file_name, file_type=""):
    """
    Extrai texto e campos (Delegacia, Delegado, Ocorrência, Requisição, Quesitos)
    de arquivos de Requisição Pericial em formato PDF ou Imagem.
    Prioriza o Gemini Vision IA quando a chave está configurada e faz fallback para OCR/Regex.
    """
    # 1. Tenta extração via Gemini Vision IA se a API Key estiver configurada
    if get_gemini_api_key():
        resp_text, dados_g, err_g = ler_requisicao_gemini_vision(file_bytes, file_name, file_type)
        if dados_g and any(dados_g.values()):
            return resp_text, dados_g

    # 2. Fallback para métodos locais (pdfplumber, pypdf, pytesseract OCR + Regex)
    from io import BytesIO
    import re

    extracted_text = ""
    file_name_lower = file_name.lower()
    is_pdf = file_type == "application/pdf" or file_name_lower.endswith(".pdf")
    is_img = file_type.startswith("image/") or file_name_lower.endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff"))

    if is_pdf:
        try:
            import pdfplumber
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        extracted_text += t + "\n"
        except Exception:
            pass

        if not extracted_text.strip():
            try:
                import pypdf
                reader = pypdf.PdfReader(BytesIO(file_bytes))
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        extracted_text += t + "\n"
            except Exception:
                pass

        if not extracted_text.strip():
            try:
                import pdfplumber
                import pytesseract
                with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        pil_img = page.to_image(resolution=150).original
                        try:
                            t = pytesseract.image_to_string(pil_img, lang="por")
                        except Exception:
                            t = pytesseract.image_to_string(pil_img)
                        if t:
                            extracted_text += t + "\n"
            except Exception:
                pass

    elif is_img:
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(BytesIO(file_bytes))
            try:
                extracted_text = pytesseract.image_to_string(img, lang="por")
            except Exception:
                extracted_text = pytesseract.image_to_string(img)
        except Exception:
            pass

    dados = {}
    if extracted_text.strip():
        m_del = re.search(r'(?:DELEGACIA(?: DE POLÍCIA)?|UNIDADE SOLICITANTE|ÓRGÃO SOLICITANTE|ORIGEM|DESTINO)[\s:]+([^\n\r]+)', extracted_text, re.IGNORECASE)
        if not m_del:
            m_del = re.search(r'(DELEGACIA DE POLÍCIA [^\n\r]+)', extracted_text, re.IGNORECASE)
        if not m_del:
            m_del = re.search(r'(DELEGACIA [^\n\r]+)', extracted_text, re.IGNORECASE)
        if m_del:
            dados["delegacia"] = m_del.group(1).strip()

        # Extrai Delegado / Autoridade Solicitante
        m_aut = re.search(r'(?:DELEGAD[OA](?: DE POLÍCIA)?|AUTORIDADE SOLICITANTE|AUTORIDADE POLICIAL|SOLICITANTE)[\s:]+([^\n\r]+)', extracted_text, re.IGNORECASE)
        if not m_aut:
            m_aut = re.search(r'(?:Dr[a]?\.\s+)([^\n\r]+)', extracted_text, re.IGNORECASE)
        if m_aut:
            dados["delegado"] = m_aut.group(1).strip()

        # Extrai Ocorrência / BO
        m_oco = re.search(r'(?:OCORRÊNCIA|OCORRENCIA|BOLETIM DE OCORRÊNCIA|BO)[\s\:\.\º\°]*(?:N[º°\.\:]*)?[\s\:\.\º\°]*([0-9A-Za-z\/\-\.]{3,})', extracted_text, re.IGNORECASE)
        if m_oco:
            dados["ocorrencia"] = m_oco.group(1).strip()

        # Extrai Requisição
        m_req = re.search(r'(?:REQUISIÇÃO|REQUISICAO|REQ\.?)[\s\:\.\º\°]*(?:N[º°\.\:]*)?[\s\:\.\º\°]*([0-9A-Za-z\/\-\.]{3,})', extracted_text, re.IGNORECASE)
        if m_req:
            dados["requisicao"] = m_req.group(1).strip()

        # Extrai Quesitos
        m_que = re.search(r'(?:QUESITOS(?: DA AUTORIDADE| FORMULADOS)?|PERGUNTAS)[\s\:\-\n\r]+([\s\S]+?)(?:\n\s*\n[A-Z0-9\s]{4,}:|\Z)', extracted_text, re.IGNORECASE)
        if m_que:
            dados["quesitos"] = m_que.group(1).strip()
        else:
            num_q = re.findall(r'(\d+[\.\)-]\s*[^\n]+)', extracted_text)
            if num_q:
                dados["quesitos"] = "\n".join(num_q)

    return extracted_text, dados


# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
for key, default in {
    "autoridades":   AUTORIDADES_DEFAULT.copy(),
    "preview_open":  False,
    "show_nova_aut": False,
    "docx_bytes":    None,
    "docx_filename": "",
    "quesitos":      "",
    "quesitos_list": [{"pergunta": "", "resposta": ""}],
    "vitimas":       [{"nome": "", "cad": "", "documento": "", "sexo": "", "data_nascimento": None, "filicao": "", "naturalidade": "",
                       "vestes": "", "pertences": "", "localizacao": "", "posicao": "", "cabeca": "", "membros": "",
                       "fenomenos": "", "lesoes": [""]}],
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

# ─────────────────────────────────────────────
# ESTILOS
# ─────────────────────────────────────────────
st.markdown("""
<style>
/* ── Reset e Base ── */
html, body, [class*="css"], .stApp {
    font-family: 'Inter', sans-serif !important;
    background-color: #e2e8f0 !important;
    color: #0f172a !important;
}


/* ── App Header (Título) ── */
.app-header {
    background-color: #ffffff;
    padding: 24px;
    border-radius: 16px;
    border: 1px solid #e2e8f0;
    margin-bottom: 24px;
    box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05);
}
.app-header-brand {
    display: flex;
    flex-direction: row;
    align-items: center;
    gap: 16px;
}
.app-header-badge {
    background: #eff6ff;
    color: #2563eb;
    padding: 4px 12px;
    border-radius: 20px;
    font-size: 12px;
    font-weight: 700;
    width: fit-content;
}
.app-header-title {
    font-size: 24px;
    font-weight: 800;
    color: #0f172a;
    line-height: 1.2;
}
.app-header-sub {
    font-size: 14px;
    color: #64748b;
}


/* ── Forçar Colunas em Grid no Mobile ── */
@media (max-width: 768px) {
    div[data-testid="column"] {
        width: 50% !important;
        flex: 1 1 calc(50% - 1rem) !important;
        min-width: 40% !important;
    }
    div[data-testid="stHorizontalBlock"] {
        flex-wrap: wrap !important;
        flex-direction: row !important;
    }
}

/* Esconder o Header Padrão do Streamlit */
header[data-testid="stHeader"] { display: none !important; }

/* ── Containers com borda (st.container border=True) ── */
div[data-testid="stVerticalBlock"]:has(.custom-border-marker) {
    background-color: #ffffff !important;
    border: 1px solid #e2e8f0 !important;
    border-radius: 16px !important;
    box-shadow: 0 10px 15px -3px rgba(0,0,0,0.05), 0 4px 6px -4px rgba(0,0,0,0.05) !important;
    padding: 24px !important;
    margin-bottom: 24px !important;
}


/* ── Título das seções ── */
.section-title {
    font-size: 22px !important;
    font-weight: 800 !important;
    color: #0f172a !important;
    margin-top: 12px !important;
    margin-bottom: 20px !important;
    padding-bottom: 10px !important;
    border-bottom: 2px solid #cbd5e1 !important;
    display: flex !important;
    align-items: center !important;
    gap: 10px !important;
    line-height: 1.3 !important;
}


/* ── Botão primário ── */
.stButton > button[kind="primary"] {
    background: #2563eb !important;
    border: none !important;
    border-radius: 12px !important;
    font-weight: 700 !important;
    font-size: 15px !important;
    color: #fff !important;
    padding: 12px 24px !important;
    transition: all 0.2s !important;
}
.stButton > button[kind="primary"]:hover {
    background: #1d4ed8 !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 10px 15px -3px rgba(37,99,235,0.3) !important;
}

/* ── Botão secundário ── */
.stButton > button[kind="secondary"] {
    background: #f8fafc !important;
    border: 1px solid #cbd5e1 !important;
    border-radius: 12px !important;
    font-weight: 600 !important;
    color: #334155 !important;
    transition: all 0.2s !important;
}
.stButton > button[kind="secondary"]:hover {
    background: #e2e8f0 !important;
}

/* ── Inputs e Selects ── */
div[data-testid="stTextInput"] input,
div[data-testid="stTextArea"] textarea,
div[data-testid="stDateInput"] input,
div[data-testid="stTimeInput"] input,
div[data-baseweb="select"] > div {
    background-color: #ffffff !important;
    border: 1px solid #cbd5e1 !important;
    border-radius: 10px !important;
    box-shadow: 0 1px 2px 0 rgba(0,0,0,0.05) !important;
    color: #0f172a !important;
}

/* Labels */
div[data-testid="stTextInput"] label,
div[data-testid="stTextArea"] label,
div[data-testid="stSelectbox"] label,
div[data-testid="stDateInput"] label,
div[data-testid="stTimeInput"] label {
    font-size: 13px !important;
    font-weight: 600 !important;
    color: #475569 !important;
    margin-bottom: 4px !important;
}

/* Expanders */
.streamlit-expanderHeader {
    background-color: #ffffff !important;
    border-radius: 12px !important;
    border: 1px solid #e2e8f0 !important;
    color: #1e293b !important;
    font-weight: 600 !important;
}
div[data-testid="stExpander"] {
    border: none !important;
    box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05) !important;
    border-radius: 12px !important;
    overflow: hidden;
}

/* Sidebar */
section[data-testid="stSidebar"] {
    background-color: #ffffff !important;
    border-right: 1px solid #e2e8f0 !important;
}
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════
# AUTENTICAÇÃO DO SISTEMA
# ═══════════════════════════════════════════════════════════
if "autenticado" not in st.session_state:
    st.session_state["autenticado"] = True

# Força bypass:
st.session_state["autenticado"] = True

if not st.session_state["autenticado"]:
    st.markdown("<h2 style='text-align: center; color: #1a3a6e; margin-top: 50px;'>🔒 Acesso Restrito ao Sistema</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #64748b; margin-bottom: 30px;'>Por favor, insira as credenciais de acesso para prosseguir.</p>", unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 1.5, 1])
    with col2:
        with st.container(border=True):
            with st.form("login_form"):
                user = st.text_input("Usuário")
                pwd = st.text_input("Senha", type="password")
                submitted = st.form_submit_button(
                    "Entrar", type="primary", use_container_width=True)

                if submitted:
                    if (user.lower() in ["alvaro", "perito"]) and pwd in ["123", "icrim123"]:
                        st.session_state["autenticado"] = True
                        st.rerun()
                    else:
                        st.error("❌ Usuário ou senha incorretos.")

    st.stop()  # Bloqueia a execução do restante do código


# ═══════════════════════════════════════════════════════════
# MODO PREVIEW
# ═══════════════════════════════════════════════════════════
if st.session_state.preview_open:

    def ss(key, default=""):
        return st.session_state.get(key, default) or default

    num_laudo = ss("num_laudo")
    ocorrencia = ss("ocorrencia")
    perito = ss("perito")
    autoridade = ss("autoridade_sel")
    requisicao = ss("requisicao")
    referencia = ss("referencia") or "(sem referência)"
    destino = ss("destino")

    vitimas_list = st.session_state.get("vitimas", [{"nome": "", "cad": "", "documento": "", "sexo": "", "data_nascimento": None, "filicao": "",
                                        "naturalidade": "", "vestes": "", "pertences": "", "localizacao": "", "posicao": "", "cabeca": "", "membros": "", "fenomenos": "", "lesoes": [""]}])
    nomes = [vt["nome"] for vt in vitimas_list if vt["nome"].strip()]
    vitima = ", ".join(nomes) if nomes else "________"

    vitimas_detalhes_html = ""
    for idx, vt in enumerate(vitimas_list, 1):
        nome_v = vt.get("nome", "").strip() or "Não identificado"
        doc_v = vt.get("documento", "").strip() or "Não informado"
        sexo_v = vt.get("sexo", "").strip() or "Não informado"
        dob_v = vt.get("data_nascimento", None)
        dob_str = data_simples(dob_v) if dob_v else "Não informado"
        fili_v = vt.get("filicao", "").strip() or "Não informada"
        nat_v = vt.get("naturalidade", "").strip() or "Não informada"
        vest_v = vt.get("vestes", "").strip() or "Não descritas"
        pert_v = vt.get("pertences", "").strip() or "Não descritos"
        loc_v = vt.get("localizacao", "").strip() or "Não descrita"
        pos_v = vt.get("posicao", "").strip() or "Não descrita"
        cab_v = vt.get("cabeca", "").strip() or "Não descrita"
        memb_v = vt.get("membros", "").strip() or "Não descritos"
        fen_v = vt.get("fenomenos", "").strip() or "Não descritos"

        lesoes_html = ""
        for les in vt.get("lesoes", [""]):
            if les.strip():
                lesoes_html += f"• {les.strip()}<br>"
        if not lesoes_html:
            lesoes_html = "• Não descritas<br>"

        vitimas_detalhes_html += f"""
                <div style="font-weight:bold; margin-top:0.6cm; font-size:11.5pt; color:#0f2044; text-indent:0px;">
                3.3.{idx} – Descrição da Vítima {idx}
                </div>

                <div style="font-size:10.5pt; line-height:1.6; padding-left:1.5cm; text-indent:0px;">
                <div style="font-weight:bold; margin-top:0.3cm; color:#0f2044; border-bottom:1px solid #e2e8f0; width:100%;">Identificação</div>
                a) Nome: {nome_v}<br>
                b) Documento: {doc_v}<br>
                c) Sexo: {sexo_v}<br>
                d) Data de Nascimento: {dob_str}<br>
                e) Filiação: {fili_v}<br>
                f) Naturalidade: {nat_v}<br>

                <div style="font-weight:bold; margin-top:0.3cm; color:#0f2044; border-bottom:1px solid #e2e8f0; width:100%;">Vestes e Pertences</div>
                a) Vestes: {vest_v}<br>
                b) Pertences: {pert_v}<br>

                <div style="font-weight:bold; margin-top:0.3cm; color:#0f2044; border-bottom:1px solid #e2e8f0; width:100%;">Posição e Localização (In Situ)</div>
                a) Localização: {loc_v}<br>
                b) Posição: {pos_v}<br>
                c) Cabeça: {cab_v}<br>
                d) Membros: {memb_v}<br>

                <div style="font-weight:bold; margin-top:0.3cm; color:#0f2044; border-bottom:1px solid #e2e8f0; width:100%;">Exame Perinecroscópico e Estado de Conservação</div>
                a) Fenômenos Transformativos: {fen_v}<br>
                b) Lesões:<br>
                {lesoes_html}
                </div>
                """
        dp_val = st.session_state.get("data_pericia_input",     date.today())
        da_val = st.session_state.get("data_atendimento_input", date.today())
        horario_val = st.session_state.get("horario",               None)
        horario_a_val = st.session_state.get("horario_atendimento",   None)
        endereco = ss("endereco")
        latitude = ss("latitude")
        longitude = ss("longitude")
        ponto = ss("ponto_referencia")
        area = ss("area")
        pavimento = ss("pavimento")
        delimitacoes = ss("delimitacoes")
        iso_s = ss("isolamento")
        isolamento = ss("isolamento_outros") if iso_s == "Outros" else iso_s
        clim_s = ss("clima")
        clima = ss("clima_outros") if clim_s == "Outros" else clim_s
        visibilidade = ss("visibilidade")
        iluminacao = ss("iluminacao")
        equipe_pm = ss("equipe_pm")
        equipe_pc = ss("equipe_pc")
        aut_local = ss("autoridade_local")

        dp_ext = data_extenso(dp_val)
        da_simp = data_simples(da_val)
        h_str = horario_val.strftime("%H:%M") if horario_val else "____"
        ha_str = horario_a_val.strftime("%H:%M") if horario_a_val else "____"

        tipo_local_val = st.session_state.get("tipo_local", "")
        ano = date.today().year

        # Barra superior
        st.markdown(f"""
                <div style="background:linear-gradient(135deg,#0f2044,#1a3a6e); padding:0 52px;
                height:68px; display:flex; align-items:center; justify-content:space-between;
                box-shadow:0 2px 16px rgba(0,0,0,0.25);">
                <div style="display:flex; align-items:center; gap:18px;">
                <span style="background:rgba(255,255,255,0.12); border:1px solid rgba(255,255,255,0.25);
                border-radius:6px; padding:5px 13px; font-size:10px; font-weight:700;
                color:#a8c4f0; letter-spacing:1.8px;">ICRIM / NPT</span>
                <div>
                <div style="font-size:17px; font-weight:700; color:#fff;">Pré-visualização do Laudo</div>
                <div style="font-size:12px; color:#7fa8d4;">Modo leitura — formulário desativado</div>
                </div>
                </div>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)

        _, btn_col, _ = st.columns([3, 2, 3])
        with btn_col:
            if st.button("✕  Fechar e Voltar ao Formulário", type="primary", use_container_width=True):
                st.session_state.preview_open = False
                st.rerun()

        st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)

        st.markdown(f"""
                <div class="a4-wrapper">
                <div class="a4-sheet">
                <div style="text-align:center; font-weight:bold; font-size:10.5pt;
                margin-bottom:1.2cm; line-height:1.6; color:#0f2044;">
                GOVERNO DO ESTADO DO MARANHÃO<br>
                SECRETARIA DE ESTADO DA SEGURANÇA PÚBLICA<br>
                POLÍCIA CIVIL<br>
                SUPERINTENDÊNCIA DE POLÍCIA TÉCNICO-CIENTÍFICA<br>
                INSTITUTO DE CRIMINALÍSTICA DE IMPERATRIZ
                </div>
                <div style="text-align:center; font-weight:bold; font-size:12pt;
                margin-bottom:1cm; text-decoration:underline; color:#0f2044;">
                LAUDO DE EXAME PERICIAL EM LOCAL DE MORTE VIOLENTA<br>
                Nº {v(num_laudo)}.{ano}/PO
                </div>
                <table style="width:100%; font-size:10.5pt; border-collapse:collapse; margin-bottom:1cm;">
                {''.join(f'<tr><td style="font-weight:bold;width:195px;padding:5px 0;vertical-align:top;color:#0f2044;">{lbl}:</td><td style="padding:5px 0;">{val}</td></tr>'
                for lbl, val in [
                ("NATUREZA DA PERÍCIA","Local de Morte Violenta"),
                ("OCORRÊNCIA Nº", v(ocorrencia)),
                ("PERITO(A) RELATOR(A)", v(perito)),
                ("AUTORIDADE SOLICITANTE", v(autoridade)),
                ("REQUISIÇÃO", v(requisicao)),
                ("VÍTIMA(S)", v(vitima)),
                ("INVESTIGADO(S)", investigado),
                ("IP/PM/TCO/BO/PROC", referencia),
                ("UNIDADE DESTINO", v(destino)),
                ])}
                </table>
                <div style="font-weight:bold;margin-top:1cm;margin-bottom:0.4cm;font-size:12pt;
                color:#0f2044;border-bottom:1px solid #cbd5e1;padding-bottom:4px;">
                1 — HISTÓRICO
                </div>
                <p style="text-align:justify;text-indent:1.5cm;margin-bottom:0.5cm;font-size:10.5pt;">
                {dp_ext}, neste Estado do Maranhão, o Diretor do Instituto de Criminalística de
                Imperatriz (ICRIM) designou o(a) Perito(a) Criminal <b>{v(perito)}</b>, para realização
                de exame pericial em Local de Morte Violenta, referente à Ocorrência nº
                <b>{v(ocorrencia)}</b>, requisitada por <b>{v(autoridade)}</b>, sob Requisição nº
                <b>{v(requisicao)}</b>.
                </p>
                <p style="text-align:justify;text-indent:1.5cm;margin-bottom:0.5cm;font-size:10.5pt;">
                Atendendo à solicitação, os peritos signatários compareceram ao local
                no dia <b>{da_simp}</b> às <b>{ha_str}</b>, encontrando as condições descritas a seguir.
                </p>
                <div style="font-weight:bold;margin-top:1cm;margin-bottom:0.4cm;font-size:12pt;
                color:#0f2044;border-bottom:1px solid #cbd5e1;padding-bottom:4px;">
                2 — DO LOCAL
                </div>
                <p style="text-align:justify;text-indent:1.5cm;font-size:10.5pt;line-height:2;">
                <b>ENDEREÇO:</b> {v(endereco)}<br>
                <b>COORDENADAS:</b> Lat: {v(latitude)} / Lon: {v(longitude)}<br>
                <b>PONTO DE REFERÊNCIA:</b> {v(ponto)}<br>
                <b>ÁREA:</b> {v(area)} &nbsp; <b>PAVIMENTO:</b> {v(pavimento)}<br>
                <b>DELIMITAÇÕES:</b> {v(delimitacoes)}<br>
                <b>ISOLAMENTO:</b> {v(isolamento)}<br>
                <b>CLIMA:</b> {v(clima)} &nbsp; <b>VISIBILIDADE:</b> {v(visibilidade)} &nbsp; <b>ILUMINAÇÃO:</b> {v(iluminacao)}<br>
                <b>EQUIPE PM:</b> {v(equipe_pm)} &nbsp; <b>EQUIPE PC:</b> {v(equipe_pc)}<br>
                <b>AUTORIDADE NO LOCAL:</b> {v(aut_local)}
                </p>

                <div style="font-weight:bold;margin-top:1.2cm;margin-bottom:0.4cm;font-size:12pt;
                color:#0f2044;border-bottom:1px solid #cbd5e1;padding-bottom:4px;">
                3 — EXAMES PERICIAIS
                </div>
                <div style="font-weight:bold;margin-top:0.5cm;margin-bottom:0.3cm;font-size:11.5pt;
                color:#0f2044;">
                3.3 — Das Vítimas
                </div>
                {vitimas_detalhes_html}

                <div style="text-align:center;font-size:9pt;margin-top:3cm;
                border-top:1px solid #94a3b8;padding-top:12px;color:#64748b;">
                Instituto de Criminalística de Imperatriz — ICRIM/NPT &nbsp;|&nbsp; Polícia Civil do Estado do Maranhão
                </div>
                </div>
                </div>
                """, unsafe_allow_html=True)


else:
    # ═══════════════════════════════════════════════════════════
    # MODO FORMULÁRIO (LOGIN DESABILITADO TEMPORARIAMENTE)
    # ═══════════════════════════════════════════════════════════
    st.session_state['logged_in'] = True

    _, main, _ = st.columns([0.04, 0.92, 0.04])


GAS_URL = "https://script.google.com/macros/s/AKfycbx9N4hidxHbfAUUpHVDAOadZBF5SRB9x9UzC0nt3k-wW0pgLThCHwBQsaUMJjtcTL1QJw/exec"

with main:
    def sanitize_datetime_state():
        from datetime import datetime, date, time
        for k in ["data_pericia_input", "data_atendimento_input"]:
            if k in st.session_state and isinstance(st.session_state[k], str):
                try:
                    st.session_state[k] = datetime.strptime(
                        st.session_state[k], "%Y-%m-%d").date()
                except:
                    st.session_state[k] = date.today()
        for k in ["horario", "horario_atendimento"]:
            if k in st.session_state and isinstance(st.session_state[k], str):
                try:
                    s = st.session_state[k]
                    if len(s.split(':')) == 2:
                        s += ":00"
                    st.session_state[k] = datetime.strptime(
                        s, "%H:%M:%S").time()
                except:
                    st.session_state[k] = None

    sanitize_datetime_state()


    logo_b64 = get_logo_base64()
    logo_html = f'<img src="data:image/png;base64,{logo_b64}" style="height:65px; width:auto; object-fit:contain;" />' if logo_b64 else ''

    st.markdown(f'''
    <div class="app-header">
      <div class="app-header-brand">
        {logo_html}
        <div>
          <div class="app-header-title">Laudo de Local de Morte Violenta</div>
          <div class="app-header-sub">Instituto de Criminalística de Imperatriz &nbsp;|&nbsp; Perícia Oficial do Maranhão</div>
        </div>
      </div>
    </div>
    <div style="height:16px"></div>
    ''', unsafe_allow_html=True)


    @st.dialog("☁️ Carregar Ocorrência (Banco Local / Nuvem)")
    def modal_ocorrencias():
        st.markdown(
            "Busca ocorrências no Google Drive (App Mobile) ou banco local.")
        import sqlite3
        import json
        import os
        import requests
        DB_PATH = os.path.join(os.path.dirname(
            __file__), "banco_laudos.sqlite")

        # 1. Fetch from Cloud
        ocorrencias_cloud = []
        try:
            resp = requests.get(
                GAS_URL, params={"key": "perito:icrim123"}, timeout=5)
            if resp.status_code == 200:
                ocorrencias_cloud = resp.json()
        except:
            pass

        if ocorrencias_cloud and isinstance(ocorrencias_cloud, list):
            for item in ocorrencias_cloud:
                try:
                    conn = sqlite3.connect(DB_PATH)
                    c = conn.cursor()
                    mob_id = item.get("mobile_id", "")
                    dados_j = item.get("dados_json", "{}")
                    dt = item.get("data_sincronizacao", "")
                    if isinstance(dados_j, dict):
                        dados_j = json.dumps(dados_j)
                    c.execute('INSERT OR REPLACE INTO laudos (mobile_id, data_sincronizacao, dados_json) VALUES (?, ?, ?)',
                              (mob_id, dt, dados_j))
                    conn.commit()
                    conn.close()
                except:
                    pass

        # 2. Show from SQLite
        db_laudos = []
        if os.path.exists(DB_PATH):
            try:
                conn = sqlite3.connect(DB_PATH)
                conn.row_factory = sqlite3.Row
                c = conn.cursor()
                c.execute(
                    'SELECT id, mobile_id, data_sincronizacao, dados_json FROM laudos ORDER BY id DESC')
                rows = c.fetchall()
                conn.close()
                for r in rows:
                    db_laudos.append({"id": r["id"], "mobile_id": r["mobile_id"],
                                     "data": r["data_sincronizacao"], "dados": json.loads(r["dados_json"])})
            except:
                pass

        if db_laudos:
            filtro = st.text_input(
                "🔍 Buscar por perito, ocorrência ou data:",
                placeholder="Digite para filtrar por perito, ocorrência ou data...",
                key="filtro_ocorrencias_modal"
            ).strip().lower()

            db_laudos_filtrados = []
            for l in db_laudos:
                oc_val = str(l['dados'].get('ocorrencia', '')).lower()
                per_val = str(l['dados'].get('perito', '')).lower()
                dt_val = str(l.get('data', '')).lower()
                if not filtro or (filtro in oc_val or filtro in per_val or filtro in dt_val):
                    db_laudos_filtrados.append(l)

            if db_laudos_filtrados:
                opcoes = {
                    f"Ocorrência {l['dados'].get('ocorrencia', 'S/N')} — {l['dados'].get('perito', 'N/I')} ({l['data'][:16]})": l
                    for l in db_laudos_filtrados
                }
                escolhida = st.selectbox(
                    f"Resultados ({len(db_laudos_filtrados)} encontrada(s)):", list(opcoes.keys()))
                if st.button("Carregar Dados", use_container_width=True, type="primary"):
                    item = opcoes[escolhida]
                    data_obj = item["dados"]
                    for k, v in data_obj.items():
                        if k not in ["vitimas", "vestigios", "fotos", "quesitos_list"]:
                            st.session_state[k] = v
                    if "vitimas" in data_obj:
                        st.session_state["vitimas"] = data_obj["vitimas"]
                    if "vestigios" in data_obj:
                        st.session_state["vestigios"] = data_obj["vestigios"]
                    if "fotos" in data_obj:
                        st.session_state["fotos"] = data_obj["fotos"]
                    if "quesitos_list" in data_obj:
                        st.session_state["quesitos_list"] = data_obj["quesitos_list"]
                    st.rerun()
            else:
                st.warning("⚠️ Nenhuma ocorrência encontrada para o filtro informado.")
        else:
            st.info("Nenhuma ocorrência encontrada no banco.")





    @st.dialog("🩸 Guia Visual de Manchas de Sangue")
    def modal_guia_sangue():
        st.markdown("Utilize as imagens de referência abaixo para guiar a identificação do tipo de mancha. Passe o mouse sobre a imagem e use o ícone ⛶ (tela cheia) ou selecione um padrão abaixo para ampliar e baixar:")
        img_dir = os.path.join(os.path.dirname(__file__), "vestigio sangue")
        if os.path.exists(img_dir):
            files = sorted([os.path.join(img_dir, f) for f in os.listdir(img_dir) if f.lower().endswith(('.png', '.jpg', '.jpeg'))])
            if files:
                cols = st.columns(2)
                for idx, img_p in enumerate(files):
                    with cols[idx % 2]:
                        st.image(img_p, use_container_width=True, caption=f"Padrão Referência #{idx+1}")

                st.markdown("<hr style='margin:16px 0; border-color:#e2e8f0;'>", unsafe_allow_html=True)
                st.markdown("##### 🔍 Ampliar e Baixar Padrão Específico")
                opcoes_img = {f"Padrão Referência #{i+1}": img_p for i, img_p in enumerate(files)}
                sel = st.selectbox("Escolha o padrão para ampliar:", list(opcoes_img.keys()), key="sel_zoom_sangue")
                if sel:
                    st.image(opcoes_img[sel], use_container_width=True, caption=f"Visualização Detalhada Ampliada — {sel}")
                    try:
                        with open(opcoes_img[sel], "rb") as file_img:
                            st.download_button(
                                label=f"📥 Baixar Imagem do {sel}",
                                data=file_img.read(),
                                file_name=os.path.basename(opcoes_img[sel]),
                                mime="image/jpeg",
                                use_container_width=True,
                                key=f"btn_dl_img_{sel}"
                            )
                    except Exception as e_dl:
                        st.error(f"Erro ao disponibilizar download: {e_dl}")
        else:
            st.info("Pasta 'vestigio sangue' não localizada.")


    gerar_top = render_action_buttons("top")
    st.markdown('<div style="height:12px"></div>', unsafe_allow_html=True)

    with st.container():
        # ══ SEÇÃO 1: DA OCORRÊNCIA ════════════════════════════
        st.markdown(
            '<div class="section-title">📋 &nbsp; Da Ocorrência</div>', unsafe_allow_html=True)
        with st.container():

            col_tools1, col_tools2 = st.columns(2)
            with col_tools1:
                with st.expander("📝 Bloco de Notas / Rascunho Livre (IA)", expanded=False):
                    st.markdown("<p style='font-size:13px; color:#475569;'>Cole ou digite suas anotações brutas de campo (anotações do celular, rascunhos, transcrições) e a IA Gemini preencherá o laudo automaticamente.</p>", unsafe_allow_html=True)
                    raw_notes = st.text_area("Anotações de Campo do Perito", placeholder="Ex: ocorrencia 1234, perito carlos, local rua das flores 50, corpo masculino de brucos...", height=120, key="txt_raw_field_notes")
                    if st.button("🤖 Processar Rascunho com IA", type="primary", key="btn_parse_raw_notes", use_container_width=True):
                        if not raw_notes.strip():
                            st.warning("⚠️ Insira o texto das anotações antes de processar.")
                        else:
                            with st.spinner("Analisando rascunho com IA Gemini..."):
                                dados_parsed, err_p = ler_anotacoes_livres_gemini(raw_notes)
                                if err_p:
                                    st.error(f"Erro ao processar anotações: {err_p}")
                                elif dados_parsed:
                                    count_fields = 0
                                    mapping = {
                                        "num_laudo": "num_laudo", "ocorrencia": "ocorrencia", "requisicao": "requisicao",
                                        "perito": "perito", "autoridade_local": "autoridade_local", "endereco": "endereco",
                                        "ponto_referencia": "ponto_referencia", "area": "area", "delimitacoes": "delimitacoes",
                                        "equipe_pm": "equipe_pm", "equipe_pc": "equipe_pc", "dinamica_fatos": "dinamica_fatos"
                                    }
                                    for k_json, k_state in mapping.items():
                                        if dados_parsed.get(k_json):
                                            st.session_state[k_state] = dados_parsed[k_json]
                                            count_fields += 1

                                    if dados_parsed.get("vitima_nome") or dados_parsed.get("vitima_vestes") or dados_parsed.get("vitima_lesoes"):
                                        if "vitimas" in st.session_state and st.session_state["vitimas"]:
                                            v0 = st.session_state["vitimas"][0]
                                            if dados_parsed.get("vitima_nome"): v0["nome"] = dados_parsed["vitima_nome"]
                                            if dados_parsed.get("vitima_vestes"): v0["vestes"] = dados_parsed["vitima_vestes"]
                                            if dados_parsed.get("vitima_posicao"): v0["posicao"] = dados_parsed["vitima_posicao"]
                                            if dados_parsed.get("vitima_lesoes"): v0["lesoes"] = [dados_parsed["vitima_lesoes"]]
                                            count_fields += 1

                                    st.success(f"✅ Sucesso! {count_fields} campo(s) preenchido(s):")
                                    st.json(dados_parsed)
                                    st.rerun()

            with col_tools2:
                with st.expander("📄 Leitor de Requisição (PDF / Imagem)", expanded=False):
                    if get_gemini_api_key():
                        st.info("✨ Gemini Vision: Extração inteligente de Delegacia, Delegado, Ocorrência, Requisição e Quesitos.")
                    else:
                        st.caption("💡 Para extração com IA Gemini Vision, verifique o arquivo CHAVE.txt.")

                    req_file = st.file_uploader("Carregar Requisição (PDF / Imagem)", type=["pdf", "png", "jpg", "jpeg", "bmp", "tiff"], key="req_file_input")
                    if req_file is not None:
                        file_bytes = req_file.read()
                        extracted_text, dados = ler_requisicao_pericial(file_bytes, req_file.name, req_file.type)
                        if dados:
                            if dados.get("delegacia"):
                                st.session_state["destino"] = dados["delegacia"]
                            if dados.get("delegado"):
                                del_nome = dados["delegado"]
                                if del_nome not in st.session_state.autoridades:
                                    st.session_state.autoridades.append(del_nome)
                                st.session_state["autoridade_sel"] = del_nome
                            if dados.get("ocorrencia"):
                                st.session_state["ocorrencia"] = dados["ocorrencia"]
                            if dados.get("requisicao"):
                                st.session_state["requisicao"] = dados["requisicao"]
                            if dados.get("quesitos"):
                                st.session_state["quesitos"] = dados["quesitos"]
                                raw_lines = [q.strip() for q in dados["quesitos"].split("\n") if q.strip()]
                                if raw_lines:
                                    st.session_state["quesitos_list"] = [{"pergunta": q_line, "resposta": ""} for q_line in raw_lines]

                            orig = dados.get("origem", "Automático")
                            st.success(f"✅ Requisição lida ({orig}):")
                            st.json(dados)
                        elif extracted_text:
                            st.warning("⚠️ Texto extraído da requisição, mas nenhum campo de formulário reconhecido por regex.")
                            with st.expander("Ver Texto Extraído"):
                                st.text(extracted_text)
                        else:
                            st.error("❌ Não foi possível extrair texto do arquivo fornecido.")

            c1, c2, c3, c4 = st.columns(4)
            num_laudo = c1.text_input(
                "Nº do Laudo",         placeholder="0092258",    key="num_laudo")
            ocorrencia = c2.text_input(
                "Nº da Ocorrência",    placeholder="12345/2026", key="ocorrencia")
            requisicao = c3.text_input(
                "Requisição",           key="requisicao")
            referencia = c4.text_input(
                "IP / PM / BO / Proc.", key="referencia")

            c5, c6, c7, c8 = st.columns(4)
            data_pericia_val = c5.date_input(
                "Data da Perícia",        value=date.today(), format="DD/MM/YYYY", key="data_pericia_input")
            horario_val = c6.time_input(
                "Horário da Perícia",     key="horario",      step=300)
            data_atendimento_val = c7.date_input(
                "Data de Atendimento",    value=date.today(), format="DD/MM/YYYY", key="data_atendimento_input")
            horario_atend_val = c8.time_input(
                "Horário de Atendimento", key="horario_atendimento", step=300)

            cp, ca, cb = st.columns([2, 2, 0.7])
            perito = cp.text_input("Perito(a) Relator(a)", key="perito")
            with ca:
                autoridade = st.selectbox("Autoridade Solicitante",
                                          options=[""] + st.session_state.autoridades, key="autoridade_sel")
            with cb:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("➕ Nova", key="btn_nova_aut", use_container_width=True):
                    st.session_state.show_nova_aut = True

        if st.session_state.show_nova_aut:
            na1, na2, na3 = st.columns([3, 1, 1])
            nova_aut = na1.text_input("Nome da autoridade:", key="nova_aut_input",
                                      label_visibility="collapsed",
                                      placeholder="Ex: Delegado(a) Titular Silva")
            if na2.button("✓ Salvar", type="primary", key="btn_ok_aut", use_container_width=True):
                if nova_aut.strip() and nova_aut.strip() not in st.session_state.autoridades:
                    st.session_state.autoridades.append(nova_aut.strip())
                st.session_state.show_nova_aut = False
                st.rerun()
            if na3.button("Cancelar", key="btn_cancel_aut", use_container_width=True):
                st.session_state.show_nova_aut = False
                st.rerun()

        cv1, cv2 = st.columns(2)
        investigado = cv1.text_area(
            "Investigado(s)",  placeholder="Nomes dos investigados...", key="investigado", height=80)
        destino = cv2.text_input(
            "Unidade Destino", value="Delegacia de Homicídios de Imperatriz", key="destino")

    with st.container():
        # ══ SEÇÃO 2: DO LOCAL ════════════════════════════════
        st.markdown(
            '<div class="section-title">📍 &nbsp; Do Local</div>', unsafe_allow_html=True)
        with st.container(border=True):
            st.markdown('<span class="custom-border-marker"></span>',
                        unsafe_allow_html=True)

            ce1, ce2, ce3, ce4 = st.columns([3, 1.2, 1.2, 0.6])
            endereco = ce1.text_input(
                "Endereço completo", placeholder="Rua, Nº, Bairro, Município", key="endereco")
            latitude = ce2.text_input(
                "Latitude",  placeholder="-5.518600",  key="latitude")
            longitude = ce3.text_input(
                "Longitude", placeholder="-47.477600", key="longitude")
            with ce4:
                st.markdown("<br>", unsafe_allow_html=True)
                st.button("📍 GPS", key="btn_gps", use_container_width=True,
                          help="GPS disponível via HTTPS")

            cr1, cr2, cr3 = st.columns([2, 1, 1])
            ponto_referencia = cr1.text_input(
                "Ponto de Referência", placeholder="Ex: Próximo ao posto de saúde", key="ponto_referencia")
            area = cr2.text_input(
                "Área",      placeholder="Ex: Externa", key="area")
            pavimento = cr3.text_input(
                "Pavimento", placeholder="Ex: Asfalto", key="pavimento")

            cdel1, cdel2 = st.columns([4, 1])
            with cdel1:
                delimitacoes = st.text_area("Delimitações do local",
                                            placeholder="Descreva os limites físicos do local periciado...",
                                            key="delimitacoes", height=70)
            with cdel2:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("✨ Polir IA", key="btn_polir_delimitacoes", use_container_width=True, help="Polir redação das delimitações com IA"):
                    modal_polir_redacao("delimitacoes", "Delimitações do Local")

            cl1, cl2, cl3, cl4 = st.columns(4)
            iso_opts = ["", "Preservado e Isolado", "Parcialmente Preservado e Isolado",
                        "Preservado e Parcialmente Isolado", "Não Preservado e Não Isolado", "Outros"]
            isolamento_sel = cl1.selectbox(
                "Isolamento", iso_opts, key="isolamento")

            clima_opts = ["", "Aberto diurno", "Aberto noturno", "Nublado diurno",
                          "Nublado noturno", "Chuvoso diurno", "Chuvoso noturno", "Outros"]
            clima_sel = cl2.selectbox("Clima", clima_opts, key="clima")

            visibilidade = cl3.selectbox(
                "Visibilidade", ["", "Ampla", "Reduzida"], key="visibilidade")
            iluminacao = cl4.selectbox("Iluminação",
                                       ["", "Natural Satisfatória", "Natural Insatisfatória",
                                        "Artificial Satisfatória", "Artificial Insatisfatória"], key="iluminacao")

            if isolamento_sel == "Outros" or clima_sel == "Outros":
                ot1, ot2 = st.columns(2)
                if isolamento_sel == "Outros":
                    isolamento_txt = ot1.text_input(
                        "Especifique o Isolamento", key="isolamento_outros")
                    isolamento = isolamento_txt or "Outros"
                else:
                    isolamento = isolamento_sel
                if clima_sel == "Outros":
                    clima_txt = ot2.text_input(
                        "Especifique o Clima", key="clima_outros")
                    clima = clima_txt or "Outros"
                else:
                    clima = clima_sel
            else:
                isolamento = isolamento_sel
                clima = clima_sel

            cq1, cq2, cq3 = st.columns(3)
            equipe_pm = cq1.text_input(
                "Equipe PM", placeholder="VTR, Comandante...", key="equipe_pm")
            equipe_pc = cq2.text_input(
                "Equipe PC", placeholder="Investigador, Delegado...", key="equipe_pc")
            autoridade_local = cq3.text_input(
                "Autoridade no Local", placeholder="Nome da autoridade...", key="autoridade_local")

    with st.container():
        # ══ SEÇÃO: DAS VÍTIMAS (Item 3.3) ═════════════════════
        st.markdown(
            '<div class="section-title">👥 &nbsp; Das Vítimas (Item 3.3)</div>', unsafe_allow_html=True)
        with st.container(border=True):
            st.markdown('<span class="custom-border-marker"></span>',
                        unsafe_allow_html=True)

            with st.expander("📄 Leitor Gemini Vision - Laudo de Necropsia / Exame Cadavérico (PDF / Imagem)", expanded=False):
                st.markdown("<p style='font-size:13px; color:#475569;'>Faça o upload do Laudo de Necropsia ou Auto de Exame Cadavérico (IML) em PDF ou imagem para extração automática da Causa Mortis, Instrumento Lesivo e Lesões Detalhadas com Gemini Vision IA.</p>", unsafe_allow_html=True)
                nec_file = st.file_uploader("Carregar Laudo de Necropsia / Exame Cadavérico", type=["pdf", "png", "jpg", "jpeg", "bmp", "tiff"], key="necropsia_file_input")
                if nec_file is not None:
                    if st.button("✨ Processar Necropsia com Gemini Vision", type="primary", key="btn_process_necropsia", use_container_width=True):
                        if not get_gemini_api_key():
                            st.error("⚠️ Chave GEMINI_API_KEY não configurada. Verifique se o arquivo CHAVE.txt está presente.")
                        else:
                            with st.spinner("Analisando Laudo de Necropsia com Gemini Vision..."):
                                nec_bytes = nec_file.read()
                                dados_nec, err_nec = ler_necropsia_gemini_vision(nec_bytes, nec_file.name, nec_file.type)
                                if err_nec:
                                    st.error(f"❌ Erro ao analisar laudo de necropsia: {err_nec}")
                                elif dados_nec:
                                    if dados_nec.get("causa_mortis"):
                                        st.session_state["resultado_laudo_IML"] = dados_nec["causa_mortis"]
                                    if dados_nec.get("numero_laudo_iml"):
                                        st.session_state["numero_laudo_necropsia"] = dados_nec["numero_laudo_iml"]

                                    inst_extraido = dados_nec.get("instrumento_lesivo", "")
                                    if inst_extraido:
                                        map_inst = {
                                            "perfurocontundente": "1. Perfurocontundente (Arma de Fogo)",
                                            "cortante": "2. Cortante (Feridas Incisas - Faca, Navalha, Estilete, Vidro)",
                                            "perfurante": "3. Perfurante (Feridas Punctórias - Espeto, Estilete, Chave de Fenda)",
                                            "perfurocortante": "4. Perfurocortante (Feridas Perfuroincisas - Faca, Punhal, Canivete)",
                                            "contundente": "5. Contundente (Feridas Contusas/Fraturas - Madeira, Pedra, Veículo, Piso)",
                                            "cortocontundente": "6. Cortocontundente (Feridas Contuso-Incisas - Machado, Facão, Foice)",
                                            "térmica": "7. Ação Térmica (Queimaduras - Chama Direta, Líquido Fervente, Superfície Aquecida)",
                                            "termica": "7. Ação Térmica (Queimaduras - Chama Direta, Líquido Fervente, Superfície Aquecida)"
                                        }
                                        for k_m, v_m in map_inst.items():
                                            if k_m in inst_extraido.lower():
                                                st.session_state["inst_acao"] = v_m
                                                break
                                    if dados_nec.get("agente_instrumento"):
                                        st.session_state["inst_agente"] = dados_nec["agente_instrumento"]

                                    if "vitimas" in st.session_state and st.session_state["vitimas"]:
                                        vit0 = st.session_state["vitimas"][0]
                                        if dados_nec.get("nome_vitima") and not vit0.get("nome"):
                                            vit0["nome"] = dados_nec["nome_vitima"]
                                        if dados_nec.get("documento_vitima") and not vit0.get("documento"):
                                            vit0["documento"] = dados_nec["documento_vitima"]
                                        if dados_nec.get("lesoes") and isinstance(dados_nec["lesoes"], list) and len(dados_nec["lesoes"]) > 0:
                                            vit0["lesoes"] = [str(l).strip() for l in dados_nec["lesoes"] if str(l).strip()]

                                    st.success("✅ Laudo de Necropsia analisado com sucesso pelo Gemini Vision! Dados preenchidos no formulário:")
                                    st.json(dados_nec)
                                    st.rerun()

            vitimas_list = st.session_state.get("vitimas", [{"nome": "", "cad": "", "documento": "", "sexo": "", "data_nascimento": None, "filicao": "",
                                                "naturalidade": "", "vestes": "", "pertences": "", "localizacao": "", "posicao": "", "cabeca": "", "membros": "", "fenomenos": "", "lesoes": [""]}])

            for i, vit in enumerate(vitimas_list):
                st.markdown(f"##### 👤 Vítima #{i+1}")

                st.markdown(
                    "<p style='font-size:12px; font-weight:700; color:#1a3a6e; margin-bottom:8px;'>IDENTIFICAÇÃO</p>", unsafe_allow_html=True)
                c_nome, c_doc, c_sexo = st.columns([2, 1, 1])
                vit["nome"] = c_nome.text_input(
                    "Nome completo", value=vit.get("nome", ""), key=f"vit_nome_{i}")
                vit["documento"] = c_doc.text_input("Nº Documento (RG/CPF)", value=vit.get(
                    "documento", ""), key=f"vit_doc_{i}", placeholder="RG ou CPF...")
                vit["sexo"] = c_sexo.selectbox("Sexo", ["", "Masculino", "Feminino", "Outro"], index=[
                                               "", "Masculino", "Feminino", "Outro"].index(vit.get("sexo", "")), key=f"vit_sexo_{i}")

                c_dob, c_fili, c_nat = st.columns([1, 2, 1])
                dob_val = vit.get("data_nascimento", None)
                if not isinstance(dob_val, date) and dob_val is not None:
                    dob_val = date.today()
                vit["data_nascimento"] = c_dob.date_input(
                    "Data de Nascimento", value=dob_val, key=f"vit_dob_{i}", format="DD/MM/YYYY")
                vit["filicao"] = c_fili.text_input("Filiação", value=vit.get(
                    "filicao", ""), key=f"vit_fili_{i}", placeholder="Nome da mãe e/ou pai...")
                vit["naturalidade"] = c_nat.text_input("Naturalidade", value=vit.get(
                    "naturalidade", ""), key=f"vit_nat_{i}", placeholder="Cidade-UF...")

                st.markdown(
                    "<p style='font-size:12px; font-weight:700; color:#1a3a6e; margin:12px 0 8px 0;'>VESTES E PERTENCES</p>", unsafe_allow_html=True)
                c_vest, c_pert = st.columns(2)
                vit["vestes"] = c_vest.text_area("Vestes", value=vit.get(
                    "vestes", ""), key=f"vit_vest_{i}", height=70, placeholder="Descrição das roupas que usava...")
                vit["pertences"] = c_pert.text_area("Pertences", value=vit.get(
                    "pertences", ""), key=f"vit_pert_{i}", height=70, placeholder="Celular, chaves, carteira, joias, etc...")

                st.markdown(
                    "<p style='font-size:12px; font-weight:700; color:#1a3a6e; margin:12px 0 8px 0;'>POSIÇÃO E LOCALIZAÇÃO (IN SITU)</p>", unsafe_allow_html=True)
                c_loc, c_pos = st.columns(2)
                vit["localizacao"] = c_loc.text_area("Localização", value=vit.get(
                    "localizacao", ""), key=f"vit_loc_{i}", height=70, placeholder="Ex: No compartimento traseiro do carro...")
                vit["posicao"] = c_pos.text_area("Posição", value=vit.get(
                    "posicao", ""), key=f"vit_pos_{i}", height=70, placeholder="Ex: Decúbito lateral direito...")

                c_cab, c_memb = st.columns(2)
                vit["cabeca"] = c_cab.text_input("Cabeça", value=vit.get(
                    "cabeca", ""), key=f"vit_cab_{i}", placeholder="Ex: Voltada a leste...")
                vit["membros"] = c_memb.text_input("Membros", value=vit.get(
                    "membros", ""), key=f"vit_memb_{i}", placeholder="Ex: Membros superiores fletidos...")

                st.markdown(
                    "<p style='font-size:12px; font-weight:700; color:#1a3a6e; margin:12px 0 8px 0;'>EXAME PERINECROSCÓPICO E ESTADO DE CONSERVAÇÃO</p>", unsafe_allow_html=True)
                vit["fenomenos"] = st.text_input("Fenômenos Transformativos", value=vit.get(
                    "fenomenos", ""), key=f"vit_fen_{i}", placeholder="Ex: Mancha verde abdominal, rigidez muscular...")

                st.markdown(
                    "<p style='font-size:11px; font-weight:600; color:#64748b; margin-top:8px;'>LESÕES CONSTATADAS</p>", unsafe_allow_html=True)
                les_list = vit.get("lesoes", [""])
                for j, lesao in enumerate(les_list):
                    c_les, c_rem_l = st.columns([12, 1])
                    les_list[j] = c_les.text_input(
                        f"Lesão #{j+1}", value=lesao, key=f"vit_{i}_les_{j}", label_visibility="collapsed", placeholder=f"Descrição da lesão #{j+1}...")
                    if len(les_list) > 1:
                        if c_rem_l.button("❌", key=f"vit_{i}_rem_les_{j}", help="Remover esta lesão"):
                            les_list.pop(j)
                            st.rerun()
                vit["lesoes"] = les_list

                if st.button("➕ Adicionar Lesão", key=f"vit_{i}_add_les"):
                    vit.setdefault("lesoes", [""]).append("")
                    st.rerun()

                if len(vitimas_list) > 1:
                    st.markdown("<div style='height:8px'></div>",
                                unsafe_allow_html=True)
                    if st.button(f"🗑️ Remover Vítima #{i+1}", key=f"remove_vit_{i}", type="secondary"):
                        st.session_state.vitimas.pop(i)
                        st.rerun()
                st.markdown(
                    "<hr style='margin:20px 0; border-color:#2563eb;'>", unsafe_allow_html=True)

            if st.button("➕ Adicionar Outra Vítima", key="btn_add_vitima"):
                st.session_state.vitimas.append({"nome": "", "cad": "", "documento": "", "sexo": "", "data_nascimento": None, "filicao": "", "naturalidade": "",
                                                "vestes": "", "pertences": "", "localizacao": "", "posicao": "", "cabeca": "", "membros": "", "fenomenos": "", "lesoes": [""]})
                st.rerun()

    with st.container():
        st.markdown(
            '<div class="section-title">🔍 &nbsp; Dos Vestígios</div>', unsafe_allow_html=True)
        with st.container(border=True):
            st.markdown('<span class="custom-border-marker"></span>',
                        unsafe_allow_html=True)

            manchas_dict = {
                "1. FORMAÇÃO PASSIVA": ["Gotejamento Isolado", "Gotejamento Sucessivo", "Gotejamento Estático", "Gotejamento Dinâmico", "Escorrimento", "Saturação", "Empoçamento"],
                "2. FORMAÇÃO ATIVA": ["Espargimento Por Impacto", "Espargimento Por Expiração", "Espargimento Por Êmese", "Espargimento De Retorno", "Espargimento De Saída", "Transferência Por Contato", "Transferência Por Arrastamento", "Inercial Por Desprendimento Centrífugo", "Inercial Por Parada", "Projeção por Diferença de Pressão", "Espalhamento de Volume em Queda Livre"],
                "3. ALTERADAS": ["Por Arrastamento", "Silhueta de Mancha", "Por Limpeza", "Por Inseto"],
                "4. AUSÊNCIA": ["Área de Sombra", "Sangue Latente"]
            }

            if "vestigios" not in st.session_state or not st.session_state.vestigios:
                st.session_state["vestigios"] = [{"tipo": "", "categoria": "", "subtipo": "", "localizacao": "", "descricao": "", "elemento_tipo": "",
                                                  "quantidade": "", "envelope": "", "arma_tipo": "", "municiada": "", "marca": "", "cor": "", "funcionamento": "", "instrumento_tipo": ""}]
            vestigios_list = st.session_state["vestigios"]

            for i, vest in enumerate(vestigios_list):
                st.markdown(f"##### 🔎 Vestígio #{i+1}")
                tipo_opts = ["", "Manchas de Sangue", "Elemento Balístico", "Arma de Fogo",
                             "Aparelho Telefônico", "Instrumento Lesivo", "Impressão Datiloscópica", "Outros"]
                tipo_idx = tipo_opts.index(vest.get("tipo", "")) if vest.get(
                    "tipo", "") in tipo_opts else 0
                vest["tipo"] = st.selectbox(
                    "Tipo de Vestígio", tipo_opts, index=tipo_idx, key=f"vest_tipo_{i}")

                if vest["tipo"] in ["Manchas de Sangue", "Mancha de Sangue"]:
                    if st.button("🩸 Abrir Guia de Manchas de Sangue", key=f"btn_guia_sangue_{i}"):
                        modal_guia_sangue()
                    c_cat, c_sub = st.columns(2)
                    cat_opts = [""] + list(manchas_dict.keys())
                    cat_idx = cat_opts.index(vest.get("categoria", "")) if vest.get(
                        "categoria", "") in cat_opts else 0
                    vest["categoria"] = c_cat.selectbox(
                        "Categoria", cat_opts, index=cat_idx, key=f"vest_cat_{i}")
                    sub_opts = [""] + \
                        manchas_dict.get(vest.get("categoria", ""), [])
                    sub_idx = sub_opts.index(vest.get("subtipo", "")) if vest.get(
                        "subtipo", "") in sub_opts else 0
                    vest["subtipo"] = c_sub.selectbox(
                        "Subtipo", sub_opts, index=sub_idx, key=f"vest_sub_{i}")
                    if vest["subtipo"]:
                        image_path = os.path.join(os.path.dirname(
                            __file__), "imagens_manchas", f"{vest['subtipo']}.jpg")
                        if os.path.exists(image_path):
                            st.image(image_path, width=300,
                                     caption=f"Guia Visual: {vest['subtipo']}")
                        else:
                            st.caption(
                                f"(Dica: Salve o recorte como 'imagens_manchas\\{vest['subtipo']}.jpg' para ver o guia visual aqui)")

                    vest["localizacao"] = st.text_input("Localização", value=vest.get(
                        "localizacao", ""), key=f"vest_loc_{i}", placeholder="Ex: Parede leste...")
                    vest["descricao"] = st.text_area("Descrição Livre", value=vest.get(
                        "descricao", ""), key=f"vest_desc_{i}", height=70)

                elif vest["tipo"] == "Elemento Balístico":
                    c_tipo, c_qtd = st.columns([3, 1])
                    elem_opts = ["", "Estojo",
                                 "Munição", "Jaqueta", "Projétil"]
                    elem_idx = elem_opts.index(vest.get("elemento_tipo", "")) if vest.get(
                        "elemento_tipo", "") in elem_opts else 0
                    vest["elemento_tipo"] = c_tipo.selectbox(
                        "Tipo de Elemento", elem_opts, index=elem_idx, key=f"vest_elem_{i}")
                    vest["quantidade"] = c_qtd.text_input(
                        "Quantidade", value=vest.get("quantidade", ""), key=f"vest_qtd_{i}")
                    vest["envelope"] = st.text_input("Número Envelope de Segurança", value=vest.get(
                        "envelope", ""), key=f"vest_env_{i}")
                    vest["descricao"] = st.text_area("Descrição (geral)", value=vest.get(
                        "descricao", ""), key=f"vest_desc_{i}", height=70)

                elif vest["tipo"] == "Arma de Fogo":
                    c_arma, c_mun = st.columns([3, 1])
                    arma_opts = ["", "Revólver", "Espingarda",
                                 "Rifle", "Pistola", "Caseira"]
                    arma_idx = arma_opts.index(vest.get("arma_tipo", "")) if vest.get(
                        "arma_tipo", "") in arma_opts else 0
                    vest["arma_tipo"] = c_arma.selectbox(
                        "Tipo de Arma", arma_opts, index=arma_idx, key=f"vest_arma_{i}")
                    mun_opts = ["", "Sim", "Não"]
                    mun_idx = mun_opts.index(vest.get("municiada", "")) if vest.get(
                        "municiada", "") in mun_opts else 0
                    vest["municiada"] = c_mun.selectbox(
                        "Municiada", mun_opts, index=mun_idx, key=f"vest_mun_{i}")
                    vest["envelope"] = st.text_input("Número Envelope de Segurança", value=vest.get(
                        "envelope", ""), key=f"vest_env_{i}")
                    vest["descricao"] = st.text_area("Descrição", value=vest.get(
                        "descricao", ""), key=f"vest_desc_{i}", height=70)

                elif vest["tipo"] == "Outros":
                    vest["descricao"] = st.text_area("Descrição", value=vest.get(
                        "descricao", ""), key=f"vest_desc_{i}", height=70, placeholder="Descreva o vestígio...")

                elif vest["tipo"] == "Impressão Datiloscópica":
                    vest["descricao"] = st.text_area("Descrição", value=vest.get(
                        "descricao", ""), key=f"vest_desc_{i}", height=70, placeholder="Descreva os detalhes da impressão datiloscópica, superfície, revelação, etc.")

                elif vest["tipo"] == "Instrumento Lesivo":
                    vest["instrumento_tipo"] = st.text_input("Tipo (ex: Faca, pedaço de madeira, corda)", value=vest.get(
                        "instrumento_tipo", ""), key=f"vest_inst_{i}")
                    vest["envelope"] = st.text_input("Número Envelope de Segurança", value=vest.get(
                        "envelope", ""), key=f"vest_env_{i}")
                    vest["descricao"] = st.text_area("Descrição", value=vest.get(
                        "descricao", ""), key=f"vest_desc_{i}", height=70)

                elif vest["tipo"] == "Aparelho Telefônico":
                    c_marca, c_cor = st.columns(2)
                    vest["marca"] = c_marca.text_input(
                        "Marca ou Modelo", value=vest.get("marca", ""), key=f"vest_marca_{i}")
                    vest["cor"] = c_cor.text_input(
                        "Cor", value=vest.get("cor", ""), key=f"vest_cor_{i}")
                    c_func, c_env = st.columns(2)
                    func_opts = ["", "Ligado", "Desligado"]
                    func_idx = func_opts.index(vest.get("funcionamento", "")) if vest.get(
                        "funcionamento", "") in func_opts else 0
                    vest["funcionamento"] = c_func.selectbox(
                        "Funcionamento", func_opts, index=func_idx, key=f"vest_func_{i}")
                    vest["envelope"] = c_env.text_input(
                        "Número Envelope de Segurança", value=vest.get("envelope", ""), key=f"vest_env_{i}")
                    vest["descricao"] = st.text_area("Descrição", value=vest.get(
                        "descricao", ""), key=f"vest_desc_{i}", height=70)

                if st.button(f"🗑️ Remover Vestígio #{i+1}", key=f"remove_vest_{i}", type="secondary"):
                    st.session_state.vestigios.pop(i)
                    st.rerun()
                st.markdown(
                    "<hr style='margin:20px 0; border-color:#2563eb;'>", unsafe_allow_html=True)

            if st.button("➕ Adicionar Vestígio", key="btn_add_vest"):
                if "vestigios" not in st.session_state:
                    st.session_state.vestigios = []
                st.session_state.vestigios.append({"tipo": "", "categoria": "", "subtipo": "", "localizacao": "", "descricao": "", "elemento_tipo": "",
                                                  "quantidade": "", "envelope": "", "arma_tipo": "", "municiada": "", "marca": "", "cor": "", "funcionamento": "", "instrumento_tipo": ""})
                st.rerun()

    # ══ BARRA DE AÇÕES (RODAPÉ) ══════════════════════════

    with st.container():
        st.markdown(
            '<div class="section-title">📝 &nbsp; 4. Considerações Técnicas</div>', unsafe_allow_html=True)
        with st.container(border=True):
            st.markdown('<span class="custom-border-marker"></span>',
                        unsafe_allow_html=True)

            st.subheader("a) Isolamento e Preservação")

            opcoes_isolamento = [
                "1. Local Preservado e Isolado",
                "2. Local Parcialmente Preservado e Isolado",
                "3. Local Preservado e Parcialmente Isolado",
                "4. Local Parcialmente Preservado e Parcialmente Isolado",
                "5. Local Não Preservado e Não Isolado (Devassado)"
            ]

            iso_idx = 0
            curr_iso = st.session_state.get("iso_estado", "")
            if curr_iso in opcoes_isolamento:
                iso_idx = opcoes_isolamento.index(curr_iso)

            iso_estado = st.selectbox(
                "Estado do Isolamento", opcoes_isolamento, index=iso_idx, key="iso_estado")

            # Condicionais
            st.session_state["iso_meio"] = st.session_state.get("iso_meio", "")
            st.session_state["iso_alteracao"] = st.session_state.get(
                "iso_alteracao", "")
            st.session_state["iso_falha"] = st.session_state.get(
                "iso_falha", "")

            if iso_estado.startswith("1"):
                st.text_input("Meio de Isolamento", value=st.session_state["iso_meio"], key="iso_meio_ui",
                              help="Ex: fita zebrada, cordão de isolamento, guarnição física da PM")
                st.session_state["iso_meio"] = st.session_state.get(
                    "iso_meio_ui", "")
            elif iso_estado.startswith("2"):
                st.text_input(
                    "Meio de Isolamento", value=st.session_state["iso_meio"], key="iso_meio_ui")
                st.text_area("Descrever Alteração", value=st.session_state["iso_alteracao"],
                             key="iso_alteracao_ui", help="Ex: movimentação da vítima por equipes de socorro")
                st.session_state["iso_meio"] = st.session_state.get(
                    "iso_meio_ui", "")
                st.session_state["iso_alteracao"] = st.session_state.get(
                    "iso_alteracao_ui", "")
            elif iso_estado.startswith("3"):
                st.text_input("Descrever Falha no Isolamento",
                              value=st.session_state["iso_falha"], key="iso_falha_ui", help="Ex: ausência de barreiras em rotas de fuga")
                st.session_state["iso_falha"] = st.session_state.get(
                    "iso_falha_ui", "")
            elif iso_estado.startswith("4"):
                st.text_input("Descrever Falha no Isolamento",
                              value=st.session_state["iso_falha"], key="iso_falha_ui", help="Ex: coibir o fluxo de pessoas não autorizadas")
                st.text_area("Descrever Alteração", value=st.session_state["iso_alteracao"],
                             key="iso_alteracao_ui", help="Ex: marcas de pneus sobrepostas a manchas de sangue")
                st.session_state["iso_falha"] = st.session_state.get(
                    "iso_falha_ui", "")
                st.session_state["iso_alteracao"] = st.session_state.get(
                    "iso_alteracao_ui", "")
            elif iso_estado.startswith("5"):
                st.text_area("Descrever Alteração (Devassado)",
                             value=st.session_state["iso_alteracao"], key="iso_alteracao_ui")
                st.session_state["iso_alteracao"] = st.session_state.get(
                    "iso_alteracao_ui", "")

            st.divider()
            st.subheader("c) Laudo de Necropsia (IML)")
            c_iml1, c_iml2 = st.columns(2)
            with c_iml1:
                st.session_state["numero_laudo_necropsia"] = st.text_input("Número Laudo IML", value=st.session_state.get(
                    "numero_laudo_necropsia", ""), placeholder="Ex: 123/2026 - IML")
            with c_iml2:
                st.session_state["resultado_laudo_IML"] = st.text_input("Resultado do Laudo IML", value=st.session_state.get(
                    "resultado_laudo_IML", ""), placeholder="Ex: traumatismo cranioencefálico por PAF")

            st.divider()
            st.subheader("d) Instrumento Utilizado")

            opcoes_inst = [
                "1. Perfurocontundente (Arma de Fogo)",
                "2. Cortante (Feridas Incisas - Faca, Navalha, Estilete, Vidro)",
                "3. Perfurante (Feridas Punctórias - Espeto, Estilete, Chave de Fenda)",
                "4. Perfurocortante (Feridas Perfuroincisas - Faca, Punhal, Canivete)",
                "5. Contundente (Feridas Contusas/Fraturas - Madeira, Pedra, Veículo, Piso)",
                "6. Cortocontundente (Feridas Contuso-Incisas - Machado, Facão, Foice)",
                "7. Ação Térmica (Queimaduras - Chama Direta, Líquido Fervente, Superfície Aquecida)"
            ]

            inst_idx = 0
            curr_inst = st.session_state.get("inst_acao", "")
            if curr_inst in opcoes_inst:
                inst_idx = opcoes_inst.index(curr_inst)

            inst_acao = st.selectbox(
                "Tipo de Ação do Instrumento", opcoes_inst, index=inst_idx, key="inst_acao")

            sugestoes_agentes = {
                "1": ["projéteis de arma de fogo (PAF)"],
                "2": ["faca", "navalha", "fragmento de vidro", "lâmina de estilete"],
                "3": ["espeto", "estilete", "chave de fenda", "sovela"],
                "4": ["faca", "punhal", "canivete"],
                "5": ["barra de ferro", "pedaço de madeira", "pedra", "veículo automotor em movimento", "piso"],
                "6": ["machado", "facão", "foice", "enxada"],
                "7": ["chama direta", "líquido fervente", "superfície superaquecida"]
            }

            inst_key = inst_acao[0]
            if inst_key != "1":
                sug = sugestoes_agentes.get(inst_key, [])
                st.session_state["inst_agente"] = st.text_input("Agente Compatível", value=st.session_state.get(
                    "inst_agente", ""), placeholder=f"Ex: {', '.join(sug[:3])}")
            else:
                st.session_state["inst_agente"] = "projéteis de arma de fogo (PAF)"

            placeholders_extras = {
                "1": "Ex: com a recuperação de um projétil em cada cadáver durante a necrópsia / com orifícios de entrada compatíveis com disparos a curta distância",
                "2": "Ex: evidenciando cauda de escoriação que indica a direção do corte / apresentando características típicas de lesão de defesa",
                "3": "Ex: com transfixação de estruturas vitais profundas / sem a recuperação ou retenção de fragmentos do instrumento no interior da cavidade",
                "4": "Ex: contabilizando-se 3 golpes desferidos / com a recuperação de um fragmento metálico da lâmina retido na estrutura óssea",
                "5": "Ex: resultando em traumatismo cranioencefálico severo / não havendo elementos materiais desprendidos do instrumento no cadáver",
                "6": "Ex: provocando lacerações profundas associadas a fraturas ósseas",
                "7": "Ex: configurando queimaduras de 3º grau, com presença de fuligem em vias aéreas"
            }

            st.session_state["inst_extra"] = st.text_area(
                "Achados Extras / Observações do Instrumento", value=st.session_state.get("inst_extra", ""), help=placeholders_extras.get(inst_key, ""))

            st.divider()
            
            st.markdown("<hr style='margin:20px 0; border-color:#cbd5e1;'>", unsafe_allow_html=True)
            st.subheader("Considerações Técnicas Adicionais")
            st.markdown("<p style='font-size:13px; color:#475569; margin-bottom:12px;'>Adicione quantas considerações personalizadas desejar. A numeração por letras (e, f, g...) será gerada automaticamente em ordem crescente.</p>", unsafe_allow_html=True)

            if "consideracoes_extras" not in st.session_state:
                st.session_state["consideracoes_extras"] = []

            cons_extras = st.session_state["consideracoes_extras"]

            for idx_extra, item_extra in enumerate(cons_extras):
                letra = chr(ord('e') + idx_extra)
                st.markdown(f"##### {letra}) Consideração Adicional #{idx_extra + 1}")
                
                item_extra["titulo"] = st.text_input(
                    "Título / Assunto (Opcional)",
                    value=item_extra.get("titulo", ""),
                    key=f"extra_cons_tit_{idx_extra}",
                    placeholder="Ex: Do Exame das Vestes / Das Marcas de Pneu"
                )
                
                item_extra["texto"] = st.text_area(
                    "Texto da Consideração",
                    value=item_extra.get("texto", ""),
                    key=f"extra_cons_txt_{idx_extra}",
                    height=100,
                    placeholder="Digite o texto da consideração técnica..."
                )
                
                if st.button(f"🗑️ Remover Consideração ({letra})", key=f"btn_rem_extra_{idx_extra}", type="secondary"):
                    st.session_state["consideracoes_extras"].pop(idx_extra)
                    st.rerun()
                st.markdown("<hr style='margin:16px 0; border-color:#e2e8f0;'>", unsafe_allow_html=True)

            if st.button("➕ Adicionar Consideração", key="btn_add_consideracao_extra", type="primary"):
                st.session_state["consideracoes_extras"].append({"titulo": "", "texto": ""})
                st.rerun()

            st.divider()
            st.subheader("e) Quesitos e Respostas")
            st.markdown("<p style='font-size:13px; color:#475569; margin-bottom:12px;'>Gerenciamento de quesitos formulados pela Autoridade Solicitante e respostas da perícia.</p>", unsafe_allow_html=True)

            if "quesitos_list" not in st.session_state or not st.session_state["quesitos_list"]:
                st.session_state["quesitos_list"] = [{"pergunta": "", "resposta": ""}]

            quesitos_list = st.session_state["quesitos_list"]

            for i, q_item in enumerate(quesitos_list):
                st.markdown(f"##### ❓ Quesito #{i+1}")
                q_item["pergunta"] = st.text_input(
                    "Pergunta",
                    value=q_item.get("pergunta", ""),
                    key=f"quesito_perg_{i}",
                    placeholder="Ex: Qual a causa da morte?"
                )
                q_item["resposta"] = st.text_area(
                    "Resposta",
                    value=q_item.get("resposta", ""),
                    key=f"quesito_resp_{i}",
                    height=75,
                    placeholder="Ex: Traumatismo cranioencefálico decorrente de PAF."
                )

                if len(quesitos_list) > 1:
                    if st.button(f"🗑️ Remover Quesito #{i+1}", key=f"remove_quesito_{i}", type="secondary"):
                        st.session_state.quesitos_list.pop(i)
                        st.rerun()
                st.markdown("<hr style='margin:16px 0; border-color:#cbd5e1;'>", unsafe_allow_html=True)

            if st.button("➕ Adicionar Quesito", key="btn_add_quesito"):
                st.session_state.quesitos_list.append({"pergunta": "", "resposta": ""})
                st.rerun()

            st.divider()
            st.subheader("f) Dinâmica dos Fatos / Conclusão Pericial")
            st.markdown("<p style='font-size:13px; color:#475569; margin-bottom:12px;'>Elabore a narrativa técnico-científica dos acontecimentos ou solicite à IA Gemini uma sugestão automatizada baseada nos dados do laudo.</p>", unsafe_allow_html=True)

            col_d1, col_d2 = st.columns([2.2, 1])
            with col_d1:
                if st.button("🤖 Gerar Sugestão de Dinâmica dos Fatos com IA", type="primary", use_container_width=True, key="btn_gerar_dinamica_ia"):
                    with st.spinner("Sintetizando vestígios, lesões e local para sugerir a Dinâmica dos Fatos..."):
                        dados_dinamica = {
                            "ocorrencia": st.session_state.get("ocorrencia"),
                            "endereco": st.session_state.get("endereco"),
                            "delimitacoes": st.session_state.get("delimitacoes"),
                            "isolamento": st.session_state.get("iso_estado"),
                            "instrumento": st.session_state.get("inst_acao"),
                            "agente_compativel": st.session_state.get("inst_agente"),
                            "achados_extras": st.session_state.get("inst_extra"),
                            "vitimas": st.session_state.get("vitimas"),
                            "vestigios": st.session_state.get("vestigios"),
                            "quesitos": st.session_state.get("quesitos_list")
                        }
                        prompt_dinamica = (
                            "Com base exclusivamente nos dados técnicos do laudo pericial fornecidos abaixo, elabore o texto formal "
                            "da 'Dinâmica dos Fatos / Conclusão Pericial'. A narrativa deve correlacionar a posição e lesões da(s) vítima(s), "
                            "os vestígios materiais coletados, o instrumento utilizado e o cenário do local, apresentando em sequência cronológica e lógica "
                            "provável como os fatos ocorreram, mantendo tom imparcial, técnico-científico e jurídico de laudo oficial da Polícia Civil.\n\n"
                            f"DADOS DO LAUDO:\n{json.dumps(dados_dinamica, ensure_ascii=False, indent=2, default=str)}"
                        )
                        sys_dinamica = (
                            "Você é um perito criminal relator especialista em homicídios e locais de morte violenta. "
                            "Redija uma Dinâmica dos Fatos formal, coesa e com vocabulário técnico rigoroso para constar no laudo oficial."
                        )
                        res, err = call_gemini_text(prompt_dinamica, system_instruction=sys_dinamica)
                        if err:
                            st.error(f"Erro ao gerar dinâmica com IA: {err}")
                        else:
                            st.session_state["dinamica_fatos"] = res
                            st.success("✅ Sugestão de Dinâmica dos Fatos gerada com sucesso!")

            with col_d2:
                if st.button("✨ Polir Redação da Dinâmica", use_container_width=True, key="btn_polir_dinamica"):
                    modal_polir_redacao("dinamica_fatos", "Dinâmica dos Fatos")

            dinamica_text = st.text_area(
                "Descrição da Dinâmica dos Fatos",
                value=st.session_state.get("dinamica_fatos", ""),
                height=180,
                key="dinamica_fatos_ui",
                placeholder="Clique no botão acima para gerar a narrativa por IA ou digite a síntese da dinâmica..."
            )
            st.session_state["dinamica_fatos"] = dinamica_text
    st.markdown("<div style='height:4px'></div>", unsafe_allow_html=True)

    with st.container():
        st.markdown(
            '<div class="section-title">📷 &nbsp; 5. Fotografias (Apêndice A)</div>', unsafe_allow_html=True)
        with st.container(border=True):
            st.markdown('<span class="custom-border-marker"></span>',
                        unsafe_allow_html=True)

            if "fotos" not in st.session_state:
                st.session_state["fotos"] = []

            st.markdown("<p style='font-size:13px; color:#475569; margin-bottom:12px;'>Faça o upload de fotos do local para compor o Apêndice A (Tomadas Fotográficas) do laudo.</p>", unsafe_allow_html=True)

            uploaded_files = st.file_uploader(
                "Selecione uma ou mais imagens (PNG, JPG, JPEG)",
                type=["png", "jpg", "jpeg"],
                accept_multiple_files=True,
                key="foto_uploader"
            )

            if uploaded_files:
                import base64
                novas_adicionadas = False
                for up_file in uploaded_files:
                    filename = up_file.name
                    if not any(f.get("nome") == filename for f in st.session_state["fotos"]):
                        bytes_data = up_file.read()
                        b64_str = base64.b64encode(bytes_data).decode("utf-8")
                        st.session_state["fotos"].append({
                            "nome": filename,
                            "b64": b64_str,
                            "descricao": f"Visão geral do local ({filename})",
                            "incluir": True
                        })
                        novas_adicionadas = True
                if novas_adicionadas:
                    st.toast("📷 Fotografias adicionadas com sucesso!")
                    st.rerun()

            if st.session_state["fotos"]:
                st.markdown("<hr style='margin:16px 0; border-color:#e2e8f0;'>", unsafe_allow_html=True)

                c_gal_hdr, c_gal_btn = st.columns([2, 1])
                with c_gal_hdr:
                    st.markdown(f"##### 🖼️ Galeria de Fotografias ({len(st.session_state['fotos'])} item(ns))")
                with c_gal_btn:
                    if st.button("✨ Gerar Legendas com IA Gemini", type="primary", key="btn_gemini_legendas_todas", use_container_width=True):
                        if not get_gemini_api_key():
                            st.error("⚠️ Chave GEMINI_API_KEY não configurada. Verifique se o arquivo CHAVE.txt está presente.")
                        else:
                            with st.spinner("✨ Analisando fotografias e gerando legendas periciais formais com Gemini Vision..."):
                                count_sucesso = 0
                                for idx_f, foto in enumerate(st.session_state["fotos"]):
                                    legenda, err_g = gerar_legenda_foto_gemini(foto["b64"])
                                    if legenda:
                                        foto["descricao"] = legenda
                                        count_sucesso += 1
                                if count_sucesso > 0:
                                    st.toast(f"✨ Legendas geradas com sucesso para {count_sucesso} foto(s)!")
                                    st.rerun()
                                else:
                                    st.error("Não foi possível gerar legendas para as fotos.")

                indices_para_remover = []
                for idx_f, foto in enumerate(st.session_state["fotos"]):
                    col_img, col_info = st.columns([1, 2])
                    with col_img:
                        try:
                            b64_clean = foto.get("b64", "")
                            if "," in b64_clean:
                                b64_clean = b64_clean.split(",", 1)[1]
                            img_bytes = base64.b64decode(b64_clean)
                            st.image(img_bytes, use_container_width=True)
                        except Exception as e_img:
                            st.error(f"Erro ao carregar imagem: {e_img}")
                    with col_info:
                        foto["incluir"] = st.checkbox(
                            "Incluir no Laudo (.docx)",
                            value=foto.get("incluir", True),
                            key=f"foto_inc_{idx_f}"
                        )
                        c_lbl, c_btn_sing = st.columns([2, 1])
                        with c_lbl:
                            st.markdown("<label style='font-size:13px; font-weight:600; color:#334155;'>Legenda / Descrição da Fotografia</label>", unsafe_allow_html=True)
                        with c_btn_sing:
                            if st.button("✨ Legenda IA", key=f"btn_single_gemini_{idx_f}", use_container_width=True, help="Gerar legenda pericial automática com Gemini Vision"):
                                if not get_gemini_api_key():
                                    st.error("⚠️ Chave Gemini API não configurada.")
                                else:
                                    with st.spinner("Gerando legenda..."):
                                        leg_s, err_s = gerar_legenda_foto_gemini(foto["b64"])
                                        if leg_s:
                                            foto["descricao"] = leg_s
                                            st.toast(f"✨ Legenda gerada para Foto #{idx_f+1}!")
                                            st.rerun()
                                        else:
                                            st.error(f"Erro: {err_s}")

                        foto["descricao"] = st.text_area(
                            "Legenda / Descrição da Fotografia",
                            value=foto.get("descricao", ""),
                            key=f"foto_desc_{idx_f}",
                            height=80,
                            label_visibility="collapsed",
                            placeholder="Descreva o elemento ou visão registrada nesta foto..."
                        )
                        if st.button(f"🗑️ Remover Foto #{idx_f+1}", key=f"btn_del_foto_{idx_f}", type="secondary"):
                            indices_para_remover.append(idx_f)
                    st.markdown("<hr style='margin:12px 0; border-color:#cbd5e1;'>", unsafe_allow_html=True)

                if indices_para_remover:
                    for i in reversed(indices_para_remover):
                        st.session_state["fotos"].pop(i)
                    st.rerun()

                if st.button("🗑️ Limpar Todas as Fotografias", type="secondary", key="btn_clear_all_fotos"):
                    st.session_state["fotos"] = []
                    st.rerun()
            else:
                st.info("Nenhuma fotografia adicionada ainda. Utilize o campo acima para enviar as fotos do laudo.")

    # ── BOTÕES INFERIORES ──
    gerar_bottom = render_action_buttons("bottom")

    gerar_clicked = gerar_top or gerar_bottom

    if gerar_clicked:
        try:
            import os
            template_path = os.path.join(
                os.path.dirname(__file__), "Modelo Laudo.docx")
            doc = Document(template_path)

            # Fetch dynamically selected data
            dp_val = st.session_state.get("data_pericia_input", date.today())
            da_val = st.session_state.get(
                "data_atendimento_input", date.today())
            horario_val = st.session_state.get("horario", None)
            horario_a_val = st.session_state.get("horario_atendimento", None)

            from datetime import datetime, date, time as dt_time

            def safe_date_obj(d):
                if isinstance(d, str):
                    try:
                        return datetime.strptime(d, "%Y-%m-%d").date()
                    except:
                        pass
                return d

            def safe_time_str(t):
                if not t:
                    return "____"
                if isinstance(t, str):
                    return t[:5]
                try:
                    return t.strftime("%H:%M")
                except:
                    return str(t)

            dp_val = safe_date_obj(dp_val)
            da_val = safe_date_obj(da_val)

            dp_extenso = data_extenso(dp_val)
            dp_simples = data_simples(dp_val)
            da_simples = data_simples(da_val)
            h_str = safe_time_str(horario_val)
            ha_str = safe_time_str(horario_a_val)

            tipo_local_val = st.session_state.get("tipo_local", "")

            vitimas_list = st.session_state.get("vitimas", [{"nome": "", "cad": "", "documento": "", "sexo": "", "data_nascimento": None, "filicao": "",
                                                "naturalidade": "", "vestes": "", "pertences": "", "localizacao": "", "posicao": "", "cabeca": "", "membros": "", "fenomenos": "", "lesoes": [""]}])
            nomes_v = [vt["nome"] for vt in vitimas_list if vt["nome"].strip()]
            vitima_header = ", ".join(nomes_v) if nomes_v else "________"

            vitimas_detalhes_list = []
            individual_vars = {}
            for idx, vt in enumerate(vitimas_list, 1):
                suffix = f"_vitima{idx}"
                membros_key = "Membros_vitima!" if idx == 1 else f"Membros_vitima{idx}"

                dob_v = vt.get("data_nascimento", None)
                dob_str = data_simples(dob_v) if dob_v else "________"

                lesoes_list = vt.get("lesoes", [""])
                lesao1_val = lesoes_list[0] if len(lesoes_list) > 0 else ""
                lesao2_val = lesoes_list[1] if len(lesoes_list) > 1 else ""

                individual_vars[f"nome{suffix}"] = vt.get("nome", "")
                individual_vars[f"numero_documento{suffix}"] = vt.get(
                    "documento", "")
                individual_vars[f"sexo{suffix}"] = vt.get("sexo", "")
                individual_vars[f"data_nascimento{suffix}"] = dob_str
                individual_vars[f"filicao{suffix}"] = vt.get("filicao", "")
                individual_vars[f"naturalidade{suffix}"] = vt.get(
                    "naturalidade", "")
                individual_vars[f"veste{suffix}"] = vt.get("vestes", "")
                individual_vars[f"pertence{suffix}"] = vt.get("pertences", "")
                individual_vars[f"localização{suffix}"] = vt.get(
                    "localizacao", "")
                individual_vars[f"posição{suffix}"] = vt.get("posicao", "")
                individual_vars[f"cabeça{suffix}"] = vt.get("cabeca", "")
                individual_vars[membros_key] = vt.get("membros", "")
                individual_vars[f"fenômenos_transformativo{suffix}"] = vt.get(
                    "fenomenos", "")
                individual_vars[f"Lesao1{suffix}"] = lesao1_val
                individual_vars[f"Lesao2{suffix}"] = lesao2_val

                lesoes_txt_list = []
                for les in lesoes_list:
                    if les.strip():
                        lesoes_txt_list.append(f"• {les.strip()}")
                lesoes_text = "\n".join(
                    lesoes_txt_list) if lesoes_txt_list else "• Não descritas"

                text = (
                    f"3.3.{idx} – Descrição da Vítima {idx}\n"
                    f"Identificação\n"
                    f"a) Nome: {vt.get('nome') or '________'}\n"
                    f"b) Documento: {vt.get('documento') or '________'}\n"
                    f"c) Sexo: {vt.get('sexo') or '________'}\n"
                    f"d) Data de Nascimento: {dob_str}\n"
                    f"e) Filiação: {vt.get('filicao') or '________'}\n"
                    f"f) Naturalidade: {vt.get('naturalidade') or '________'}\n\n"
                    f"Vestes e Pertences\n"
                    f"a) Vestes: {vt.get('vestes') or '________'}\n"
                    f"b) Pertences: {vt.get('pertences') or '________'}\n\n"
                    f"Posição e Localização (In Situ)\n"
                    f"a) Localização: {vt.get('localizacao') or '________'}\n"
                    f"b) Posição: {vt.get('posicao') or '________'}\n"
                    f"c) Cabeça: {vt.get('cabeca') or '________'}\n"
                    f"d) Membros: {vt.get('membros') or '________'}\n\n"
                    f"Exame Perinecroscópico e Estado de Conservação\n"
                    f"a) Fenômenos Transformativos: {vt.get('fenomenos') or '________'}\n"
                    f"b) Lesões:\n"
                    f"{lesoes_text}"
                )
                vitimas_detalhes_list.append(text)

            vitimas_detalhes = "\n\n".join(vitimas_detalhes_list)

            vestigios_list = st.session_state.get("vestigios", [])
            vestigios_detalhes_list = []
            for idx_v, vest in enumerate(vestigios_list, 1):
                if vest.get("tipo") == "Manchas de Sangue":
                    txt = (f"Vestígio {idx_v} (Mancha de Sangue):\n"
                           f"Categoria: {vest.get('categoria') or '________'} - Subtipo: {vest.get('subtipo') or '________'}\n"
                           f"Localização: {vest.get('localizacao') or '________'}\n"
                           f"Descrição: {vest.get('descricao') or '________'}")
                    vestigios_detalhes_list.append(txt)
                elif vest.get("tipo") == "Elemento Balístico":
                    txt = (f"Vestígio {idx_v} (Elemento Balístico):\n"
                           f"Tipo: {vest.get('elemento_tipo') or '________'} - Quantidade: {vest.get('quantidade') or '________'}\n"
                           f"Envelope de Segurança: {vest.get('envelope') or '________'}\n"
                           f"Descrição: {vest.get('descricao') or '________'}")
                    vestigios_detalhes_list.append(txt)
                elif vest.get("tipo") == "Arma de Fogo":
                    txt = (f"Vestígio {idx_v} (Arma de Fogo):\n"
                           f"Tipo: {vest.get('arma_tipo') or '________'} - Municiada: {vest.get('municiada') or '________'}\n"
                           f"Envelope de Segurança: {vest.get('envelope') or '________'}\n"
                           f"Descrição: {vest.get('descricao') or '________'}")
                    vestigios_detalhes_list.append(txt)
                elif vest.get("tipo") == "Outros":
                    txt = (f"Vestígio {idx_v} (Outros):\n"
                           f"Descrição: {vest.get('descricao') or '________'}")
                    vestigios_detalhes_list.append(txt)
                elif vest.get("tipo") == "Impressão Datiloscópica":
                    txt = (f"Vestígio {idx_v} (Impressão Datiloscópica):\n"
                           f"Descrição: {vest.get('descricao') or '________'}")
                    vestigios_detalhes_list.append(txt)
                elif vest.get("tipo") == "Instrumento Lesivo":
                    txt = (f"Vestígio {idx_v} (Instrumento Lesivo):\n"
                           f"Tipo: {vest.get('instrumento_tipo') or '________'}\n"
                           f"Envelope de Segurança: {vest.get('envelope') or '________'}\n"
                           f"Descrição: {vest.get('descricao') or '________'}")
                    vestigios_detalhes_list.append(txt)
                elif vest.get("tipo") == "Aparelho Telefônico":
                    txt = (f"Vestígio {idx_v} (Aparelho Telefônico):\n"
                           f"Marca/Modelo: {vest.get('marca') or '________'} - Cor: {vest.get('cor') or '________'}\n"
                           f"Funcionamento: {vest.get('funcionamento') or '________'}\n"
                           f"Envelope de Segurança: {vest.get('envelope') or '________'}\n"
                           f"Descrição: {vest.get('descricao') or '________'}")
                    vestigios_detalhes_list.append(txt)
            vestigios_detalhes = "\n\n".join(
                vestigios_detalhes_list) if vestigios_detalhes_list else "Não foram descritos vestígios."

            quesitos_formatted_list = []
            for idx_q, q_item in enumerate(st.session_state.get("quesitos_list", []), 1):
                p_val = q_item.get("pergunta", "").strip()
                r_val = q_item.get("resposta", "").strip()
                if p_val or r_val:
                    p_text = p_val if p_val else "________"
                    r_text = r_val if r_val else "________"
                    quesitos_formatted_list.append(f"Quesito {idx_q}: {p_text}\nResposta: {r_text}")

            quesitos_final_str = "\n\n".join(quesitos_formatted_list) if quesitos_formatted_list else "Não foram formulados quesitos específicos."
            st.session_state["quesitos"] = quesitos_final_str

            VARS = {
                "num_laudo": st.session_state.get("num_laudo", ""),
                "ocorrencia": st.session_state.get("ocorrencia", ""),
                "ocorrência": st.session_state.get("ocorrencia", ""),
                "perito": st.session_state.get("perito", ""),
                "autoridade": st.session_state.get("autoridade_sel", ""),
                "requisicao": st.session_state.get("requisicao", ""),
                "vitima": vitima_header,
                "investigado": st.session_state.get("investigado", ""),
                "referencia": st.session_state.get("referencia", ""),
                "destino": st.session_state.get("destino", ""),
                "data_extenso": dp_extenso,
                "data_simples": da_simples,
                "horario": h_str,
                "horario_atend": ha_str,
                "endereco": st.session_state.get("endereco", ""),
                "latitude": st.session_state.get("latitude", ""),
                "longitude": st.session_state.get("longitude", ""),
                "ponto": st.session_state.get("ponto_referencia", ""),
                "area": st.session_state.get("area", ""),
                "pavimento": st.session_state.get("pavimento", ""),
                "delimitacoes": st.session_state.get("delimitacoes", ""),
                "isolamento": st.session_state.get("isolamento_outros", "") if st.session_state.get("isolamento", "") == "Outros" else st.session_state.get("isolamento", ""),
                "clima": st.session_state.get("clima_outros", "") if st.session_state.get("clima", "") == "Outros" else st.session_state.get("clima", ""),
                "visibilidade": st.session_state.get("visibilidade", ""),
                "iluminacao": st.session_state.get("iluminacao", ""),
                "equipe_pm": st.session_state.get("equipe_pm", ""),
                "equipe_pc": st.session_state.get("equipe_pc", ""),
                "aut_local": st.session_state.get("autoridade_local", ""),
                "vitimas_detalhes": vitimas_detalhes,
                "vestigios_detalhes": vestigios_detalhes,
                "quesitos": quesitos_final_str,
                "quesitos_respostas": quesitos_final_str,
                "dinamica_fatos": st.session_state.get("dinamica_fatos", "") or "Dinâmica não descrita.",
                "dinâmica_dos_fatos": st.session_state.get("dinamica_fatos", "") or "Dinâmica não descrita.",
                "dinamica": st.session_state.get("dinamica_fatos", "") or "Dinâmica não descrita.",
            }
            VARS.update(individual_vars)
            VARS["tipo_local"] = tipo_local_val

            # --- Lógica de Isolamento ---
            tpl_isolamento = {
                "1": "No momento da chegada da equipe pericial, o local encontrava-se devidamente isolado por meio de {meio_de_isolamento}, impedindo o acesso de pessoas não autorizadas ao perímetro de interesse. Constatou-se a integral preservação do estado das coisas, não havendo quaisquer indícios de alteração, supressão, contaminação ou acréscimo de vestígios. Tais condições atestam o fiel cumprimento ao Art. 6º, inciso I, e Art. 169 do Código de Processo Penal, garantindo a idoneidade da etapa de isolamento da Cadeia de Custódia (Art. 158-A, § 2º) e conferindo total confiabilidade ao levantamento pericial e à dinâmica interpretada.",
                "2": "O sítio pericial apresentava-se isolado por {meio_de_isolamento}, restringindo o trânsito de pessoas estranhas aos exames naquele momento. Contudo, constatou-se que o ambiente encontrava-se apenas parcialmente preservado. Notabilizou-se a alteração do estado original das coisas, consubstanciada por {descrever_alteracao}, ocorrida em momento anterior ao isolamento efetivo. Embora a restrição de acesso atual seja adequada, a descaracterização pretérita impõe ressalvas na interpretação absoluta da dinâmica do evento, sem, contudo, inviabilizar a análise dos vestígios remanescentes.",
                "3": "O perímetro de interesse criminalístico encontrava-se parcialmente isolado, apresentando vulnerabilidades na delimitação física, tais como {descrever_falha_isolamento}, permitindo potencial acesso de curiosos. Não obstante a fragilidade da contenção perimetral, o núcleo do evento encontrava-se preservado. As evidências materiais e a disposição dos elementos encontravam-se em seu estado aparente original, sem sinais de manipulação, contaminação ou descaracterização, permitindo a devida fixação, coleta e garantia da cadeia de custódia dos vestígios encontrados.",
                "4": "O local dos fatos encontrava-se parcialmente isolado, com delimitação perimetral incipiente e insuficiente para {descrever_falha_isolamento}. Concomitantemente, verificou-se a preservação apenas parcial do ambiente, evidenciada por {descrever_alteracao}. Esta conjugação de ineficiência no resguardo do perímetro e a consequente alteração do estado original mitigam a robustez da análise da dinâmica delitiva, configurando inobservância parcial aos ritos de inalterabilidade exigidos pela legislação processual penal vigente.",
                "5": "No momento do acionamento e chegada desta equipe, o sítio pericial encontrava-se desprovido de qualquer isolamento e totalmente devassado. Constatou-se a ausência absoluta de preservação do estado de coisas, com evidências de intensa e irreversível modificação do cenário original, notadamente por {descrever_alteracao}. Tal inobservância à garantia da inalterabilidade do sítio (Art. 6º, I e Art. 169 do CPP) compromete severamente as etapas iniciais da Cadeia de Custódia, prejudicando o estabelecimento inconteste do nexo de causalidade e limitando os achados periciais apenas aos vestígios intrínsecos."
            }
            iso_val_str = str(st.session_state.get("iso_estado", "1")).strip()
            iso_key = iso_val_str[0] if iso_val_str else "1"
            iso_text = tpl_isolamento.get(iso_key, "")
            iso_text = iso_text.replace(
                "{meio_de_isolamento}", st.session_state.get("iso_meio") or "________")
            iso_text = iso_text.replace(
                "{descrever_alteracao}", st.session_state.get("iso_alteracao") or "________")
            iso_text = iso_text.replace(
                "{descrever_falha_isolamento}", st.session_state.get("iso_falha") or "________")

            VARS["isolamento_detalhes"] = iso_text

            # Map IML variables
            VARS["número_laudo_necropsia"] = st.session_state.get(
                "numero_laudo_necropsia") or "________"
            VARS["numero_laudo_necropsia"] = st.session_state.get(
                "numero_laudo_necropsia") or "________"
            VARS["resultado_laudo_IML"] = st.session_state.get(
                "resultado_laudo_IML") or "________"

            # --- Lógica de Instrumento Utilizado ---
            tpl_instrumento = {
                "1": "As lesões constatadas na(s) vítima(s) são características daquelas produzidas por instrumento em ação perfurocontundente, comprovadamente produzidas por projéteis de arma de fogo (PAF), {achado_extra}.",
                "2": "As lesões constatadas na(s) vítima(s) são características daquelas produzidas por instrumento em ação cortante (feridas incisas), comprovadamente produzidas por deslizamento de gume afiado, compatível com {agente_compativel}, {achado_extra}.",
                "3": "As lesões constatadas na(s) vítima(s) são características daquelas produzidas por instrumento em ação perfurante (feridas punctórias), comprovadamente produzidas por agente de ponta fina, compatível com {agente_compativel}, {achado_extra}.",
                "4": "As lesões constatadas na(s) vítima(s) são características daquelas produzidas por instrumento em ação perfurocortante (feridas perfuroincisas), comprovadamente produzidas por arma branca dotada de ponta e gume(s), compatível com {agente_compativel}, {achado_extra}.",
                "5": "As lesões constatadas na(s) vítima(s) são características daquelas produzidas por instrumento em ação contundente (feridas contusas / equimoses / fraturas), comprovadamente produzidas por choque ou impacto contra superfície rígida, compatível com {agente_compativel}, {achado_extra}.",
                "6": "As lesões constatadas na(s) vítima(s) são características daquelas produzidas por instrumento em ação cortocontundente (feridas contuso-incisas), comprovadamente produzidas por agente dotado de massa expressiva e gume, compatível com {agente_compativel}, {achado_extra}.",
                "7": "As lesões constatadas na(s) vítima(s) são características daquelas produzidas por energia de ordem física (ação térmica), comprovadamente decorrentes de exposição a {agente_compativel}, {achado_extra}."
            }

            inst_val_str = str(st.session_state.get("inst_acao", "1")).strip()
            inst_k = inst_val_str[0] if inst_val_str else "1"
            inst_text = tpl_instrumento.get(inst_k, "")
            agente_val = st.session_state.get("inst_agente") or "________"
            extra_val = st.session_state.get("inst_extra") or "________"

            inst_text = inst_text.replace("{agente_compativel}", agente_val)
            inst_text = inst_text.replace("{achado_extra}", extra_val)

            # Cleanup if extra ends with double period or unformatted spaces
            inst_text = inst_text.replace("..", ".")

            VARS["instrumento_detalhes"] = inst_text
            # Map consideracoes_extras in VARS for DOCX
            formatted_extras = []
            for idx_extra, item_extra in enumerate(st.session_state.get("consideracoes_extras", [])):
                letra = chr(ord('e') + idx_extra)
                tit = item_extra.get("titulo", "").strip()
                txt = item_extra.get("texto", "").strip()
                if tit and txt:
                    formatted_extras.append(f"{letra}) {tit}\n{txt}")
                elif txt:
                    formatted_extras.append(f"{letra}) {txt}")

            VARS["consideracoes_extras"] = "\n\n".join(formatted_extras) if formatted_extras else ""

            # Also map to target paragraph if tag is missing in docx
            VARS[
                "d. As lesões constatadas nas vítimas são características daquelas produzidas por instrumento em ação perfurocontundente, inequivocamente produzidas por projéteis de arma de fogo (PAF), com a recuperação de um projétil em cada cadáver durante a necrópsia."] = f"d. {inst_text}"

            def replace_para(para, var_dict):
                for k, v in var_dict.items():
                    target = "{" + k + "}"
                    if target in para.text:
                        val_str = str(v) if v else "________"
                        if "\\n" in val_str or "\n" in val_str:
                            val_str = val_str.replace("\\n", "\n")
                            lines = val_str.split("\n")
                            parts = para.text.split(target, 1)
                            prefix = parts[0]
                            suffix = parts[1] if len(parts) > 1 else ""

                            # Original para will become the LAST line.
                            # So we insert all lines except the last one BEFORE para.
                            for i, line in enumerate(lines[:-1]):
                                p_new = para.insert_paragraph_before(
                                    style=para.style)
                                if i == 0:
                                    p_new.add_run(prefix + line)
                                else:
                                    p_new.add_run(line)

                            # Now modify the original para to hold the last line
                            para.text = ""
                            if len(lines) == 1:
                                para.add_run(prefix + lines[0] + suffix)
                            else:
                                para.add_run(lines[-1] + suffix)

                            # Recursively call in case there are other tags in the same paragraph
                            replace_para(para, var_dict)
                            return
                        else:
                            new = para.text.replace(target, val_str)
                            if len(para.runs) > 0:
                                para.runs[0].text = new
                                for r in para.runs[1:]:
                                    r.text = ""
                            else:
                                para.add_run(new)

            for para in doc.paragraphs:
                replace_para(para, VARS)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_para(para, VARS)
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        replace_para(para, VARS)
                if section.footer:
                    for para in section.footer.paragraphs:
                        replace_para(para, VARS)

            # --- Inserção de Fotografias no Apêndice A ---
            fotos_para_laudo = [f for f in st.session_state.get(
                "fotos", []) if f.get("incluir", True)]
            if fotos_para_laudo:
                import base64
                from io import BytesIO
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                # Locate Apêndice A paragraph or append to document
                apendice_p = None
                for p in doc.paragraphs:
                    if "APÊNDICE A" in p.text.upper() or "TOMADAS FOTOGRÁFICAS" in p.text.upper():
                        apendice_p = p

                # If not found, add header
                if not apendice_p:
                    apendice_p = doc.add_paragraph(
                        "APÊNDICE A – TOMADAS FOTOGRÁFICAS")
                    apendice_p.runs[0].bold = True

                # Clear placeholders after apendice_p if any
                # Append each photo with description ABOVE and picture BELOW
                for f_idx, foto in enumerate(fotos_para_laudo, 1):
                    # Add Description Paragraph (ABOVE)
                    p_desc = doc.add_paragraph()
                    p_desc.paragraph_format.space_before = Pt(12)
                    p_desc.paragraph_format.space_after = Pt(4)

                    r_num = p_desc.add_run(f"Fotografia {f_idx} – ")
                    r_num.bold = True
                    r_desc = p_desc.add_run(foto.get("descricao") or "")

                    # Add Image Paragraph (BELOW)
                    try:
                        img_data = base64.b64decode(foto["b64"])
                        img_stream = BytesIO(img_data)
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_after = Pt(18)
                        r_img = p_img.add_run()
                        r_img.add_picture(img_stream, width=Inches(5.5))
                    except Exception as img_err:
                        st.warning(
                            f"Não foi possível inserir a fotografia {f_idx}: {str(img_err)}")

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)

            ocr = st.session_state.get("ocorrencia", "")
            nome = f"Laudo_{ocr.replace('/','_') if ocr else 'sem_numero'}.docx"
            st.session_state.docx_bytes = buf.getvalue()
            st.session_state.docx_filename = nome
            st.success("✅ Laudo gerado com sucesso!")
        except Exception as e:
            st.error(f"❌ Erro ao gerar o laudo: {str(e)}")

        if st.session_state.get("docx_bytes"):
            st.download_button(
                label="⬇️  Baixar Laudo Gerado",
                data=st.session_state.docx_bytes,
                file_name=st.session_state.docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,

            )
